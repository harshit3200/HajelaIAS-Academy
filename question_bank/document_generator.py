from django.shortcuts import render, redirect, get_object_or_404
from django.core.files.storage import FileSystemStorage
from django.contrib import messages
import pandas as pd
from .models import QuestionBank, InputSuggestion, InputSuggestionImage, InputSuggestionDocument, Area, PartName, ChapterName, TopicName, QuoteIdiomPhrase
from django.db.models import Max, Count, Value, Q
from .forms import UploadFileForm, InputSuggestionForm
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image as PILImage
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from django.http import FileResponse, HttpResponse
from django.contrib.auth.decorators import login_required
from io import BytesIO
from django.contrib.auth import get_user_model
from django.conf import settings
import re


User = get_user_model()

def clean_text(text):
    """Utility function to clean and format text by removing extra newlines and spaces."""
    if not text:
        return ''
    # Strip leading and trailing spaces
    text = text.strip()
    # Replace multiple newlines or newline + spaces with a single space
    text = re.sub(r'\s*\n\s*', ' ', text)
    # Ensure no multiple spaces remain after replacements
    text = re.sub(r'\s+', ' ', text)
    return text

def set_no_border(cell):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tcBorders.append(border)
    tcPr.append(tcBorders)

def generate_filtered_questions_document(request, filters):
    try:
        # Setup directory and document file to save the generated Word file
        base_dir = os.path.join(settings.MEDIA_ROOT, 'word_file')
        os.makedirs(base_dir, exist_ok=True)
        today = datetime.today().strftime('%Y-%m-%d')
        file_name = f'class_plus_filtered_questions_{today}.docx'
        file_path = os.path.join(base_dir, file_name)
        document = Document()

        # Filter questions based on user input
        questions = QuestionBank.objects.filter(**filters)

        # Generate the Word document
        for question in questions:
            table = document.add_table(rows=0, cols=3)
            table.style = 'Table Grid'

            question_text = clean_text(question.question_part or question.question_part_first or '')

            q_row = table.add_row().cells
            q_row[0].text = 'Question'
            q_row[1].text = question_text
            q_row[1].merge(q_row[2])

            if question.image:
                image_path = question.image.path
                pil_img = PILImage.open(image_path)

                if pil_img.mode == 'RGBA':
                    pil_img = pil_img.convert('RGB')

                img_io = BytesIO()
                pil_img.save(img_io, format='JPEG')
                img_io.seek(0)

                paragraph = q_row[1].add_paragraph()
                run = paragraph.add_run()
                run.add_picture(img_io, width=Inches(1.5))

            marks_row = table.add_row().cells
            marks_row[0].text = 'Marks'
            marks_row[1].text = str(question.marks)
            marks_row[2].text = str(question.negative_marks)

            document.add_paragraph()

        document.save(file_path)

        response = FileResponse(open(file_path, 'rb'), as_attachment=True, filename=file_name)
        return response

    except Exception as e:
        return HttpResponse(f"An error occurred: {str(e)}", status=500)
