import sys
sys.stdout.reconfigure(encoding='utf-8')
sys.stderr.reconfigure(encoding='utf-8')

from django.shortcuts import render, redirect, get_object_or_404
from django.core.files.storage import FileSystemStorage
from django.contrib import messages
import pandas as pd
from .models import QuestionBank,InputSuggestion,InputSuggestionImage, InputSuggestionDocument, Area, PartName, ChapterName,TopicName, QuoteIdiomPhrase
from django.db.models import Max, Count, Value, Q
from .forms import UploadFileForm, InputSuggestionForm, QuestionFilterForm
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image as PILImage
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from django.http import FileResponse, HttpResponse
from PIL import Image as PILImage
from django.http import JsonResponse
from django.contrib.auth.decorators import login_required
import json
from django.utils import timezone
from datetime import timedelta
from django.db import transaction
from accounts.models import User
import csv
from io import BytesIO
from django.http import HttpResponse
from .models import Report
from django.template.loader import render_to_string
from django.db.models.functions import Concat
from django.contrib.auth import get_user_model
from django.conf import settings
from .document_generator import generate_filtered_questions_document  # Assuming you have a separate file for document generation
User = get_user_model()


# ************************* Generate Test Word file End *********************************************

# Clean text utility
def clean_text(text):
    """Utility function to clean and format text by removing extra newlines and spaces."""
    if not text:
        return ''
    # Strip leading and trailing spaces
    text = text.strip()
    # Replace multiple newlines or newline + spaces with a single space
    text = re.sub(r'\s*\n\s*', ' ', text)
    # Ensure no multiple spaces remain after replacements
    text = re.sub(r'\s+', ' ', text)
    return text

# Generate Word file
def generate_questions(request):
    try:
        buffer = BytesIO()  # In-memory buffer to store the document
        today = datetime.today().strftime('%Y-%m-%d')
        file_name = f'all_questions_{today}.docx'
        document = Document()

        for question in QuestionBank.objects.all():
            if question.question_sub_type == 'simple_type':
                add_simple_type(question, document)
            elif question.question_sub_type == 'r_and_a_type':
                add_r_and_a_type(question, document)
            elif question.question_sub_type == 'list_type_1':
                add_list_type_1(question, document)
            elif question.question_sub_type == 'list_type_2':
                add_list_type_2(question, document)
            
            # Add a space between questions
            document.add_paragraph()

        # Save the document to the in-memory file object
        document.save(buffer)
        buffer.seek(0)

        # Return the generated file as a downloadable response
        response = FileResponse(buffer, as_attachment=True, filename=file_name)
        response['Content-Disposition'] = f'attachment; filename={file_name}'
        return response

    except Exception as e:
        return HttpResponse(f"An error occurred: {str(e)}", status=500)

# Add question types
def add_simple_type(question, document):
    """Add simple type question to the document."""
    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    # Left column: Question and options
    left_col = table.cell(0, 0)
    left_col.text = f"({question.question_number}). {clean_text(question.question_part)}"
    add_options(left_col, question)

    # Right column: Solution, marks, etc.
    right_col = table.cell(0, 1)
    add_question_details(right_col, question)

def add_r_and_a_type(question, document):
    """Add reason and assertion type question to the document."""
    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    # Left column: Assertion and reason
    left_col = table.cell(0, 0)
    left_col.text = f"({question.question_number}). {clean_text(question.question_part_first)}"
    left_col.text += f"\n{clean_text(question.assertion)}"
    left_col.text += f"\n{clean_text(question.reason)}"
    left_col.text += f"\n{clean_text(question.question_part_third)}"
    add_options(left_col, question)

    # Right column: Solution, marks, etc.
    right_col = table.cell(0, 1)
    add_question_details(right_col, question)

def add_list_type_1(question, document):
    """Add list type 1 question to the document."""
    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    # Left column: Question and list options
    left_col = table.cell(0, 0)
    left_col.text = f"({question.question_number}). {clean_text(question.question_part_first)}"
    
    for i in range(1, 9):  # 8 list rows
        list_row = getattr(question, f'list_1_row{i}', None)
        if list_row:
            left_col.text += f"\n{i}. {clean_text(list_row)}"
    
    if question.question_part_third:
        left_col.text += f"\n\n{clean_text(question.question_part_third)}"
    add_options(left_col, question)

    # Right column: Solution, marks, etc.
    right_col = table.cell(0, 1)
    add_question_details(right_col, question)

def add_list_type_2(question, document):
    """Add list type 2 question to the document."""
    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    # Left column: Question and list options (2 lists)
    left_col = table.cell(0, 0)
    left_col.text = f"({question.question_number}). {clean_text(question.question_part_first)}"

    # Create another table within the left column for LIST-I and LIST-II
    inner_table = left_col.add_table(rows=1, cols=2)
    inner_table.style = 'Table Grid'
    inner_table.cell(0, 0).text = f"LIST - I ({clean_text(question.list_1_name)})"
    inner_table.cell(0, 1).text = f"LIST - II ({clean_text(question.list_2_name)})"
    
    for i in range(1, 9):
        list_1_item = getattr(question, f'list_1_row{i}', '')
        list_2_item = getattr(question, f'list_2_row{i}', '')
        row = inner_table.add_row().cells
        row[0].text = f"{chr(64+i)}. {clean_text(list_1_item)}"
        row[1].text = f"{i}. {clean_text(list_2_item)}"

    # Right column: Solution, marks, etc.
    right_col = table.cell(0, 1)
    add_question_details(right_col, question)

# Add options to the question
def add_options(cell, question):
    """Add options (A, B, C, D) to a table cell."""
    for opt in ['a', 'b', 'c', 'd']:
        option_text = getattr(question, f'answer_option_{opt}', None)
        if option_text:
            cell.text += f"\n({opt.upper()}) {clean_text(option_text)}"

# Add details like solution, marks, and metadata
def add_question_details(cell, question):
    """Add details like the solution, marks, and metadata with bold headings."""
    # Get a paragraph object to start adding runs with formatting
    paragraph = cell.add_paragraph()

    # Correct Answer
    run = paragraph.add_run('Correct Answer: ')
    run.bold = True
    paragraph.add_run(f"{clean_text(question.correct_answer_choice)}")

    # Solution
    paragraph = cell.add_paragraph()  # Create a new paragraph for each item
    run = paragraph.add_run('Solution: ')
    run.bold = True
    paragraph.add_run(f"{clean_text(question.correct_answer_description)}")

    # Marks
    paragraph = cell.add_paragraph()
    run = paragraph.add_run('Marks: ')
    run.bold = True
    paragraph.add_run(f"{question.marks}")
    
    run = paragraph.add_run('\nNegative Marks: ')
    run.bold = True
    paragraph.add_run(f"{question.negative_marks}")

    # Degree of Difficulty
    paragraph = cell.add_paragraph()
    run = paragraph.add_run('Degree of Difficulty: ')
    run.bold = True
    paragraph.add_run(f"{clean_text(question.degree_of_difficulty)}")

    # Evergreen Index
    paragraph = cell.add_paragraph()
    run = paragraph.add_run('Evergreen Index: ')
    run.bold = True
    paragraph.add_run(f"{question.evergreen_index}")

    # Created At
    created_at_str = question.created_at.strftime('%Y-%m-%d %H:%M:%S')
    paragraph = cell.add_paragraph()
    run = paragraph.add_run('Created At: ')
    run.bold = True
    paragraph.add_run(f"{created_at_str}")

    # Created By
    if question.created_by:
        created_by_str = question.created_by.get_full_name() or question.created_by.username
        paragraph = cell.add_paragraph()
        run = paragraph.add_run('Created By: ')
        run.bold = True
        paragraph.add_run(f"{clean_text(created_by_str)}")


# # ************************* Generate Test Word file End *********************************************


# ************************* Generate Clas Plus Word file Start *********************************************
from django.shortcuts import render, redirect
from django.http import JsonResponse
from django.contrib.auth.decorators import login_required
from django.template.loader import render_to_string
from django.contrib import messages
from .models import Area, Section, PartName, ChapterName, TopicName, SubTopicName, QuestionBank


# *********************** Dynamic Data Fetching Views *********************** #

# Function to dynamically fetch areas based on selected subjects
def get_areas_list(request):
    subject_ids = request.GET.getlist('subject_ids[]')
    if not subject_ids:
        return JsonResponse({'areas': []})  # Return an empty list if no subject IDs are provided
    areas = Area.objects.filter(subject_name__id__in=subject_ids).values('id', 'name')
    return JsonResponse({'areas': list(areas)})

# Function to dynamically fetch sections based on selected areas
def get_sections_list(request):
    area_ids = request.GET.getlist('area_ids[]')
    if not area_ids:
        return JsonResponse({'sections': []})  # Return an empty list if no area IDs are provided
    sections = Section.objects.filter(area__id__in=area_ids).values('id', 'name')
    return JsonResponse({'sections': list(sections)})

# Function to dynamically fetch parts based on selected sections
def get_parts_list(request):
    section_ids = request.GET.getlist('section_ids[]')
    if not section_ids:
        return JsonResponse({'parts': []})  # Return an empty list if no section IDs are provided
    parts = PartName.objects.filter(section__id__in=section_ids).values('id', 'name')
    return JsonResponse({'parts': list(parts)})

# Function to dynamically fetch chapters based on selected parts
def get_chapters_list(request):
    part_ids = request.GET.getlist('part_ids[]')
    if not part_ids:
        return JsonResponse({'chapters': []})  # Return an empty list if no part IDs are provided
    chapters = ChapterName.objects.filter(part__id__in=part_ids).values('id', 'name')
    return JsonResponse({'chapters': list(chapters)})

# Function to dynamically fetch topics based on selected chapters
def get_topics_list(request):
    chapter_ids = request.GET.getlist('chapter_ids[]')
    if not chapter_ids:
        return JsonResponse({'topics': []})  # Return an empty list if no chapter IDs are provided
    topics = TopicName.objects.filter(chapter__id__in=chapter_ids).values('id', 'name')
    return JsonResponse({'topics': list(topics)})

# Function to dynamically fetch subtopics based on selected topics
def get_subtopics_list(request):
    topic_ids = request.GET.getlist('topic_ids[]')
    if not topic_ids:
        return JsonResponse({'subtopics': []})  # Return an empty list if no topic IDs are provided
    subtopics = SubTopicName.objects.filter(topic__id__in=topic_ids).values('id', 'name')
    return JsonResponse({'subtopics': list(subtopics)})


# ************************** Question Filtering ************************** #


# ************************** Question Filtering ************************** #

from django.shortcuts import render, get_object_or_404
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.template.loader import render_to_string
from django.contrib.auth.decorators import login_required
from decouple import config
from openai import OpenAI

from .models import QuestionBank, Area, Section, PartName, ChapterName, TopicName, SubTopicName

# ✅ Initialize OpenAI client
client = OpenAI(api_key=config("OPENAI_API_KEY"))


# views.py

from collections import defaultdict
from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from .models import (
    QuestionBank, Area, Section, PartName,
    ChapterName, TopicName, SubTopicName
)

@login_required
def question_filter_view(request):
    """
    Filter questions by hierarchy (Area → Section → Part → Chapter → Topic → Subtopic)
    and group them by base_question_id into English/Hindi pairs.
    """

    # ---------- Dropdown data ----------
    areas = Area.objects.all().order_by("area_SI_Code")
    sections = Section.objects.all().order_by("section_Unit_SI")
    parts = PartName.objects.all().order_by("part_serial")
    chapters = ChapterName.objects.all().order_by("chapter_number")
    topics = TopicName.objects.all().order_by("topic_SI_number")
    subtopics = SubTopicName.objects.all().order_by("sub_topic_SI_Number")

    # ---------- Apply filters ----------
    filters = {
        "area_name__area_SI_Code": request.GET.get("area"),
        "section_name__section_Unit_SI": request.GET.get("section"),
        "part_name__part_serial": request.GET.get("part"),
        "chapter_name__chapter_number": request.GET.get("chapter"),
        "topic_name__topic_SI_number": request.GET.get("topic"),
        "subtopic_name__sub_topic_SI_Number": request.GET.get("subtopic"),
        "created_at__date": request.GET.get("created_at"),
        "question_sub_type": request.GET.get("question_sub_type"),
    }

    # Keep only non-empty values
    filters = {k: v for k, v in filters.items() if v}

    questions = QuestionBank.objects.all().select_related("created_by").prefetch_related(
        "area_name", "section_name", "part_name", "chapter_name", "topic_name", "subtopic_name"
    )
    if filters:
        questions = questions.filter(**filters)

    # ---------- Group by base_question_id ----------
    grouped_questions = defaultdict(lambda: {"english": None, "hindi": None})

    for q in questions:
        if not q.base_question_id:
            continue
        if q.language == "e":
            grouped_questions[q.base_question_id]["english"] = q
        elif q.language == "h":
            grouped_questions[q.base_question_id]["hindi"] = q

    # ---------- Context ----------
    context = {
        "grouped_questions": dict(grouped_questions),  # force JSON-serializable
        "areas": areas,
        "sections": sections,
        "parts": parts,
        "chapters": chapters,
        "topics": topics,
        "subtopics": subtopics,
        "question_types": QuestionBank.QUESTION_TYPES,  # ✅ pass all types (incl. statement_type)
    }

    return render(request, "question_bank/question_filter.html", context)


from django.http import JsonResponse
from django.shortcuts import get_object_or_404
from django.views.decorators.csrf import csrf_exempt
from django.template.loader import render_to_string
from openai import OpenAI
import traceback, json, re, random

from .models import (
    QuestionBank, KeywordName, SubTopicName,
    Area, Section, PartName, ChapterName, TopicName
)

client = OpenAI(api_key="YOUR_OPENAI_API_KEY")


@csrf_exempt
def generate_alternate_question(request):
    """
    🔹 Generate bilingual (English + Hindi) alternate questions
    🔹 Auto-map each generated question to subtopics + keywords
    🔹 Return rendered HTML + tagging summary for UI display
    🔹 (Preview-only: does not save until Save button clicked)
    """
    if request.method != "POST":
        return JsonResponse({"status": "error", "message": "Invalid request method"}, status=400)

    question_id = request.POST.get("question_id")
    if not question_id:
        return JsonResponse({"status": "error", "message": "Missing question ID"}, status=400)

    try:
        # ==========================================================
        # 1️⃣ Fetch source question
        # ==========================================================
        question = get_object_or_404(QuestionBank, id=question_id)
        subtype = question.question_sub_type

        # ==========================================================
        # 2️⃣ Build source text
        # ==========================================================
        def build_source(q):
            try:
                return f"""
Question: {q.question_part_first or q.question_part or ''}
Options:
A. {q.answer_option_a}
B. {q.answer_option_b}
C. {q.answer_option_c}
D. {q.answer_option_d}
Correct Answer: {q.correct_answer_choice}
Explanation: {q.correct_answer_description or ''}
"""
            except Exception:
                return q.question_part_first or ""
        source_text = build_source(question)

        # ==========================================================
        # 3️⃣ Prompt for alternate question generation
        # ==========================================================
        bilingual_prompt = f"""
You are an expert UPSC/MPPSC bilingual question creator.

SOURCE QUESTION:
{source_text}

TASK:
1. Generate 4 alternate questions in ENGLISH.
   - Start with: ### ENGLISH
   - Separate each question with ---.

2. Then generate the SAME 4 alternate questions in HINDI.
   - Start with: ### HINDI
   - Separate each question with ---.
"""

        completion = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "Follow the given bilingual structure exactly."},
                {"role": "user", "content": bilingual_prompt},
            ],
            temperature=0.7,
            max_tokens=2000,
        )
        ai_text = completion.choices[0].message.content.strip()
        if not ai_text:
            return JsonResponse({"status": "error", "message": "OpenAI returned empty response"}, status=500)

        # ==========================================================
        # 4️⃣ Parse English + Hindi pairs
        # ==========================================================
        english_blocks, hindi_blocks = [], []
        for block in ai_text.split("###"):
            b = block.strip()
            if b.lower().startswith("english"):
                english_blocks = [x.strip() for x in b.replace("ENGLISH", "").split("---") if x.strip()]
            elif b.lower().startswith("hindi"):
                hindi_blocks = [x.strip() for x in b.replace("HINDI", "").split("---") if x.strip()]
        paired_blocks = [
            {"english": english_blocks[i] if i < len(english_blocks) else "",
             "hindi": hindi_blocks[i] if i < len(hindi_blocks) else ""}
            for i in range(max(len(english_blocks), len(hindi_blocks)))
        ]

        # ==========================================================
        # 5️⃣ Load all subtopics once
        # ==========================================================
        all_subtopics = list(
            SubTopicName.objects.select_related("topic__chapter__part__section__area").values(
                "name", "sub_topic_short_Code", "topic__chapter__part__section__area__name"
            )
        )
        subtopic_list = [
            f"{s['topic__chapter__part__section__area__name']} → {s['name']} ({s['sub_topic_short_Code']})"
            for s in all_subtopics
        ][:1709]

        auto_results = []

        # ==========================================================
        # 6️⃣ For each generated question → AI map subtopics + keywords
        # ==========================================================
        for pair in paired_blocks:
            text_for_ai = pair["english"].strip()
            if not text_for_ai:
                continue

            # ------------------- SUBTOPIC MAPPING -------------------
            try:
                mapping_prompt = f"""
You are a UPSC/MPPSC content mapping expert at Hajela’s IAS Academy.

Analyse this question (including options and meaning).
Map it to 2–5 conceptually relevant subtopics from the provided list.

Rules:
- Match conceptually, not by word overlap.
- Include "confidence" (0–1) and "reason" in JSON.
- Output valid JSON only.

Question:
{text_for_ai}

Available Subtopics:
{subtopic_list[:1709]}
"""
                completion = client.chat.completions.create(
                    model="gpt-4o",
                    messages=[
                        {"role": "system", "content": "Output valid JSON only — no markdown."},
                        {"role": "user", "content": mapping_prompt},
                    ],
                    temperature=0.3,
                    max_tokens=1200,
                )

                raw_output = completion.choices[0].message.content.strip()
                cleaned = re.sub(r"```(?:json)?", "", raw_output).replace("```", "").strip()

                try:
                    parsed_json = json.loads(cleaned)
                    if isinstance(parsed_json, dict):
                        matched_subtopics = (
                            parsed_json.get("matches")
                            or parsed_json.get("relevant_subtopics")
                            or parsed_json.get("mapped_subtopics")
                            or parsed_json.get("high_confidence_mappings")
                            or []
                        )
                    else:
                        matched_subtopics = parsed_json
                except Exception:
                    matched_subtopics = []

                verified_matches, seen = [], set()
                for st in matched_subtopics:
                    name_ai, conf, reason = "", 0.9, "AI thematic mapping"
                    if isinstance(st, dict):
                        name_ai = st.get("subtopic", "").strip().lower()
                        conf = float(st.get("confidence", 0.9))
                        reason = st.get("reason", reason)
                    elif isinstance(st, str):
                        name_ai = st.strip().lower()

                    name_ai = re.sub(r"→", " ", name_ai)
                    name_ai = re.sub(r"\(.*?\)", "", name_ai)
                    name_ai = re.sub(
                        r"\b(history|geography|economy|polity|environment|science|gs|general studies|other)\b",
                        "",
                        name_ai,
                        flags=re.IGNORECASE,
                    ).strip()

                    if not name_ai or name_ai in seen:
                        continue
                    seen.add(name_ai)

                    sub_obj = SubTopicName.objects.filter(name__icontains=name_ai).select_related(
                        "topic__chapter__part__section__area"
                    ).first()
                    if not sub_obj:
                        continue

                    topic = sub_obj.topic
                    chapter = topic.chapter if topic else None
                    part = chapter.part if chapter else None
                    section = part.section if part else None
                    area = section.area if section else None

                    verified_matches.append({
                        "area": getattr(area, "name", "—"),
                        "area_color": getattr(area, "area_Colour_Hex", "#555"),
                        "section": getattr(section, "name", "—"),
                        "part": getattr(part, "name", "—"),
                        "chapter": getattr(chapter, "name", "—"),
                        "topic": getattr(topic, "name", "—"),
                        "subtopic": sub_obj.name,
                        "code": sub_obj.sub_topic_short_Code,
                        "reason": reason,
                        "confidence": round(conf, 2),
                    })

                if not verified_matches:
                    verified_matches.append({
                        "area": "General Studies",
                        "area_color": "#777",
                        "subtopic": "General Studies",
                        "code": "GEN",
                        "reason": "No match found — fallback used.",
                        "confidence": 0.5,
                    })

            except Exception as inner_e:
                print("⚠️ Subtopic mapping error:", inner_e)
                verified_matches = [{
                    "area": "General Studies",
                    "area_color": "#777",
                    "subtopic": "General Studies",
                    "code": "GEN",
                    "reason": "Error during AI mapping — fallback.",
                    "confidence": 0.5,
                }]

            # ------------------- KEYWORD GENERATION -------------------
            try:
                keyword_prompt = f"""
Generate 5–10 English and Hindi keywords for this question.

Format:
### ENGLISH
keyword1, keyword2, ...
---
### HINDI
कीवर्ड1, कीवर्ड2, ...

Question:
{text_for_ai}
"""
                kw_ai = client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[
                        {"role": "system", "content": "Return bilingual keywords only."},
                        {"role": "user", "content": keyword_prompt},
                    ],
                    temperature=0.4,
                    max_tokens=300,
                )

                kw_text = kw_ai.choices[0].message.content.strip()
                en_kw, hi_kw = "", ""
                for blk in kw_text.split("###"):
                    blk = blk.strip()
                    if blk.lower().startswith("english"):
                        en_kw = re.sub(r"(?i)english", "", blk).replace("---", "").strip()
                    elif blk.lower().startswith("hindi"):
                        hi_kw = re.sub(r"(?i)hindi", "", blk).replace("---", "").strip()

                en_clean = [k.strip() for k in en_kw.split(",") if k.strip()]
                hi_clean = [k.strip() for k in hi_kw.split(",") if k.strip()]

            except Exception as inner_e:
                print("⚠️ Keyword generation error:", inner_e)
                en_clean, hi_clean = [], []

            auto_results.append({
                "question_preview": text_for_ai,
                "subtopics": verified_matches,
                "keywords_en": en_clean,
                "keywords_hi": hi_clean,
                "db_keywords": [],
            })

        # ==========================================================
        # 7️⃣ Render Final HTML (preview)
        # ==========================================================
        html = render_to_string(
            "partials/ai_que_result_card.html",
            {
                "paired_blocks": paired_blocks,
                "question_id": question.id,
                "auto_tagged": auto_results,
            },
        )

        return JsonResponse({"status": "success", "html": html, "auto_tagged": auto_results})

    except Exception:
        print("❌ generate_alternate_question Error:", traceback.format_exc())
        return JsonResponse({"status": "error", "message": "Internal error during generation"}, status=500)



import re, random, traceback
from django.http import JsonResponse
from django.shortcuts import get_object_or_404
from django.contrib.auth.decorators import login_required
from django.db.models import Q
from .models import QuestionBank, SubTopicName


@login_required
def save_alternate_question(request):
    """
    ✅ Save bilingual alternate questions
    ✅ Auto-sync Area → Section → Part → Chapter → Topic → Subtopic hierarchy
    ✅ Save English + Hindi fields exactly as used in template
    """
    if request.method != "POST":
        return JsonResponse({"status": "error", "message": "Invalid request method"}, status=400)

    ai_text = request.POST.get("content", "").strip()
    original_question_id = request.POST.get("source_id")
    english_keywords = request.POST.get("english_keywords", "").strip()
    hindi_keywords = request.POST.get("hindi_keywords", "").strip()
    selected_subtopics = request.POST.getlist("subtopics[]")

    if not ai_text or not original_question_id:
        return JsonResponse({"status": "error", "message": "Missing data"}, status=400)

    try:
        source_question = get_object_or_404(QuestionBank, id=original_question_id)
        subtype = source_question.question_sub_type

        # Split English/Hindi blocks
        english_blocks, hindi_blocks = [], []
        for chunk in ai_text.split("###"):
            c = chunk.strip()
            if c.lower().startswith("english"):
                english_blocks = [b.strip() for b in c.replace("ENGLISH", "").split("---") if b.strip()]
            elif c.lower().startswith("hindi"):
                hindi_blocks = [b.strip() for b in c.replace("HINDI", "").split("---") if b.strip()]
        if not english_blocks and ai_text:
            english_blocks = [ai_text]

        # Helpers
        def safe_search(pattern, text, flags=0, group=1):
            m = re.search(pattern, text, flags)
            return m.group(group).strip() if m else ""

        def safe_findall(pattern, text, flags=0):
            return [x.strip() for x in re.findall(pattern, text, flags)]

        option_pattern = r"[\(\[]?[a-dA-D][\)\]]?\s*[\.\-–:]?\s*(.+)"

        def has_meaningful_data(data):
            return any(v and str(v).strip() for v in data.values())

        # Parser
        def parse_block(text, lang="e"):
            if not text:
                return {}

            fields = {}
            keywords = safe_search(r"Keywords\s*[:：]\s*(.+)", text, re.IGNORECASE)

            # Normalize synonyms for Hindi/English consistency
            text = (
                text.replace("Explanation", "Solution")
                    .replace("Description", "Solution")
                    .replace("उत्तर व्याख्या", "Solution")
                    .replace("व्याख्या", "Solution")
                    .replace("समाधान", "Solution")
                    .replace("उत्तर:", "Answer:")
            )

            # --- Simple Type ---
            if subtype in ["simple_type", "true_and_false_type", "fill_in_the_blank_type"]:
                q = safe_search(
                    r"(?:Question|Q[\.\-:]|प्रश्न|प्रश्‍न)\s*[:：\-]?\s*(.+?)(?=(?:Options|विकल्प|Answer|उत्तर|Solution|$))",
                    text, re.DOTALL | re.IGNORECASE
                )
                opts = safe_findall(option_pattern, text)
                ans = safe_search(
                    r"(?:Answer|उत्तर)\s*[:：]?\s*\(?([a-dA-Dअ-ड])\)?",
                    text, re.IGNORECASE
                )
                desc = safe_search(
                    r"(?:Solution|Explanation|Description|उत्तर\s*व्याख्या|व्याख्या|समाधान)\s*[:：]?\s*(.+)",
                    text, re.DOTALL | re.IGNORECASE
                )

                fields = {
                    "question_part": q or "",
                    **{f"answer_option_{chr(97+i)}": opts[i] if i < len(opts) else "" for i in range(4)},
                    "correct_answer_choice": ans.lower() if ans else "",
                    "correct_answer_description": desc or "",
                    "key_words": keywords,
                }

            # --- Assertion & Reason ---
            elif subtype == "r_and_a_type":
                assertion = safe_search(r"Assertion\s*[:：]?\s*(.+?)(?=Reason|Options|Answer|उत्तर|$)", text, re.DOTALL)
                reason = safe_search(r"Reason\s*[:：]?\s*(.+?)(?=Options|Answer|उत्तर|$)", text, re.DOTALL)
                opts = safe_findall(option_pattern, text)
                ans = safe_search(r"(?:Answer|उत्तर)\s*[:：]?\s*\(?([a-dA-Dअ-ड])\)?", text, re.IGNORECASE)
                desc = safe_search(r"(?:Solution|Explanation|Description|उत्तर\s*व्याख्या|व्याख्या|समाधान)\s*[:：]?\s*(.+)", text, re.DOTALL)

                fields = {
                    "assertion": assertion,
                    "reason": reason,
                    **{f"answer_option_{chr(97+i)}": opts[i] if i < len(opts) else "" for i in range(4)},
                    "correct_answer_choice": ans.lower() if ans else "",
                    "correct_answer_description": desc or "",
                    "key_words": keywords,
                }

            # --- List Type 1 ---
            elif subtype == "list_type_1":
                q1 = safe_search(r"(?:First\s*Part|पहला\s*भाग)\s*[:：]?\s*(.+)", text)
                q3 = safe_search(r"(?:Third\s*Part|तीसरा\s*भाग)\s*[:：]?\s*(.+)", text)
                items = safe_findall(r"\d+\.\s*(.+)", text)
                opts = safe_findall(option_pattern, text)
                ans = safe_search(r"(?:Answer|उत्तर)\s*[:：]?\s*\(?([a-dA-Dअ-ड])\)?", text, re.IGNORECASE)
                desc = safe_search(r"(?:Solution|Explanation|Description|उत्तर\s*व्याख्या|व्याख्या|समाधान)\s*[:：]?\s*(.+)", text, re.DOTALL)

                fields = {
                    "question_part_first": q1 or "",
                    "question_part_third": q3 or "",
                    "correct_answer_choice": ans.lower() if ans else "",
                    "correct_answer_description": desc or "",
                    "key_words": keywords,
                    **{f"list_1_row{i+1}": items[i] if i < len(items) else "" for i in range(8)},
                    **{f"answer_option_{chr(97+i)}": opts[i] if i < len(opts) else "" for i in range(4)},
                }

            # --- List Type 2 ---
            elif subtype == "list_type_2":
                # Capture question intro part before the tables
                q_intro = safe_search(
                    r"(?:Question|Q[\.\-:]|प्रश्न|प्रश्‍न|Match|मिलान)\s*[:：\-]?\s*(.+?)(?=(?:List[\-–]I|List[\-–]II|Options|विकल्प|Answer|उत्तर|$))",
                    text, re.DOTALL | re.IGNORECASE
                )

                # Capture List-I and List-II separately
                list1 = safe_findall(r"[A-Da-d][\.\)]\s*(.+)", text)
                list2 = safe_findall(r"\d+[\.\)]\s*(.+)", text)

                # Capture options below the lists (ensure not to confuse with list items)
                opts_block = re.split(r"(?:Options|विकल्प)", text, flags=re.IGNORECASE)
                opts_text = opts_block[-1] if len(opts_block) > 1 else text
                opts = safe_findall(option_pattern, opts_text)

                # Correct Answer and Description / Solution
                ans = safe_search(r"(?:Answer|उत्तर)\s*[:：]?\s*\(?([a-dA-Dअ-ड])\)?", text, re.IGNORECASE)
                desc = safe_search(
                    r"(?:Solution|Explanation|Description|उत्तर\s*व्याख्या|व्याख्या|समाधान)\s*[:：]?\s*(.+)",
                    text, re.DOTALL | re.IGNORECASE
                )

                fields = {
                    "question_part_first": q_intro or "",
                    "correct_answer_choice": ans.lower() if ans else "",
                    "correct_answer_description": desc or "",
                    "key_words": keywords,
                    **{f"list_1_row{i+1}": list1[i] if i < len(list1) else "" for i in range(8)},
                    **{f"list_2_row{i+1}": list2[i] if i < len(list2) else "" for i in range(9)},
                    **{f"answer_option_{chr(97+i)}": opts[i] if i < len(opts) else "" for i in range(4)},
                }


            # --- Statement Type ---
            elif subtype == "statement_type":
                q1 = safe_search(r"(?:First\s*Part|पहला\s*भाग)\s*[:：]?\s*(.+)", text)
                stmts = safe_findall(r"(?:Statement|वक्तव्य)\s*\d+\s*[:：]?\s*(.+)", text)
                q3 = safe_search(r"(?:Third\s*Part|तीसरा\s*भाग)\s*[:：]?\s*(.+)", text)
                opts = safe_findall(option_pattern, text)
                ans = safe_search(r"(?:Answer|उत्तर)\s*[:：]?\s*\(?([a-dA-Dअ-ड])\)?", text, re.IGNORECASE)
                desc = safe_search(r"(?:Solution|Explanation|Description|उत्तर\s*व्याख्या|व्याख्या|समाधान)\s*[:：]?\s*(.+)", text, re.DOTALL)

                fields = {
                    "question_part_first": q1 or "",
                    "question_part_third": q3 or "",
                    "correct_answer_choice": ans.lower() if ans else "",
                    "correct_answer_description": desc or "",
                    "key_words": keywords,
                    **{f"stmt_line_row{i+1}": stmts[i] if i < len(stmts) else "" for i in range(9)},
                    **{f"answer_option_{chr(97+i)}": opts[i] if i < len(opts) else "" for i in range(4)},
                }

            # --- If Hindi, mirror with _hi ---
            if lang == "h":
                hi_fields = {}
                for k, v in fields.items():
                    if k == "correct_answer_description":
                        hi_fields["correct_answer_description_hi"] = v
                    elif k in ["question_part_first", "question_part"]:
                        hi_fields["question_part_first_hi"] = v
                    elif k in ["correct_answer_choice", "key_words"]:
                        hi_fields[k] = v
                    else:
                        hi_fields[k + "_hi"] = v
                fields = hi_fields

            return fields

        # Save bilingual pairs
        created_ids, linked_names = [], []
        total_pairs = max(len(english_blocks), len(hindi_blocks))

        for idx in range(total_pairs):
            en_text = english_blocks[idx] if idx < len(english_blocks) else ""
            hi_text = hindi_blocks[idx] if idx < len(hindi_blocks) else en_text
            en_data, hi_data = parse_block(en_text, "e"), parse_block(hi_text, "h")
            if not has_meaningful_data(en_data) and not has_meaningful_data(hi_data):
                continue

            base_id = random.randint(100000, 999999)
            while QuestionBank.objects.filter(base_question_id=base_id).exists():
                base_id = random.randint(100000, 999999)

            last_q = QuestionBank.objects.order_by("-question_number").first()
            next_number = (last_q.question_number + 1) if last_q else 1

            q_en = QuestionBank.objects.create(
                base_question_id=base_id,
                question_number=next_number,
                question_sub_type=subtype,
                language="e",
                type_of_question="moq",
                created_by=request.user,
                degree_of_difficulty="Medium",
                elim_tactics_degree="Normal",
                current_relevance="General",
                key_words_en=english_keywords,
                key_words_hi=hindi_keywords,
                **en_data,
            )

            q_hi = QuestionBank.objects.create(
                base_question_id=base_id,
                question_number=next_number,
                question_sub_type=subtype,
                language="h",
                type_of_question="moq",
                created_by=request.user,
                degree_of_difficulty="Medium",
                elim_tactics_degree="Normal",
                current_relevance="General",
                key_words_en=english_keywords,
                key_words_hi=hindi_keywords,
                **hi_data,
            )

            # Copy hierarchy
            for q in [q_en, q_hi]:
                q.area_name.set(source_question.area_name.all())
                q.section_name.set(source_question.section_name.all())
                q.part_name.set(source_question.part_name.all())
                q.chapter_name.set(source_question.chapter_name.all())
                q.topic_name.set(source_question.topic_name.all())
                q.subtopic_name.set(source_question.subtopic_name.all())
                q.exam_name.set(source_question.exam_name.all())

                if english_keywords:
                    q.key_words = english_keywords
                if hindi_keywords:
                    q.key_words_hi = hindi_keywords

            if selected_subtopics:
                clean_names = []
                for s in selected_subtopics:
                    s_clean = s.split("→")[-1].split("(")[0].strip()
                    if s_clean and s_clean not in clean_names:
                        clean_names.append(s_clean)

                query = Q()
                for n in clean_names:
                    query |= Q(name__iexact=n) | Q(name__icontains=n)

                matched = SubTopicName.objects.filter(query).select_related(
                    "topic__chapter__part__section__area"
                ).distinct()

                for st in matched:
                    for q in [q_en, q_hi]:
                        q.subtopic_name.add(st)
                    linked_names.append(st.name)

                    topic = getattr(st, "topic", None)
                    chapter = getattr(topic, "chapter", None)
                    part = getattr(chapter, "part", None)
                    section = getattr(part, "section", None)
                    area = getattr(section, "area", None)

                    for q in [q_en, q_hi]:
                        if area: q.area_name.add(area)
                        if section: q.section_name.add(section)
                        if part: q.part_name.add(part)
                        if chapter: q.chapter_name.add(chapter)
                        if topic: q.topic_name.add(topic)

            q_en.save()
            q_hi.save()
            created_ids.extend([q_en.id, q_hi.id])

        print(f"✅ Saved {len(created_ids)} questions and {len(linked_names)} subtopics")

        return JsonResponse({
            "status": "success",
            "message": f"✅ {len(created_ids)} Questions saved with {len(linked_names)} Subtopics.",
            "new_ids": created_ids,
            "linked_subtopics": linked_names,
            "subtype": subtype,
        })

    except Exception as e:
        print("❌ save_alternate_question Error:", traceback.format_exc())
        return JsonResponse({"status": "error", "message": str(e)}, status=500)



from django.http import JsonResponse
from django.shortcuts import get_object_or_404
from django.views.decorators.csrf import csrf_exempt
from django.template.loader import render_to_string
from openai import OpenAI
import traceback, json, re

from .models import (
    QuestionBank, KeywordName, SubTopicName,
    Area, Section, PartName, ChapterName, TopicName
)

# ✅ Initialize OpenAI Client
client = OpenAI(api_key="YOUR_OPENAI_API_KEY")

# 🧹 Ignore junk/common words
IGNORED_KEYWORDS = {
    "india", "indian", "delhi", "bihar", "up", "mp", "maharashtra",
    "country", "state", "district", "world", "nation",
    "january", "february", "march", "april", "may", "june",
    "july", "august", "september", "october", "november", "december",
    "none", "all", "both", "each", "year", "century", "data", "time"
}

# ============================================================
# 🔹 Safe Keyword Creation Helper
# ============================================================
def safe_get_or_create_keyword(name, lang="EN", source="ai", score=0.9, subtopics=None):
    """
    Ensures uniqueness (case-insensitive) and safely links to one or more subtopics.
    Supports multiple subtopics per keyword.
    """
    try:
        clean_name = (name or "").strip()
        if not clean_name:
            return None

        # Case-insensitive fetch or create
        obj = KeywordName.objects.filter(name__iexact=clean_name, language__iexact=lang).first()
        if not obj:
            obj = KeywordName.objects.create(
                name=clean_name,
                language=lang,
                source=source,
                relevance_score=score,
            )

        # ✅ Link multiple subtopics if provided
        if subtopics:
            if isinstance(subtopics, (list, tuple, set)):
                for st in subtopics:
                    if st:
                        obj.subtopics.add(st)
            else:
                obj.subtopics.add(subtopics)

        obj.save(update_fields=[])
        return obj

    except Exception as e:
        print("⚠️ Keyword Creation Error:", e)
        return None


# ============================================================
# 🔹 Generate AI-based Subtopics & Keywords
# ============================================================
@csrf_exempt
def generate_pyq_keywords(request):
    """
    🔹 AI-based Subtopic & Keyword Mapping
    - Maps Question → SubTopics (many-to-many)
    - Generates bilingual keywords
    - Classifies each keyword to correct SubTopic
    - Saves keyword–subtopic links
    - Also detects Unmapped (General) Keywords not linked to any subtopic
    - Returns rendered HTML for frontend
    """
    if request.method != "POST":
        return JsonResponse({"status": "error", "message": "Invalid request"}, status=400)

    qid = request.POST.get("question_id")
    if not qid:
        return JsonResponse({"status": "error", "message": "Missing question ID"}, status=400)

    try:
        # 1️⃣ Fetch Question
        question = get_object_or_404(QuestionBank, id=qid)
        q_text = f"""
Question: {question.question_part_first or question.question_part or ''}
Options:
A. {question.answer_option_a}
B. {question.answer_option_b}
C. {question.answer_option_c}
D. {question.answer_option_d}
Correct Answer: {question.correct_answer_choice}
Explanation: {question.correct_answer_description or ''}
"""

        # 2️⃣ Load available Subtopics
        subtopics = list(
            SubTopicName.objects.select_related("topic__chapter__part__section__area").values(
                "name", "sub_topic_short_Code", "topic__chapter__part__section__area__name",
            )
        )
        subtopic_list = [
            f"{s['topic__chapter__part__section__area__name']} → {s['name']} ({s['sub_topic_short_Code']})"
            for s in subtopics
        ]

        # 3️⃣ Map Question → Subtopics (AI)
        mapping_prompt = f"""
You are a UPSC/MPPSC content mapping expert at Hajela’s IAS Academy.

Analyse this question (options + explanation).
Map it to 2–5 conceptually relevant subtopics from the provided list.

Rules:
- Match by conceptual link, not word overlap.
- Include "confidence" (0–1) and "reason" in JSON.
- Output valid JSON only.

Question:
{q_text}

Available Subtopics:
{subtopic_list[:1709]}
"""
        completion = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "Output valid JSON only — no commentary or markdown."},
                {"role": "user", "content": mapping_prompt},
            ],
            temperature=0.3,
            max_tokens=1200,
        )

        raw_output = completion.choices[0].message.content.strip()
        cleaned = re.sub(r"```(?:json)?", "", raw_output).replace("```", "").strip()

        try:
            parsed_json = json.loads(cleaned)
            if isinstance(parsed_json, dict):
                matched_subtopics = (
                    parsed_json.get("matches")
                    or parsed_json.get("relevant_subtopics")
                    or parsed_json.get("mapped_subtopics")
                    or parsed_json.get("high_confidence_mappings")
                    or []
                )
            else:
                matched_subtopics = parsed_json
        except Exception:
            matched_subtopics = []

        # 4️⃣ Verify & Link Subtopics
        verified_matches, linked_subtopics, seen = [], [], set()
        question.subtopic_name.clear()

        for st in matched_subtopics:
            name_ai, conf, reason = "", 0.9, "AI thematic mapping"
            if isinstance(st, dict):
                name_ai = st.get("subtopic", "").strip().lower()
                conf = float(st.get("confidence", 0.9))
                reason = st.get("reason", reason)
            elif isinstance(st, str):
                name_ai = st.strip().lower()

            name_ai = re.sub(r"→", " ", name_ai)
            name_ai = re.sub(r"\(.*?\)", "", name_ai)
            name_ai = re.sub(
                r"\b(history|geography|economy|polity|environment|science|gs|general studies|other)\b",
                "",
                name_ai,
                flags=re.IGNORECASE,
            ).strip()

            if not name_ai or name_ai in seen:
                continue
            seen.add(name_ai)

            sub_obj = SubTopicName.objects.filter(name__icontains=name_ai).select_related(
                "topic__chapter__part__section__area"
            ).first()
            if not sub_obj:
                continue

            topic = sub_obj.topic
            chapter = topic.chapter if topic else None
            part = chapter.part if chapter else None
            section = part.section if part else None
            area = section.area if section else None

            if area: question.area_name.add(area)
            if section: question.section_name.add(section)
            if part: question.part_name.add(part)
            if chapter: question.chapter_name.add(chapter)
            if topic: question.topic_name.add(topic)

            question.subtopic_name.add(sub_obj)
            linked_subtopics.append(sub_obj)

            verified_matches.append({
                "area": getattr(area, "name", "—"),
                "area_color": getattr(area, "area_Colour_Hex", "#555"),
                "section": getattr(section, "name", "—"),
                "part": getattr(part, "name", "—"),
                "chapter": getattr(chapter, "name", "—"),
                "topic": getattr(topic, "name", "—"),
                "subtopic": sub_obj.name,
                "code": sub_obj.sub_topic_short_Code,
                "reason": reason,
                "confidence": round(conf, 2),
            })

        if linked_subtopics:
            question.ai_matched_subtopic = linked_subtopics[0]
            question.save(update_fields=["ai_matched_subtopic"])

        # 5️⃣ Generate Bilingual Keywords (AI)
        all_existing_keywords = set(KeywordName.objects.values_list("name", flat=True))
        existing_keywords_lower = {k.lower().strip() for k in all_existing_keywords}

        keyword_prompt = f"""
You are an expert UPSC/MPPSC keyword curator following Hajela’s IAS Keyword Rulebook v1.1.

Generate 15–20 precise, examinable keywords in English and Hindi for this question.
Avoid generic/country/state words and duplicates from this list:
{', '.join(list(existing_keywords_lower)[:200])}

Format strictly:
### ENGLISH
keyword1, keyword2, keyword3
---
### HINDI
कीवर्ड1, कीवर्ड2, कीवर्ड3

Question Context:
{q_text}
"""
        kw_completion = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "Return bilingual keywords only — no commentary."},
                {"role": "user", "content": keyword_prompt},
            ],
            temperature=0.4,
            max_tokens=900,
        )

        ai_text = kw_completion.choices[0].message.content.strip()
        en_keywords, hi_keywords = "", ""
        for block in ai_text.split("###"):
            block = block.strip().replace("---", "").strip()
            if block.lower().startswith("english"):
                en_keywords = re.sub(r"(?i)english", "", block).strip()
            elif block.lower().startswith("hindi") or "कीवर्ड" in block:
                hi_keywords = re.sub(r"(?i)hindi", "", block).strip()

        def clean_keywords(text):
            words = [k.strip() for k in text.split(",") if k.strip()]
            return [
                w for w in words
                if w.lower() not in IGNORED_KEYWORDS and w.lower() not in existing_keywords_lower
            ]

        en_clean = clean_keywords(en_keywords)
        hi_clean = clean_keywords(hi_keywords)

        # 6️⃣ Distribute Keywords by Subtopic (AI)
        distribution_prompt = f"""
You are a UPSC/MPPSC content classifier at Hajela’s IAS Academy.
Distribute the following keywords into the most relevant subtopics based on conceptual meaning.

Question Context:
{q_text}

Subtopics:
{[s['area'] + ' → ' + s['subtopic'] for s in verified_matches]}

English Keywords: {', '.join(en_clean)}
Hindi Keywords: {', '.join(hi_clean)}

Output valid JSON ONLY:
{{
  "subtopic_keyword_map": [
    {{
      "subtopic": "Economy → Major Crops and Cropping Patterns",
      "english_keywords": ["Agriculture Price Policy", "Farmer Distress"],
      "hindi_keywords": ["कृषि मूल्य नीति", "किसान संकट"]
    }}
  ]
}}
"""
        try:
            dist_completion = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "Output valid JSON only — no commentary."},
                    {"role": "user", "content": distribution_prompt},
                ],
                temperature=0.2,
                max_tokens=800,
            )
            dist_raw = dist_completion.choices[0].message.content.strip()
            dist_clean = re.sub(r"```(?:json)?", "", dist_raw).replace("```", "").strip()
            dist_json = json.loads(dist_clean)
            subtopic_map = {d["subtopic"].lower(): d for d in dist_json.get("subtopic_keyword_map", [])}
        except Exception as e:
            print("⚠️ Keyword distribution error:", e)
            subtopic_map = {}

        # 7️⃣ Attach distributed + DB keywords
        for st in verified_matches:
            st_key = f"{st['area']} → {st['subtopic']}".lower()
            mapped = subtopic_map.get(st_key, {})
            st["keywords_en"] = ", ".join(mapped.get("english_keywords", en_clean))
            st["keywords_hi"] = ", ".join(mapped.get("hindi_keywords", hi_clean))

            # ✅ Attach DB keywords from KeywordName table
            try:
                sub_obj = SubTopicName.objects.filter(name__iexact=st["subtopic"]).first()
                if sub_obj:
                    db_keywords = KeywordName.objects.filter(subtopics=sub_obj).values_list("name", flat=True)
                    st["db_keywords"] = ", ".join(sorted(set(db_keywords)))
                else:
                    st["db_keywords"] = ""
            except Exception as e:
                print("⚠️ DB keyword fetch error:", e)
                st["db_keywords"] = ""

        # 🟡 7B: Identify Unmapped (General) Keywords
        mapped_en = set()
        mapped_hi = set()
        for st in verified_matches:
            mapped_en.update([k.strip().lower() for k in st["keywords_en"].split(",") if k.strip()])
            mapped_hi.update([k.strip().lower() for k in st["keywords_hi"].split(",") if k.strip()])

        unmapped_en = [k for k in en_clean if k.lower() not in mapped_en]
        unmapped_hi = [k for k in hi_clean if k.lower() not in mapped_hi]

        general_keywords = {
            "english": ", ".join(unmapped_en),
            "hindi": ", ".join(unmapped_hi),
        }

        # 8️⃣ Save keywords to DB (EN + HI)
        for st in verified_matches:
            sub_obj = SubTopicName.objects.filter(name__iexact=st["subtopic"]).first()
            if not sub_obj:
                continue

            for kw in st["keywords_en"].split(","):
                if kw.strip():
                    safe_get_or_create_keyword(kw.strip(), "EN", "ai", 0.9, [sub_obj])

            for kw in st["keywords_hi"].split(","):
                if kw.strip():
                    safe_get_or_create_keyword(kw.strip(), "HI", "ai", 0.9, [sub_obj])

        # 9️⃣ Save overall keywords in QuestionBank
        question.key_words = f"EN: {', '.join(en_clean)} | HI: {', '.join(hi_clean)}"
        question.key_words_en = ", ".join(en_clean)
        question.key_words_hi = ", ".join(hi_clean)
        question.save(update_fields=["key_words", "key_words_en", "key_words_hi"])

        # 🔟 Render Result for Frontend
        html = render_to_string(
            "partials/ai_result_card.html",
            {
                "paired_blocks": [{
                    "english": question.question_part_first or "",
                    "hindi": question.question_part_first_hi or "",
                    "matched_subtopics": verified_matches or [],
                    "general_keywords": general_keywords,  # 🟡 Include unmapped keywords
                    "source": "AI-classified Keywords per Subtopic",
                }],
                "question_id": question.id,
            },
            request=request,
        )

        return JsonResponse({
            "status": "success",
            "html": html,
            "matched_count": len(verified_matches),
            "linked_subtopics": [s.sub_topic_short_Code for s in linked_subtopics],
            "hierarchy_tagged": True,
        })

    except Exception as e:
        print("❌ Keyword Generation Error:", traceback.format_exc())
        return JsonResponse({"status": "error", "message": str(e)}, status=500)



from django.views.decorators.csrf import csrf_exempt
from django.http import JsonResponse
from .models import Area, PartName, ChapterName, TopicName

@csrf_exempt
def get_scope_list(request):
    """Returns all Areas, Parts, Chapters, or Topics for dropdown."""
    try:
        scope = request.GET.get("scope", "").strip().lower()
        data = []

        if scope == "area":
            data = list(Area.objects.all().order_by("name").values("id", "name"))
        elif scope == "part":
            data = list(PartName.objects.all().order_by("name").values("id", "name"))
        elif scope == "chapter":
            data = list(ChapterName.objects.all().order_by("name").values("id", "name"))
        elif scope == "topic":
            data = list(TopicName.objects.all().order_by("name").values("id", "name"))
        else:
            return JsonResponse({"status": "error", "message": "Invalid scope"}, status=400)

        return JsonResponse({"status": "success", "items": data})

    except Exception as e:
        import traceback
        print("❌ get_scope_list error:", traceback.format_exc())
        return JsonResponse({"status": "error", "message": str(e)}, status=500)

import json, re, traceback
from django.shortcuts import get_object_or_404
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from openai import OpenAI
from .models import SubTopicName, QuestionBank, KeywordName

client = OpenAI(api_key="YOUR_OPENAI_API_KEY")

@csrf_exempt
def refine_subtopics_by_scope(request):
    """
    ✅ Enhanced:
    - Refines subtopics within selected scope (Area/Part/Chapter/Topic)
    - Tags refined subtopics with relevant keywords
    - Also saves question-level general keywords not mapped to any subtopic
    """
    try:
        payload = json.loads(request.body.decode("utf-8"))
        qid = payload.get("question_id")
        scope_type = payload.get("scope_type")
        selected_ids = payload.get("selected_ids", [])

        if not qid or not scope_type or not selected_ids:
            return JsonResponse({
                "status": "error",
                "message": "Missing required fields (question_id, scope_type, selected_ids)."
            }, status=400)

        question = get_object_or_404(QuestionBank, id=qid)

        # 🧭 Scope filtering
        if scope_type == "area":
            subtopics_qs = SubTopicName.objects.filter(
                topic__chapter__part__section__area__area_SI_Code__in=selected_ids
            )
        elif scope_type == "part":
            subtopics_qs = SubTopicName.objects.filter(
                topic__chapter__part__part_serial__in=selected_ids
            )
        elif scope_type == "chapter":
            subtopics_qs = SubTopicName.objects.filter(
                topic__chapter__chapter_number__in=selected_ids
            )
        elif scope_type == "topic":
            subtopics_qs = SubTopicName.objects.filter(
                topic__topic_SI_number__in=selected_ids
            )
        else:
            return JsonResponse({"status": "error", "message": "Invalid scope"}, status=400)

        if not subtopics_qs.exists():
            return JsonResponse({"status": "success", "data": [], "message": f"No subtopics found for {scope_type}."})

        # 🧠 Build question text
        q_text = f"""
Question: {question.question_part_first or question.question_part or ''}
Options:
A. {question.answer_option_a}
B. {question.answer_option_b}
C. {question.answer_option_c}
D. {question.answer_option_d}
Correct Answer: {question.correct_answer_choice}
Explanation: {question.correct_answer_description or ''}
"""

        subtopic_list = [
            f"{s.topic.chapter.part.section.area.name} → {s.name} ({s.sub_topic_short_Code})"
            for s in subtopics_qs.select_related("topic__chapter__part__section__area")
        ]

        prompt = f"""
You are an expert UPSC content classifier at Hajela’s IAS Academy.
Map this question to the most relevant subtopics ONLY from the filtered list below.
Also suggest 5–10 precise English and Hindi keywords for each mapped subtopic.

Return valid JSON like:
[
  {{
    "subtopic": "XYZ",
    "code": "ECO-203",
    "confidence": 0.91,
    "reason": "Conceptually linked to question",
    "keywords_en": ["Agriculture Policy", "MSP"],
    "keywords_hi": ["कृषि नीति", "न्यूनतम समर्थन मूल्य"]
  }}
]

Question:
{q_text}

Filtered Subtopics:
{subtopic_list[:1700]}
"""

        # 🔹 Run AI
        completion = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "Return valid JSON only — no commentary."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.4,
            max_tokens=1200,
        )

        raw_output = completion.choices[0].message.content.strip()
        print("\n🔍 RAW REFINEMENT OUTPUT:\n", raw_output, "\n")
        cleaned = re.sub(r"```(?:json)?", "", raw_output).replace("```", "").strip()

        try:
            parsed = json.loads(cleaned)
        except Exception:
            parsed = []

        response_data = []
        all_keywords_en, all_keywords_hi = set(), set()

        # 🧩 Process and save each refined subtopic + keywords
        for st in parsed:
            sub_name = st.get("subtopic", "").strip()
            if not sub_name:
                continue

            sub_obj = (
                SubTopicName.objects.filter(name__icontains=sub_name)
                .select_related("topic__chapter__part__section__area")
                .first()
            )

            if sub_obj:
                area = sub_obj.topic.chapter.part.section.area.name
                color = sub_obj.topic.chapter.part.section.area.area_Colour_Hex
                question.subtopic_name.add(sub_obj)
            else:
                area, color = "General", "#999"

            keywords_en = st.get("keywords_en", [])
            keywords_hi = st.get("keywords_hi", [])
            all_keywords_en.update(keywords_en)
            all_keywords_hi.update(keywords_hi)

            # 🗂 Save keywords in DB linked to subtopic
            for kw in keywords_en:
                KeywordName.objects.get_or_create(
                    name=kw.strip(),
                    language="EN",
                    defaults={"source": "ai"},
                )[0].subtopics.add(sub_obj)
            for kw in keywords_hi:
                KeywordName.objects.get_or_create(
                    name=kw.strip(),
                    language="HI",
                    defaults={"source": "ai"},
                )[0].subtopics.add(sub_obj)

            response_data.append({
                "subtopic": sub_name,
                "area": area,
                "area_color": color,
                "confidence": float(st.get("confidence", 0.9)),
                "reason": st.get("reason", "Conceptually linked"),
                "keywords_en": keywords_en,
                "keywords_hi": keywords_hi,
            })

        # 🟢 Save unmapped general keywords at question-level
        question.key_words_en = ", ".join(sorted(all_keywords_en))
        question.key_words_hi = ", ".join(sorted(all_keywords_hi))
        question.key_words = f"EN: {question.key_words_en} | HI: {question.key_words_hi}"
        question.save(update_fields=["key_words", "key_words_en", "key_words_hi"])

        print(f"✅ Refined {len(response_data)} subtopics and saved keywords to DB.\n")

        return JsonResponse({
            "status": "success",
            "data": response_data,
            "count": len(response_data),
        })

    except Exception as e:
        print("❌ refine_subtopics_by_scope Error:", traceback.format_exc())
        return JsonResponse({"status": "error", "message": str(e)}, status=500)


@csrf_exempt
def update_question_subtopics(request):
    """
    ✅ Saves both original matched and refined subtopics
    and updates hierarchy (Area → Section → Part → Chapter → Topic).
    """
    try:
        payload = json.loads(request.body.decode("utf-8"))
        qid = payload.get("question_id")
        selected_names = payload.get("selected_subtopics", [])

        if not qid:
            return JsonResponse({"status": "error", "message": "Missing question_id"}, status=400)

        question = get_object_or_404(QuestionBank, id=qid)

        # 🧠 Clean names from "Polity → ..." and "(564)"
        cleaned = []
        for name in selected_names:
            if not name: continue
            name_clean = name.split("→")[-1].split("(")[0].strip()
            cleaned.append(name_clean)

        # 🔍 Flexible name matching
        from django.db.models import Q
        query = Q()
        for n in cleaned:
            query |= Q(name__iexact=n) | Q(name__icontains=n)

        subtopic_objs = SubTopicName.objects.filter(query).select_related(
            "topic__chapter__part__section__area"
        ).distinct()

        # 🔄 Clear old & add new
        question.subtopic_name.clear()
        linked = []
        for st in subtopic_objs:
            question.subtopic_name.add(st)
            linked.append(st.name)

            # ✅ auto-link hierarchy
            topic = st.topic
            chapter = topic.chapter if topic else None
            part = chapter.part if chapter else None
            section = part.section if part else None
            area = section.area if section else None

            if area: question.area_name.add(area)
            if section: question.section_name.add(section)
            if part: question.part_name.add(part)
            if chapter: question.chapter_name.add(chapter)
            if topic: question.topic_name.add(topic)

        question.save()

        return JsonResponse({
            "status": "success",
            "count": subtopic_objs.count(),
            "linked": linked,
            "message": f"✅ {subtopic_objs.count()} subtopics saved successfully (Matched + Refined)."
        })

    except Exception as e:
        print("❌ update_question_subtopics Error:", traceback.format_exc())
        return JsonResponse({"status": "error", "message": str(e)}, status=500)


from django.views.decorators.http import require_POST
from django.http import JsonResponse
from django.shortcuts import get_object_or_404
from django.views.decorators.csrf import csrf_exempt
import json
from .models import QuestionBank, KeywordName, SubTopicName

import re, random, json, traceback
from django.views.decorators.csrf import csrf_exempt
from django.http import JsonResponse
from django.shortcuts import get_object_or_404
from django.db.models import Q
from django.contrib.auth.decorators import login_required
from .models import QuestionBank, SubTopicName


@csrf_exempt
@login_required
def update_question_keywords(request):
    """
    ✅ Bilingual Version:
    - Creates BOTH English & Hindi alternate questions
    - Saves keywords, subtopics, and copies hierarchy
    - Mimics format and logic of save_alternate_question()
    """
    if request.method != "POST":
        return JsonResponse({"status": "error", "message": "Invalid request method"}, status=400)

    try:
        # ---------- Parse Data ----------
        data = json.loads(request.body.decode("utf-8"))
        qid = data.get("question_id")
        english_keywords = data.get("english_keywords", "").strip()
        hindi_keywords = data.get("hindi_keywords", "").strip()
        selected_subtopics = data.get("subtopics", [])

        if not qid:
            return JsonResponse({"status": "error", "message": "Missing question_id"}, status=400)

        # ---------- Source & Meta ----------
        source_question = get_object_or_404(QuestionBank, id=qid)
        subtype = source_question.question_sub_type

        # ---------- Generate Base ID + Question Number ----------
        last_q = QuestionBank.objects.order_by("-question_number").first()
        next_number = (last_q.question_number + 1) if last_q else 1

        while True:
            base_id = random.randint(100000, 999999)
            if not QuestionBank.objects.filter(base_question_id=base_id).exists():
                break

        linked_names = []

        # ---------- Common Function to Create a Question ----------
        def create_alt_question(lang="e", keywords_en="", keywords_hi=""):
            q = QuestionBank.objects.create(
                base_question_id=base_id,
                question_number=next_number,
                question_sub_type=subtype,
                language=lang,
                type_of_question="moq",
                created_by=request.user,
                degree_of_difficulty="Medium",
                elim_tactics_degree="Normal",
                current_relevance="General",
                key_words=keywords_en if lang == "e" else "",
                key_words_hi=keywords_hi if lang == "h" else "",
            )

            # Copy Hierarchy from Source
            q.area_name.set(source_question.area_name.all())
            q.section_name.set(source_question.section_name.all())
            q.part_name.set(source_question.part_name.all())
            q.chapter_name.set(source_question.chapter_name.all())
            q.topic_name.set(source_question.topic_name.all())

            return q

        # ---------- Create EN + HI versions ----------
        q_en = create_alt_question(lang="e", keywords_en=english_keywords)
        q_hi = create_alt_question(lang="h", keywords_hi=hindi_keywords)

        # ---------- Match Subtopics ----------
        if selected_subtopics:
            clean_names = []
            for s in selected_subtopics:
                s_clean = s.split("→")[-1].split("(")[0].strip()
                if s_clean and s_clean not in clean_names:
                    clean_names.append(s_clean)

            if clean_names:
                query = Q()
                for n in clean_names:
                    query |= Q(name__iexact=n) | Q(name__icontains=n)

                matched = SubTopicName.objects.filter(query).select_related(
                    "topic__chapter__part__section__area"
                ).distinct()

                for st in matched:
                    q_en.subtopic_name.add(st)
                    q_hi.subtopic_name.add(st)
                    linked_names.append(st.name)

                    topic = getattr(st, "topic", None)
                    chapter = getattr(topic, "chapter", None)
                    part = getattr(chapter, "part", None)
                    section = getattr(part, "section", None)
                    area = getattr(section, "area", None)

                    if area:
                        q_en.area_name.add(area)
                        q_hi.area_name.add(area)
                    if section:
                        q_en.section_name.add(section)
                        q_hi.section_name.add(section)
                    if part:
                        q_en.part_name.add(part)
                        q_hi.part_name.add(part)
                    if chapter:
                        q_en.chapter_name.add(chapter)
                        q_hi.chapter_name.add(chapter)
                    if topic:
                        q_en.topic_name.add(topic)
                        q_hi.topic_name.add(topic)

        q_en.save()
        q_hi.save()

        # ---------- Response ----------
        return JsonResponse({
            "status": "success",
            "message": f"✅ Bilingual Alternate Questions saved successfully with {len(linked_names)} Subtopics.",
            "linked_subtopics": linked_names,
            "new_ids": [q_en.id, q_hi.id],
            "base_question_id": base_id,
            "subtype": subtype,
        })

    except Exception as e:
        print("❌ update_question_keywords Error:", traceback.format_exc())
        return JsonResponse({"status": "error", "message": str(e)}, status=500)



from django.shortcuts import render, get_object_or_404
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.conf import settings
from openai import OpenAI
from .models import SubTopicName, MicroSubTopicName

# client = OpenAI(api_key=settings.OPENAI_API_KEY)

def subtopic_list_view(request):
    """Show all SubTopics with button to generate keywords"""
    subtopics = SubTopicName.objects.select_related('topic').all()

    # Convert MicroSubTopicName objects into {subtopic_id: [list of keywords]}
    existing_keywords = {}
    for ms in MicroSubTopicName.objects.select_related('subtopics'):
        kw_list = [kw.strip() for kw in ms.name.split(',') if kw.strip()]
        existing_keywords[ms.subtopics.sub_topic_SI_Number] = kw_list

    return render(
        request,
        'question_bank/subtopics_list.html',
        {
            'subtopics': subtopics,
            'existing_keywords': existing_keywords
        }
    )


@csrf_exempt
def generate_keywords_view(request, subtopic_id):
    """Generate conceptual or PYQ-based keywords for a SubTopic."""
    subtopic = get_object_or_404(SubTopicName, pk=subtopic_id)
    mode = request.GET.get('mode', 'conceptual').lower()

    if mode == "pyq":
        # 📘 PYQ-based prompt (exam trend-aware)
        prompt = f"""
You are an expert UPSC and State PSC exam analyst.

TASK:
Generate 15–20 exam-relevant keywords for the following Sub-Topic,
based on previous year questions (PYQs) asked in UPSC, MPPSC and other State PSC exams.

Sub-Topic: {subtopic.name}
Topic: {subtopic.topic.name if subtopic.topic else 'N/A'}
Short Code: {subtopic.sub_topic_short_Code}

Guidelines:
- Base your output on PYQ trends.
- Focus on frequently tested entities, acts, articles, committees, reports, and institutions.
- Return only comma-separated keywords (no numbering or explanation).
"""
    else:
        # 🧠 CONCEPTUAL PROMPT (Full Hajela’s IAS Rulebook v1.7)
        prompt = f"""
You are an expert UPSC/MPPSC keyword curator.

TASK: 
1. Generate 20–30 high-signal keywords in English. The rules (Rulebook for Sub Topics MCQ Keywording (v1.7 – 2025-10-04 IST)) to be followed while creating keywords are:
   a. Scope & Philosophy
      i. Hajela's IAS (HIA) is a Bhopal-based coaching institution for various Civil Services examinations (UPSC, MPPSC, etc.). Its database contains:
         1. Topics (basic unit: "Sub_topic")
         2. Questions (PYQs, Model Questions)
         3. Content (lecture videos, PDFs, notes)
      ii. The database is divided into Areas → Parts → Chapters → Topics → Sub-topics.
      iii. Areas:
         1. History
         2. Geography
         3. Polity
         4. Economy
         5. Environment
         6. Science
         7. General Knowledge (GK)
         8. Current Affairs
         9. International Relations
         10. Ethics
         11. HIA Special
         12. Art & Culture
         13. Hindi
         14. English
         15. Internal Security & Disaster Management
         16. Humanities & Management
         17. MP History
         18. MP Geography
         19. Information and Communication Technology (ICT)
         20. MP Tribes
         21. MP Polity
         22. MP Economy
         23. GS2 (General Studies 2) consisting of Comprehension, Logical Reasoning etc.
      iv. Parts (≈140):
         1. Ancient India
         2. Medieval India
         3. The 18th Century India
         4. The 19th Century India
         5. Indian National Movement
         6. Modern Indian History
         7. World History
         8. Physical Geography Astronomy
         9. World Area Geography
         10. Physical Geography Earth
         11. Physical Geography Lithosphere
         12. Physical Geography Atmosphere
         13. Physical Geography Climatology
         14. Physical Geography Hydrosphere
         15. World Human Geography
         16. World Economic Geography
         17. India Area Geography
         18. Indian Human Geography
         19. Indian Physical Geography
         20. Indian Economic Geography
         21. Geographical Terms
         22. Geo-informatics
         23. Introduction to Polity
         24. System of Governance
         25. Constitutional Framework
         26. Constitutional Details
         27. Details of Governance
         28. Welfare State
         29. Constitutional Bodies
         30. Working of the Constitution
         31. Bodies other than Constitutional Bodies
         32. Other Constitutional Dimensions
         33. Elections
         34. Other Political Dynamics
         35. Comparison of the Constitution
         36. Polity-Others
         37. Introduction & Basic Concepts
         38. Macroeconomics
         39. Microeconomics
         40. General Economics
         41. Health
         42. Education
         43. Sociology & Economy
         44. Government Policies and Interventions
         45. Economy Current Issues
         46. General Environment
         47. Biodiversity
         48. Biodiversity Chemistry
         49. Climate Change
         50. Science Tech
         51. Space Tech
         52. Defence Tech
         53. Nuclear Tech
         54. Nano Tech
         55. Intellectual Property
         56. Biotechnology
         57. Biosciences
         58. Physics
         59. Chemistry
         60. Geometry & Mensuration
         61. Arithmetic
         62. Statistics & Mathematics
         63. Permutation, Combination & Probability
         64. GK
         65. International Current Affairs
         66. National Current Affairs
         67. Madhya Pradesh Current Affairs
         68. Current Affairs - General
         69. Indian Foreign Relations
         70. Groupings & International Institutions
         71. International Relations Current Issues
         72. Ethics
         73. Psychology
         74. Ethics & Philosophy
         75. Ethics in Public Administration
         76. Ethics, Life Style and Counter Force MP
         77. Art Forms
         78. Painting
         79. Literature
         80. Architecture
         81. Music
         82. World Art & Culture
         83. MP Art & Culture
         84. General Hindi GS UPSC
         85. Hindi General MP Mains
         86. Hindi Essay MP Mains
         87. English UPSC
         88. Internal Security
         89. Disaster Management
         90. Sociology
         91. Entrepreneurship & Management MP
         92. MP History
         93. MP Geography
         94. Information Technology
         95. M P Tribes
         96. MP Polity
         97. MP Education
         98. MP Economy
         99. MP Health
         100. Comprehension
         101. Interpersonal & Communication Skills
         102. Logical Reasoning and Analytical Ability
         103. Decision Making
         104. Mental Ability
         105. Numeracy
         106. Data Interpretation
         107. Comprehension MP
         108. Communication Skills MP
         109. Logical Reasoning and Analytical Ability MP
         110. Decision Making MP
         111. Mental Ability MP
         112. Numeracy MP
         113. Data Interpretation MP
         114. Essay UPSC
         115. Test
         116. Mock Interview
      v. Next levels:
         1. ≈450 Chapters
         2. ≈1200 Topics
         3. ≈1700 Sub-topics
         4. The details are available at the database ...............
      vi. Keywords are the bridge between Sub-topics and Questions:
         1. Keywords must be generated for both Questions and Sub-topics.
         2. Matching keywords enable tagging Questions ↔ Sub-topics.
         3. Target: 15–20 high-signal, examinable, precise phrases per Sub Topic.
         4. General properties:
            a. Case-insensitive
            b. Sequence-sensitive
            c. Keywords (backend) vs. Hashtags (frontend) → system relies on keywords.
      vii. Presently, keywords need to be generated only for Sub-topics:
         1. For generation of Keywords, the Area, Part, Chapter, Topic and Sub Topic need to be considered together.
   b. Hard Deletions (do not create such keywords unless exempted)
      i. Axis tails (generic): polity, economy, society, religion, culture, art & architecture, literature, chronology, sources.
      ii. Era mislabels: standalone "Ancient India", "Medieval India", "Modern India".
      iii. Noisy stand-alones:
         1. Generic nouns: development, reforms, rule, law, events, contribution, emergence, knowledge, mapping, field, records, rise, factors, responsible.
         2. Unqualified geographies: Asia, Europe, Africa, America, World.
         3. Bare science terms: geomagnetic field, rocks, theory, effect, eruption.
         4. Filler adjectives: important, major, various, different, key.
         5. Meta-words: overview, background, miscellaneous, types.
         6. Roman numerals (if alone).
         7. Initials/honorifics: Dr., Pt., Sri, Netaji, Mahatma, Lokmanya, Begum, Rao, etc.
         8. Duplicate forms (case/punctuation variants).
         9. Standalone years (normalize or drop).
         10. Example deletions: stone, india, post, indo, dynasties, salient events, empires, across, world, century, wars, scenario, india 1.
   c. Using High-signal Terms (examples)
      i. Era anchoring:
         1. "Ancient India (Pre-Mauryan)", "Medieval India (Delhi Sultanate)""Medieval India (1000-1700 AD)".
      ii. Geographic anchoring:
         1. Tie "India" with period (e.g., Modern India).
         2. Foreign example: "European Renaissance (15th–16th c.)".
      iii. Dynasty/Era/Text expansion:
         1. Maurya → Mauryan Empire
         2. Gupta → Gupta Empire
         3. Rigveda → Rigveda (Vedic Text)
      iv. Conceptual enrichment:
         1. Architecture → Nagara temple style / Stupa–Chaitya–Vihara
         2. Literature → Sangam corpus / Arthashastra (text)
      v. Year normalization:
         1. 1200 → 1200 AD; 600 BCE
         2. Use en dash: 1775–82
      vi. Relation expansion:
         1. US USSR relations → US–USSR relations (Cold War)
         2. America formation → Formation of the USA (1776–1789)
   d. Facet Templates
      i. Persons/Leaders: role, tenure, key battles, treaties, allies/enemies.
      ii. Wars/Battles: battle name, year, commanders, outcomes, treaties.
      iii. Constitutional Acts/Missions: provisions, criticisms, institutions formed.
      iv. Mass movements: aims, slogans, sessions, withdrawals, spread.
      v. Civilisations/Cultural: polity, economy, society/religion, art, literature, sources.
      vi. Foreign relations: bind to era + geography.
   e. Question Keyword Rules (APQO/RQK)
      i. Base context: Area + Part (+ Chapter/Topic if available).
      ii. Stem: key entities/ideas.
      iii. Options: include all 4 answer options.
      iv. Correct answer description: enrich for concept precision.
   f. QA Checklist
      i. Era alignment with Area/Part/Chapter/Topic.
      ii. ≥70% high-signal entities/events/ideas.
      iii. De-duplicate by case/punctuation.
      iv. Typography: en dash, proper names, full forms.
   g. Other Rules (evolving)
      i. Core matching semantics: case-insensitive, order-sensitive, lowercase internal, Title Case display.
      ii. Noise deletion: drop generic singletons, unqualified labels, naked years, filler tokens, exam-instruction artifacts.
      iii. Concept enrichment & canonicalization:
         1. Acts/Sessions/Movements with year (e.g., Government of India Act 1919 – Dyarchy).
         2. Person


2.  Output: comma-separated, Title Case where needed.

3. Translate the keywords generated also into Hindi (keywords only, no sentences).

Now, generate keywords for the following:

Sub-Topic: {subtopic.name}  
Topic: {subtopic.topic.name if subtopic.topic else 'N/A'}  
Short Code: {subtopic.sub_topic_short_Code}

Output Format:
English Keywords (comma-separated)
Hindi Keywords (comma-separated)
"""

    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a civil services keyword tagging expert following Hajela’s IAS Keyword Rulebook v1.7."},
                {"role": "user", "content": prompt},
            ],
            max_tokens=500,
            temperature=0.7,
        )

        keywords = response.choices[0].message.content.strip()

        # Save separately for conceptual / pyq
        MicroSubTopicName.objects.update_or_create(
            micro_sub_topic_SI_number=f"{subtopic.sub_topic_SI_Number}_{mode.upper()}",
            defaults={"name": keywords, "subtopics": subtopic}
        )

        return JsonResponse({"success": True, "keywords": keywords, "mode": mode})
    except Exception as e:
        return JsonResponse({"success": False, "error": str(e)})



from django.contrib.auth.decorators import login_required
from .models import QuestionBank

@login_required
def view_ai_generated_questions(request):
    questions = QuestionBank.objects.filter(
        correct_answer_description__isnull=False,
        question_sub_type='simple_type'  # You can remove this filter if you want all types
    ).order_by('-created_at')  # Shows newest first
    return render(request, 'question_bank/ai_questions.html', {'questions': questions})



import os
import re
from datetime import datetime
from io import BytesIO

from PIL import Image as PILImage
from django.conf import settings
from django.http import FileResponse, HttpResponse
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches

from .models import QuestionBank

def clean_text(text):
    """Utility function to clean and format text by removing extra newlines and spaces."""
    if not text:
        return ''
    text = text.strip()
    text = re.sub(r'\s*\n\s*', ' ', text)
    text = re.sub(r'\s+', ' ', text)
    return text

def set_no_border(cell):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tcBorders.append(border)
    tcPr.append(tcBorders)

def generate_questions_document(request):
    try:
        # ✅ Get selected IDs from POST
        selected_ids = request.POST.getlist('selected_questions')  # ✅ correct
  # JavaScript sends this list via AJAX
        if not selected_ids:
            return HttpResponse("No questions selected!", status=400)

        questions = QuestionBank.objects.filter(id__in=selected_ids)

        # Setup document
        base_dir = os.path.join(settings.MEDIA_ROOT, 'word_file')
        os.makedirs(base_dir, exist_ok=True)
        today = datetime.today().strftime('%Y-%m-%d')
        file_name = f'class_plus_questions_{today}.docx'
        file_path = os.path.join(base_dir, file_name)
        document = Document()

        for question in questions:
            table = document.add_table(rows=0, cols=3)
            table.style = 'Table Grid'

            if question.question_sub_type == "r_and_a_type":
                question_text = (
                    clean_text(question.question_part_first or '') + "\n" +
                    (clean_text(question.assertion) if question.assertion else '') + "\n" +
                    (clean_text(question.reason) if question.reason else '') + "\n" +
                    clean_text(question.question_part_third or '')
                )
            else:
                if question.question_part and question.question_part.strip():
                    question_text = clean_text(question.question_part)
                else:
                    question_text = (
                        clean_text(question.question_part_first or '') + "\n" +
                        ("Assertion (A): " + clean_text(question.assertion) if question.assertion else '') + "\n" +
                        ("Reason (R): " + clean_text(question.reason) if question.reason else '') + "\n" +
                        clean_text(question.question_part_third or '')
                    )

            q_row = table.add_row().cells
            q_row[0].text = 'Question'

            if question.question_sub_type == 'list_type_1':
                q_row[1].text = f"{clean_text(question.question_part_first or '')}\n"
                q_row[1].text += "\n".join([
                    f"{i}. {getattr(question, f'list_1_row{i}', '')}"
                    for i in range(1, 9) if getattr(question, f'list_1_row{i}', '')
                ]) + "\n" + clean_text(question.question_part_third or '')
            elif question.list_1_name and question.list_2_name:
                sub_table = document.add_table(rows=1, cols=2)
                sub_table.style = 'Table Grid'
                sub_hdr_cells = sub_table.rows[0].cells
                sub_hdr_cells[0].text = f"LIST - I\n({clean_text(question.list_1_name)})"
                sub_hdr_cells[1].text = f"LIST - II\n({clean_text(question.list_2_name)})"

                for i in range(1, 9):
                    list_1_option = getattr(question, f'list_1_row{i}', '')
                    list_2_option = getattr(question, f'list_2_row{i}', '')
                    if not list_2_option:
                        break
                    row_cells = sub_table.add_row().cells
                    row_cells[0].text = f"{chr(64+i)}. {list_1_option}"
                    row_cells[1].text = f"{i}. {list_2_option}"

                q_row[1]._element.clear_content()
                p = q_row[1].add_paragraph()
                p.add_run(clean_text(question.question_part_first or '') + "\n")
                q_row[1]._element.append(sub_table._element)
                q_row[1].add_paragraph("\nCodes:\t A\t B\t C\t D")
            else:
                q_row[1].text = question_text

            q_row[1].merge(q_row[2])

            # Insert Image
            if question.image:
                image_path = question.image.path
                pil_img = PILImage.open(image_path)
                if pil_img.mode == 'RGBA':
                    pil_img = pil_img.convert('RGB')
                img_io = BytesIO()
                pil_img.save(img_io, format='JPEG')
                img_io.seek(0)
                q_row[1].add_paragraph().add_run().add_picture(img_io, width=Inches(1.5))

            # Handle table options (A-D)
            valid_options = ['a', 'b', 'c', 'd']
            correct_answer = question.correct_answer_choice.lower() if question.correct_answer_choice else None
            for opt in valid_options:
                option_text = getattr(question, f"answer_option_{opt}", None)
                if option_text:
                    opt_row = table.add_row().cells
                    opt_row[0].text = 'Option'
                    opt_row[1].text = f"{opt.upper()}. {option_text}"
                    opt_row[2].text = 'correct' if opt == correct_answer else 'incorrect'

            if question.question_sub_type == 'fill_in_the_blank_type':
                option_row = table.add_row().cells
                option_row[0].text = 'Option'
                option_row[1].text = question.correct_answer_choice or '___'

            if question.question_sub_type == 'true_and_false_type':
                answer_row = table.add_row().cells
                answer_row[0].text = 'Answer'
                correct_answer_letter = question.correct_answer_choice.lower()
                correct_answer_text = getattr(question, f"answer_option_{correct_answer_letter}", None)
                answer_row[1].text = correct_answer_text or "Answer not available"
                answer_row[1].merge(answer_row[2])

            # Additional Info
            table.add_row().cells[0].text = 'Type'
            type_row = table.rows[-1].cells
            type_row[1].text = (
                'true_false' if question.question_sub_type == 'true_and_false_type'
                else 'fill_ups' if question.question_sub_type == 'fill_in_the_blank_type'
                else 'multiple_choice'
            )
            type_row[1].merge(type_row[2])

            sol_row = table.add_row().cells
            sol_row[0].text = 'Solution'
            sol_row[1].text = clean_text(question.correct_answer_description)
            sol_row[1].merge(sol_row[2])

            marks_row = table.add_row().cells
            marks_row[0].text = 'Marks'
            marks_row[1].text = str(question.marks)
            marks_row[2].text = str(question.negative_marks)

            diff_row = table.add_row().cells
            diff_row[0].text = 'Degree of Difficulty'
            diff_row[1].text = clean_text(question.degree_of_difficulty or '')
            diff_row[1].merge(diff_row[2])

            document.add_paragraph()  # space between questions

        document.save(file_path)
        return FileResponse(open(file_path, 'rb'), as_attachment=True, filename=file_name)

    except Exception as e:
        return HttpResponse(f"An error occurred: {str(e)}", status=500)


# views.py
from django.shortcuts import render
from .models import QuestionBank

def generate_classplus_doc_view(request):
    questions = QuestionBank.objects.all()
    return render(request, 'question_bank/generate_classplus_doc.html', {'questions': questions})


# ************************* Generate Clas Plus Word file Start *********************************************


# ************************* Upload Excel file Start *********************************************

# Function to safely convert exam years to integers
def safe_int(value, default=0):
    try:
        return int(value)
    except (ValueError, TypeError):
        return default  # Return a default value if conversion fails

# Function to safely convert numeric values to float
def safe_float(value, default=0.0):
    try:
        return float(value) if value not in [None, '', ' '] else default
    except (ValueError, TypeError):
        return default  # Return default if conversion fails

import random
from django.db import transaction

import re
import pandas as pd
from django.shortcuts import render, redirect
from django.contrib import messages
from django.core.files.storage import FileSystemStorage
from django.db.models import Max

from .models import QuestionBank, Area, PartName
from .forms import UploadFileForm


def safe_int(value, default=0):
    try:
        return int(value)
    except (ValueError, TypeError):
        return default


def safe_float(value, default=0.0):
    try:
        return float(value)
    except (ValueError, TypeError):
        return default


def _pick(row, *keys):
    """Return the first non-empty value from the given candidate keys (strings)."""
    for k in keys:
        v = row.get(k, '')
        if v is None:
            continue
        # pandas can give floats for empty cells; coerce to str safely
        s = str(v).strip()
        if s != '' and s.lower() != 'nan':
            return s
    return ''


# ************************* Upload Excel file Start *********************************************
# ************************* Upload Excel file (updated) *********************************************

import re
import pandas as pd
import random
from django.db import transaction
from django.shortcuts import render, redirect
from django.contrib import messages
from django.core.files.storage import FileSystemStorage
from django.db.models import Max
from django.contrib.auth import get_user_model

from .forms import UploadFileForm
from .models import (
    QuestionBank,
    Area, Section, PartName, ChapterName, TopicName, SubTopicName,
    ExamName, EvergreenIndexName
)

User = get_user_model()

# ---------- tiny helpers ----------

def safe_int(value, default=0):
    try:
        return int(value)
    except (ValueError, TypeError):
        return default

def safe_float(value, default=0.0):
    try:
        if value in (None, '', ' '):
            return default
        return float(value)
    except (ValueError, TypeError):
        return default

def _pick(row, *keys):
    """First non-empty (non-NaN) string from candidate headers."""
    for k in keys:
        if k in row:
            v = row.get(k, '')
            if v is None:
                continue
            s = str(v).strip()
            if s and s.lower() != 'nan':
                return s
    return ''

def _split_multi(val):
    """Split on | or , and trim; ignore blanks/-."""
    if val is None:
        return []
    s = str(val).replace(',', '|')
    return [p.strip() for p in s.split('|') if p and p.strip() and p.strip() != '-']

def _allowed_field_names():
    return {f.name for f in QuestionBank._meta.get_fields()}

def _only_allowed(d: dict, allowed: set) -> dict:
    return {k: v for k, v in d.items() if k in allowed and v not in (None, '')}

def _normalize_source(val: str) -> str:
    v = (val or '').strip().lower()
    if v in {'modern', 'model', 'model question', 'model_question', 'moq'}:
        return 'moq'
    if v in {'pyq', 'previous', 'previous year', 'previous year question', 'prev'}:
        return 'pyq'
    if v in {'osq', 'other', 'other source', 'other source question'}:
        return 'osq'
    return 'moq'

def _qs_by_codes_or_names(model, code_field, name_field, codes_value, names_value):
    codes = _split_multi(codes_value)
    if codes:
        return model.objects.filter(**{f"{code_field}__in": codes})
    names = _split_multi(names_value)
    if names:
        return model.objects.filter(**{f"{name_field}__in": names})
    return model.objects.none()

# ---------- the view ----------

def upload_file(request):
    if request.method == 'POST':
        form = UploadFileForm(request.POST, request.FILES)
        if form.is_valid():
            # store file
            xfile = request.FILES['file']
            fs = FileSystemStorage()
            filename = fs.save(xfile.name, xfile)

            # read Excel
            data = pd.read_excel(fs.path(filename))
            data = data.fillna({
                'marks': 0,
                'negative_marks': 0,
                'exam_year1': 0,
                'exam_year2': 0,
                'exam_year3': 0,
                'exam_year4': 0,
            }).fillna('')

            # next question_number (shared for EN+HI)
            max_qno = QuestionBank.objects.aggregate(Max('question_number'))['question_number__max']
            next_qno = int(max_qno) + 1 if max_qno else 1

            allowed = _allowed_field_names()

            with transaction.atomic():
                for _, row in data.iterrows():
                    # keep question_number unique across table
                    while QuestionBank.objects.filter(question_number=next_qno).exists():
                        next_qno += 1

                    # ---- source & year ----
                    type_of_question = _normalize_source(_pick(row, 'question_type_source', 'type_of_question'))
                    years = [safe_int(row.get('exam_year1', 0)),
                             safe_int(row.get('exam_year2', 0)),
                             safe_int(row.get('exam_year3', 0)),
                             safe_int(row.get('exam_year4', 0))]
                    exam_year = next((y for y in years if y != 0), None)

                    # ---- M2Ms via *codes or name columns ----
                    areas_qs = _qs_by_codes_or_names(
                        Area, 'area_SI_Code', 'name',
                        row.get('area_codes'), _pick(row, 'area_name')
                    )
                    sections_qs = _qs_by_codes_or_names(
                        Section, 'section_Unit_SI', 'name',
                        row.get('section_codes'), _pick(row, 'section_name')
                    )
                    parts_qs = _qs_by_codes_or_names(
                        PartName, 'part_serial', 'name',
                        row.get('part_codes'), _pick(row, 'part_name')
                    )
                    chapters_qs = _qs_by_codes_or_names(
                        ChapterName, 'chapter_number', 'name',
                        row.get('chapter_codes'), _pick(row, 'chapter_name')
                    )
                    topics_qs = _qs_by_codes_or_names(
                        TopicName, 'topic_SI_number', 'name',
                        row.get('topic_codes'), _pick(row, 'topic_name')
                    )
                    subtopics_qs = _qs_by_codes_or_names(
                        SubTopicName, 'sub_topic_short_Code', 'name',
                        row.get('subtopic_codes'), _pick(row, 'subtopic_name')
                    )
                    exams_qs = _qs_by_codes_or_names(
                        ExamName, 'exam_SI_Number', 'name',
                        row.get('exam_codes'), _pick(row, 'exam_name')
                    )
                    evergreen_qs = _qs_by_codes_or_names(
                        EvergreenIndexName, 'evergreen_index_SI_Number', 'name',
                        row.get('evergreenindex_codes'), _pick(row, 'evergreenindex_name')
                    )

                    # ---- subtype & stems/options/RA ----
                    question_sub_type = _pick(row, 'question_sub_type') or 'simple_type'

                    # EN
                    qpf_en = _pick(row, 'question_part_first_part', 'question_part_first', 'question_stem_en')
                    qpt_en = _pick(row, 'question_part_third_part', 'question_part_third', 'question_tail_en')
                    ans_a = _pick(row, 'answer_option_a')
                    ans_b = _pick(row, 'answer_option_b')
                    ans_c = _pick(row, 'answer_option_c')
                    ans_d = _pick(row, 'answer_option_d')

                    # HI
                    qpf_hi = _pick(row, 'question_part_first_part_hindi', 'question_part_first_hi', 'question_stem_hi', 'question_part_first_part_hi')
                    qpt_hi = _pick(row, 'question_part_third_part_hindi', 'question_part_third_hi', 'question_tail_hi', 'question_part_third_part_hi')
                    ans_a_hi = _pick(row, 'answer_option_a_hi', 'answer_option_a_hindi')
                    ans_b_hi = _pick(row, 'answer_option_b_hi', 'answer_option_b_hindi')
                    ans_c_hi = _pick(row, 'answer_option_c_hi', 'answer_option_c_hindi')
                    ans_d_hi = _pick(row, 'answer_option_d_hi', 'answer_option_d_hindi')

                    # RA
                    assertion_en = _pick(row, 'assertion')
                    reason_en    = _pick(row, 'reason')
                    assertion_hi = _pick(row, 'assertion_hindi', 'assertion_hi')
                    reason_hi    = _pick(row, 'reason_hindi', 'reason_hi')

                    # list names
                    list_1_name_en = _pick(row, 'list_1_name')
                    list_2_name_en = _pick(row, 'list_2_name')
                    list_1_name_hi = _pick(row, 'list_1_name_hi', 'list_1_name_hindi')
                    list_2_name_hi = _pick(row, 'list_2_name_hi', 'list_2_name_hindi')

                    # dynamic rows (EN/HI)
                    list1_en = {f'list_1_row{i}': _pick(row, f'list_1_row{i}') for i in range(1, 9)}
                    list1_hi = {f'list_1_row{i}_hi': _pick(row, f'list_1_row{i}_hi', f'list_1_row{i}_hindi') for i in range(1, 9)}
                    list2_en = {f'list_2_row{i}': _pick(row, f'list_2_row{i}') for i in range(1, 10)}   # includes row9
                    list2_hi = {f'list_2_row{i}_hi': _pick(row, f'list_2_row{i}_hi', f'list_2_row{i}_hindi') for i in range(1, 10)}
                    stmt_en  = {f'stmt_line_row{i}': _pick(row, f'stmt_line_row{i}', f'statement_row{i}', f'statement_{i}') for i in range(1, 10)}
                    stmt_hi  = {f'stmt_line_row{i}_hi': _pick(row, f'stmt_line_row{i}_hi', f'statement_row{i}_hi', f'statement_{i}_hi', f'statement_row{i}_hindi') for i in range(1, 10)}

                    correct_choice = _pick(row, 'correct_answer_choice')
                    sol_en = _pick(row, 'correct_answer_description')
                    sol_hi = _pick(row, 'correct_answer_description_hindi', 'correct_answer_description_hi')

                    marks_val = safe_float(_pick(row, 'exam_year1_marks', 'marks'))
                    neg_val   = safe_float(_pick(row, 'exam_year1_negative_marks', 'negative_marks'))

                    # ---- other scalar fields ----
                    elim_tactics_degree = _pick(row, 'elim_tactics_degree')
                    current_relevance   = _pick(row, 'current_relevance')
                    current_rel_topic   = _pick(row, 'current_relevance_topic')
                    script_en = _pick(row, 'script', 'script_en')
                    script_hi = _pick(row, 'script_hindi', 'script_hi')

                    # ---------- create EN ----------
                    payload_en = {
                        'question_number': next_qno,
                        'type_of_question': type_of_question,
                        'exam_stage': _pick(row, 'exam_stage'),
                        'exam_year': exam_year,
                        'language': 'e',
                        'script': script_en,
                        'marks': marks_val,
                        'negative_marks': neg_val,
                        'degree_of_difficulty': _pick(row, 'degree_of_difficulty'),
                        'elim_tactics_degree': elim_tactics_degree,
                        'current_relevance': current_relevance,
                        'current_relevance_topic': current_rel_topic,
                        'question_sub_type': question_sub_type,

                        'question_part': _pick(row, 'question_part'),
                        'question_part_first': qpf_en,
                        'question_part_third': qpt_en,
                        'assertion': assertion_en,
                        'reason': reason_en,

                        'list_1_name': list_1_name_en,
                        'list_2_name': list_2_name_en,

                        'answer_option_a': ans_a,
                        'answer_option_b': ans_b,
                        'answer_option_c': ans_c,
                        'answer_option_d': ans_d,

                        'correct_answer_choice': correct_choice,
                        'correct_answer_description': sol_en,
                        'created_by': request.user,
                    }
                    payload_en.update(_only_allowed(list1_en, allowed))
                    payload_en.update(_only_allowed(list2_en, allowed))
                    payload_en.update(_only_allowed(stmt_en, allowed))

                    q_en = QuestionBank.objects.create(**payload_en)
                    # set M2Ms
                    if areas_qs.exists():      q_en.area_name.set(areas_qs)
                    if sections_qs.exists():   q_en.section_name.set(sections_qs)
                    if parts_qs.exists():      q_en.part_name.set(parts_qs)
                    if chapters_qs.exists():   q_en.chapter_name.set(chapters_qs)
                    if topics_qs.exists():     q_en.topic_name.set(topics_qs)
                    if subtopics_qs.exists():  q_en.subtopic_name.set(subtopics_qs)
                    if exams_qs.exists():      q_en.exam_name.set(exams_qs)
                    if evergreen_qs.exists():  q_en.evergreenindex_name.set(evergreen_qs)

                    # ---------- create HI (same base_question_id & question_number) ----------
                    payload_hi = {
                        'base_question_id': q_en.base_question_id,
                        'question_number': next_qno,
                        'type_of_question': type_of_question,
                        'exam_stage': _pick(row, 'exam_stage'),
                        'exam_year': exam_year,
                        'language': 'h',
                        'script': script_hi,
                        'marks': marks_val,
                        'negative_marks': neg_val,
                        'degree_of_difficulty': _pick(row, 'degree_of_difficulty'),
                        'elim_tactics_degree': elim_tactics_degree,
                        'current_relevance': current_relevance,
                        'current_relevance_topic': current_rel_topic,
                        'question_sub_type': question_sub_type,

                        'question_part': _pick(row, 'question_part_hindi', 'question_part_hi'),
                        'question_part_first': qpf_hi,
                        'question_part_third': qpt_hi,
                        'assertion': assertion_hi,
                        'reason': reason_hi,

                        'list_1_name': list_1_name_hi,
                        'list_2_name': list_2_name_hi,

                        'answer_option_a_hi': ans_a_hi,
                        'answer_option_b_hi': ans_b_hi,
                        'answer_option_c_hi': ans_c_hi,
                        'answer_option_d_hi': ans_d_hi,

                        'correct_answer_description_hi': sol_hi,
                        'created_by': request.user,
                    }
                    payload_hi.update(_only_allowed(list1_hi, allowed))
                    payload_hi.update(_only_allowed(list2_hi, allowed))
                    payload_hi.update(_only_allowed(stmt_hi, allowed))

                    q_hi = QuestionBank.objects.create(**payload_hi)

                    # set M2Ms (same as EN)
                    if areas_qs.exists():      q_hi.area_name.set(areas_qs)
                    if sections_qs.exists():   q_hi.section_name.set(sections_qs)
                    if parts_qs.exists():      q_hi.part_name.set(parts_qs)
                    if chapters_qs.exists():   q_hi.chapter_name.set(chapters_qs)
                    if topics_qs.exists():     q_hi.topic_name.set(topics_qs)
                    if subtopics_qs.exists():  q_hi.subtopic_name.set(subtopics_qs)
                    if exams_qs.exists():      q_hi.exam_name.set(exams_qs)
                    if evergreen_qs.exists():  q_hi.evergreenindex_name.set(evergreen_qs)

                    # increment for next pair
                    next_qno += 1

            messages.success(request, "File uploaded and questions created successfully!")
            return redirect('upload-file')
    else:
        form = UploadFileForm()

    return render(request, 'question_bank/upload.html', {'form': form})


from django.http import JsonResponse
from django.views.decorators.http import require_GET
from .models import Area, Section, PartName, ChapterName, TopicName, SubTopicName, ExamName

# ✅ Helper function for consistent JSON responses
def json_response(data, status=200):
    return JsonResponse(data, status=status)

# ✅ Get all Areas (used when loading the form)
@require_GET
def get_areas(request):
    areas = Area.objects.values('area_SI_Code', 'name')
    return json_response({'areas': list(areas)})

# ✅ Get Areas by selected Subject(s) (optional - if subject filtering is needed)
@require_GET
def get_areas_list(request):
    subject_ids = request.GET.getlist('subject_ids[]')
    if not subject_ids:
        return json_response({'areas': []})
    areas = Area.objects.filter(subject_name__id__in=subject_ids).values('id', 'name')
    return json_response({'areas': list(areas)})

# ✅ Get Sections by selected Area(s)
@require_GET
def get_sections_list(request):
    area_ids = [str(a) for a in request.GET.getlist('area_ids[]')]
    if not area_ids:
        return json_response({'sections': []})
    sections = Section.objects.filter(area__area_SI_Code__in=area_ids).values('section_Unit_SI', 'name')
    return json_response({'sections': list(sections)})

# ✅ Get Parts by selected Section(s)
@require_GET
def get_parts_list(request):
    section_ids = [str(s) for s in request.GET.getlist('section_ids[]')]
    if not section_ids:
        return json_response({'parts': []})
    parts = PartName.objects.filter(section__section_Unit_SI__in=section_ids).values('part_serial', 'name')
    return json_response({'parts': list(parts)})

# ✅ Get Chapters by selected Part(s)
@require_GET
def get_chapters_list(request):
    part_ids = [str(p) for p in request.GET.getlist('part_ids[]')]
    if not part_ids:
        return json_response({'chapters': []})
    chapters = ChapterName.objects.filter(part__part_serial__in=part_ids).values('chapter_number', 'name')
    return json_response({'chapters': list(chapters)})

# ✅ Get Topics by selected Chapter(s)
@require_GET
def get_topics_list(request):
    chapter_ids = [str(c) for c in request.GET.getlist('chapter_ids[]')]
    if not chapter_ids:
        return json_response({'topics': []})
    topics = TopicName.objects.filter(chapter__chapter_number__in=chapter_ids).values('topic_SI_number', 'name')
    return json_response({'topics': list(topics)})

# ✅ Get Subtopics by selected Topic(s)
@require_GET
def get_subtopics_list(request):
    topic_ids = [str(t) for t in request.GET.getlist('topic_ids[]')]
    if not topic_ids:
        return json_response({'subtopics': []})

    subtopics = SubTopicName.objects.filter(
        topic__topic_SI_number__in=topic_ids
    ).values('sub_topic_SI_Number', 'name')

    return json_response({'subtopics': list(subtopics)})


# ✅ Get Exams by selected Subtopic(s)
@require_GET
def get_exams_list(request):
    subtopic_ids = [str(t) for t in request.GET.getlist('subtopic_ids[]')]
    if not subtopic_ids:
        return json_response({'related_exams': [], 'all_exams': []})

    related_exams_qs = ExamName.objects.filter(
        subtopics__sub_topic_SI_Number__in=subtopic_ids
    ).distinct()

    all_exams_qs = ExamName.objects.all()

    related_exams = [
        {
            'exam_SI_Number': exam.exam_SI_Number,
            'name': exam.name,
            'exam_code': exam.exam_code,
        }
        for exam in related_exams_qs
    ]

    all_exams = [
        {
            'exam_SI_Number': exam.exam_SI_Number,
            'name': exam.name,
            'exam_code': exam.exam_code,
        }
        for exam in all_exams_qs
    ]

    return json_response({
        'related_exams': related_exams,
        'all_exams': all_exams
    })


# ✅ NEW: Get Parts directly by selected Area(s) for "Modern" questions

@require_GET
def get_parts_by_area(request):
    area_ids = request.GET.getlist('area_ids[]')
    if not area_ids:
        return JsonResponse({'parts': []})

    parts = PartName.objects.filter(section__area__area_SI_Code__in=area_ids).values('part_serial', 'name')
    return JsonResponse({'parts': list(parts)})

from django.views.decorators.http import require_GET
from django.http import JsonResponse
from .models import HashtagsName, SubTopicName

# ✅ Helper for JSON
def json_response(data, status=200):
    return JsonResponse(data, status=status)

from django.shortcuts import render, get_object_or_404
from django.views.decorators.http import require_GET
from .models import HashtagsName, InputSuggestion
from django.http import JsonResponse

# Helper function for JSON response
def json_response(data):
    return JsonResponse(data, safe=True)


# ✅ Get Hashtags by selected Subtopic(s)
@require_GET
def get_hashtags_list(request):
    subtopic_ids = [str(t) for t in request.GET.getlist('subtopic_ids[]')]
    if not subtopic_ids:
        return json_response({'related_hashtags': [], 'all_hashtags': []})

    # Get hashtags related to selected subtopics
    related_hashtags_qs = HashtagsName.objects.filter(
        subtopics__sub_topic_SI_Number__in=subtopic_ids
    ).order_by("name").distinct()  # distinct ensures no duplicates

    # Get all hashtags (distinct)
    all_hashtags_qs = HashtagsName.objects.all().order_by("name").distinct()

    # Serialize queryset
    related_hashtags = [
        {'hashtags_SI_Number': h.hashtags_SI_Number, 'name': h.name}
        for h in related_hashtags_qs
    ]
    all_hashtags = [
        {'hashtags_SI_Number': h.hashtags_SI_Number, 'name': h.name}
        for h in all_hashtags_qs
    ]

    return json_response({
        'related_hashtags': related_hashtags,
        'all_hashtags': all_hashtags
    })


# Example view: hashtag_detail
def hashtag_detail(request, slug):
    hashtag = get_object_or_404(HashtagsName, hashtags_SI_Number=slug)
    suggestions = InputSuggestion.objects.filter(hashtags=hashtag).distinct()
    return render(request, 'question_bank/hashtag_detail.html', {
        'hashtag': hashtag,
        'suggestions': suggestions,
    })


# ✅ Add to your views.py or a separate api_views.py
from django.views.decorators.http import require_GET
from .models import SubTopicName
from django.http import JsonResponse

from django.views.decorators.http import require_GET
from django.http import JsonResponse
from .models import SubTopicName

@require_GET
def get_hierarchy_from_shortcode(request):
    """
    Accepts multiple shortcodes via:
      - ?codes[]=CODE1&codes[]=CODE2
      - or ?code=CODE1,CODE2
      - or space separated ?code=CODE1 CODE2
    Returns merged hierarchy lists for multi-select filling.
    """
    # ✅ Collect all codes
    codes = request.GET.getlist('codes[]')
    if not codes:
        code_param = request.GET.get("code", "").strip()
        if code_param:
            codes = [c.strip() for c in code_param.replace(',', ' ').split() if c.strip()]

    if not codes:
        return JsonResponse({'status': 'error', 'message': 'No code(s) provided'}, status=400)

    # ✅ Query subtopics
    subtopics = SubTopicName.objects.select_related(
        'topic__chapter__part__section__area'
    ).filter(sub_topic_short_Code__in=codes)

    if not subtopics.exists():
        return JsonResponse({'status': 'error', 'message': 'No matching subtopics found'}, status=404)

    # ✅ Helper to build unique lists while preserving order
    def unique_list(seq):
        seen = set()
        result = []
        for item in seq:
            if item not in seen:
                seen.add(item)
                result.append(item)
        return result

    # ✅ Build merged lists
    area_ids, area_names = [], []
    section_ids, section_names = [], []
    part_ids, part_names = [], []
    chapter_ids, chapter_names = [], []
    topic_ids, topic_names = [], []
    subtopic_ids, subtopic_names = [], []

    for st in subtopics:
        # Area
        area_ids.append(st.topic.chapter.part.section.area.area_SI_Code)
        area_names.append(st.topic.chapter.part.section.area.name)
        # Section
        section_ids.append(st.topic.chapter.part.section.section_Unit_SI)
        section_names.append(st.topic.chapter.part.section.name)
        # Part
        part_ids.append(st.topic.chapter.part.part_serial)
        part_names.append(st.topic.chapter.part.name)
        # Chapter
        chapter_ids.append(st.topic.chapter.chapter_number)
        chapter_names.append(st.topic.chapter.name)
        # Topic
        topic_ids.append(st.topic.topic_SI_number)
        topic_names.append(st.topic.name)
        # Subtopic
        subtopic_ids.append(st.sub_topic_SI_Number)
        subtopic_names.append(st.name)

    data = {
        'status': 'success',
        'data': {
            'area_ids': unique_list(area_ids),
            'area_names': unique_list(area_names),
            'section_ids': unique_list(section_ids),
            'section_names': unique_list(section_names),
            'part_ids': unique_list(part_ids),
            'part_names': unique_list(part_names),
            'chapter_ids': unique_list(chapter_ids),
            'chapter_names': unique_list(chapter_names),
            'topic_ids': unique_list(topic_ids),
            'topic_names': unique_list(topic_names),
            'subtopic_ids': unique_list(subtopic_ids),
            'subtopic_names': unique_list(subtopic_names),
        }
    }

    return JsonResponse(data)

from django.shortcuts import render, redirect
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.db import transaction
from .models import (
    Area, Section, PartName, ChapterName, TopicName, SubTopicName,
    QuestionBank, ExamName   # ✅ import ExamName
)
import random

@login_required
def add_simple_type_question(request):
    return save_simple_type_question(request)

from django.shortcuts import render, redirect
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.db import transaction
from .models import (
    Area, Section, PartName, ChapterName, TopicName, SubTopicName,
    QuestionBank, ExamName
)
import random

@login_required
def add_simple_type_question(request):
    return save_simple_type_question(request)

@login_required
def save_simple_type_question(request, question=None):
    # ---------- Defaults for preselected data (from shortcode) ----------
    preselected = {
        'area_ids': [],
        'section_ids': [],
        'part_ids': [],
        'chapter_ids': [],
        'topic_ids': [],
        'subtopic_ids': [],
    }

    # Multi-shortcode handling via GET
    codes = request.GET.getlist('codes[]')
    if not codes:
        shortcode_param = (request.GET.get('shortcode') or '').strip()
        if shortcode_param:
            codes = [c.strip() for c in shortcode_param.replace(',', ' ').split() if c.strip()]

    # Preselect hierarchy when editing/adding via shortcode(s)
    if codes:
        subtopics = SubTopicName.objects.select_related(
            'topic__chapter__part__section__area'
        ).filter(sub_topic_short_Code__in=codes)

        if subtopics.exists():
            preselected['area_ids']      = list({st.topic.chapter.part.section.area.area_SI_Code for st in subtopics})
            preselected['section_ids']   = list({st.topic.chapter.part.section.section_Unit_SI for st in subtopics})
            preselected['part_ids']      = list({st.topic.chapter.part.part_serial for st in subtopics})
            preselected['chapter_ids']   = list({st.topic.chapter.chapter_number for st in subtopics})
            preselected['topic_ids']     = list({st.topic.topic_SI_number for st in subtopics})
            preselected['subtopic_ids']  = list({st.sub_topic_SI_Number for st in subtopics})
        else:
            messages.warning(request, "No matching subtopics found for provided shortcode(s).")

    # ---------- Handle POST ----------
    if request.method == "POST":
        try:
            # Core fields
            question_type        = request.POST.get('questionType', 'simple_type')
            languages            = request.POST.getlist('language')
            script               = request.POST.get('script', '')

            # Question text
            question_part_first    = request.POST.get('question_part_first', '')
            question_part_first_hi = request.POST.get('question_part_first_hi', '')

            # Options
            answer_option_a    = request.POST.get('answer_option_a', '')
            answer_option_a_hi = request.POST.get('answer_option_a_hi', '')
            answer_option_b    = request.POST.get('answer_option_b', '')
            answer_option_b_hi = request.POST.get('answer_option_b_hi', '')
            answer_option_c    = request.POST.get('answer_option_c', '')
            answer_option_c_hi = request.POST.get('answer_option_c_hi', '')
            answer_option_d    = request.POST.get('answer_option_d', '')
            answer_option_d_hi = request.POST.get('answer_option_d_hi', '')

            # Answer meta
            correct_answer_choice        = request.POST.get('correct_answer_choice', '')
            correct_answer_description   = request.POST.get('correct_answer_description', '')
            correct_answer_description_hi= request.POST.get('correct_answer_description_hi', '')

            # Hierarchy (accept both legacy & new names for subtopic)
            area_ids     = request.POST.getlist('area_name[]')
            section_ids  = request.POST.getlist('section_name[]')
            part_ids     = request.POST.getlist('part_name[]')
            chapter_ids  = request.POST.getlist('chapter_name[]')
            topic_ids    = request.POST.getlist('topic_name[]')
            subtopic_ids = request.POST.getlist('subtopic_name[]') or request.POST.getlist('sub_topic_name[]')

            # Exams
            exam_ids = request.POST.getlist('exam_name[]')

            # Scoring / meta
            marks                 = float(request.POST.get('marks') or 0.0)
            negative_marks        = float(request.POST.get('negative_marks') or 0.0)
            degree_of_difficulty  = (request.POST.get('degree_of_difficulty') or '').strip()
            elim_tactics_degree   = (request.POST.get('elim_tactics_degree') or '').strip()
            # evergreen_index       = (request.POST.get('evergreen_index') or '').strip()
            current_relevance     = (request.POST.get('current_relevance') or '').strip()
            current_relevance_topic = request.POST.get('current_relevance_topic', '')

            # Source → type_of_question
            source = (request.POST.get('question_type_source') or '').strip()
            if source in ('pyq', 'osq'):
                type_of_question = source
            else:
                # treat anything else (incl. 'modern') as MO(Q)
                type_of_question = 'moq'

            # Exam year (optional int)
            exam_year_raw = request.POST.get('exam_year')
            if exam_year_raw in (None, '', 'null'):
                exam_year = None
            else:
                try:
                    exam_year = int(exam_year_raw)
                except ValueError:
                    messages.error(request, "Please provide a valid exam year.")
                    return redirect(request.path_info)

            # Basic validation
            if not languages or not question_part_first:
                messages.error(request, "Please fill in all required fields.")
                return redirect(request.path_info)

            # ---------- DB writes ----------
            with transaction.atomic():
                # Resolve M2M querysets
                areas     = Area.objects.filter(area_SI_Code__in=area_ids)
                sections  = Section.objects.filter(section_Unit_SI__in=section_ids)
                parts     = PartName.objects.filter(part_serial__in=part_ids)
                chapters  = ChapterName.objects.filter(chapter_number__in=chapter_ids)
                topics    = TopicName.objects.filter(topic_SI_number__in=topic_ids)
                subtopics = SubTopicName.objects.filter(sub_topic_SI_Number__in=subtopic_ids)
                exams     = ExamName.objects.filter(exam_SI_Number__in=exam_ids)

                # Next question number
                last_q = QuestionBank.objects.order_by('-question_number').first()
                next_question_number = (last_q.question_number + 1) if last_q else 1

                # Unique base_question_id
                while True:
                    base_question_id = random.randint(100000, 999999)
                    if not QuestionBank.objects.filter(base_question_id=base_question_id).exists():
                        break

                created_questions = []

                for lang in languages:
                    is_en = (lang == 'e')
                    is_hi = (lang == 'h')

                    q = QuestionBank.objects.create(
                        question_number=next_question_number,
                        question_sub_type=question_type,
                        base_question_id=base_question_id,
                        language=lang,
                        script=script,

                        # Question text (per language)
                        question_part_first = question_part_first if is_en else '',
                        question_part_first_hi = question_part_first_hi if is_hi else '',

                        # Options (per language)
                        answer_option_a = answer_option_a if is_en else '',
                        answer_option_a_hi = answer_option_a_hi if is_hi else '',
                        answer_option_b = answer_option_b if is_en else '',
                        answer_option_b_hi = answer_option_b_hi if is_hi else '',
                        answer_option_c = answer_option_c if is_en else '',
                        answer_option_c_hi = answer_option_c_hi if is_hi else '',
                        answer_option_d = answer_option_d if is_en else '',
                        answer_option_d_hi = answer_option_d_hi if is_hi else '',

                        # Answers
                        correct_answer_choice = correct_answer_choice,
                        correct_answer_description = correct_answer_description if is_en else '',
                        correct_answer_description_hi = correct_answer_description_hi if is_hi else '',

                        # Scoring/meta
                        marks=marks,
                        negative_marks=negative_marks,
                        degree_of_difficulty=degree_of_difficulty,
                        elim_tactics_degree=elim_tactics_degree,
                        # evergreen_index=evergreen_index,
                        current_relevance=current_relevance,
                        current_relevance_topic=current_relevance_topic,
                        exam_year=exam_year,

                        # Source
                        type_of_question=type_of_question,

                        created_by=request.user,
                    )

                    # Set M2M
                    q.area_name.set(areas)
                    q.section_name.set(sections)
                    q.part_name.set(parts)
                    q.chapter_name.set(chapters)
                    q.topic_name.set(topics)
                    q.subtopic_name.set(subtopics)
                    q.exam_name.set(exams)

                    created_questions.append(q)

                messages.success(request, "Question(s) added successfully!")
                return redirect('view_questions')

        except Exception as e:
            messages.error(request, f"An error occurred: {e}")
            return redirect(request.path_info)

    # ---------- Render form (GET) ----------
    template = 'question_bank/add_question/simple_type_form.html'
    return render(request, template, {
        'question': question,
        'areas': Area.objects.all(),
        'sections': Section.objects.all(),
        'parts': PartName.objects.all(),
        'chapters': ChapterName.objects.all(),
        'topics': TopicName.objects.all(),
        'subtopics': SubTopicName.objects.all(),
        **preselected,  # Preselect values based on shortcode if passed
    })



# ************************* Edit Simple Type Question Start *********************************************
from django.shortcuts import get_object_or_404, render, redirect
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.db import transaction
from .models import (
    Area, Section, PartName, ChapterName, TopicName, SubTopicName, ExamName, QuestionBank
)

@login_required
def edit_simple_type_question(request, pk):
    english_q = get_object_or_404(QuestionBank, pk=pk, language='e')
    hindi_q = QuestionBank.objects.filter(base_question_id=english_q.base_question_id, language='h').first()

    areas = Area.objects.all()
    sections = Section.objects.all()
    parts = PartName.objects.all()
    chapters = ChapterName.objects.all()
    topics = TopicName.objects.all()
    subtopics = SubTopicName.objects.all()
    exams = ExamName.objects.all()

    if request.method == "POST":
        try:
            with transaction.atomic():
                question_type = request.POST.get('questionType', 'simple_type')
                script = request.POST.get('script', '')

                question_part_first = request.POST.get('question_part_first', '')
                question_part_first_hi = request.POST.get('question_part_first_hi', '')

                answer_option_a = request.POST.get('answer_option_a', '')
                answer_option_a_hi = request.POST.get('answer_option_a_hi', '')

                answer_option_b = request.POST.get('answer_option_b', '')
                answer_option_b_hi = request.POST.get('answer_option_b_hi', '')

                answer_option_c = request.POST.get('answer_option_c', '')
                answer_option_c_hi = request.POST.get('answer_option_c_hi', '')

                answer_option_d = request.POST.get('answer_option_d', '')
                answer_option_d_hi = request.POST.get('answer_option_d_hi', '')

                correct_answer_choice = request.POST.get('correct_answer_choice', '')
                correct_answer_description = request.POST.get('correct_answer_description', '')
                correct_answer_description_hi = request.POST.get('correct_answer_description_hi', '')

                area_ids = request.POST.getlist('area_name[]')
                section_ids = request.POST.getlist('section_name[]')
                part_ids = request.POST.getlist('part_name[]')
                chapter_ids = request.POST.getlist('chapter_name[]')
                topic_ids = request.POST.getlist('topic_name[]')
                subtopic_ids = request.POST.getlist('sub_topic_name[]')
                exam_ids = request.POST.getlist('exam_name[]')

                new_topic_name = request.POST.get('new_topic_name', None)

                marks = float(request.POST.get('marks') or 0.0)
                negative_marks = float(request.POST.get('negative_marks') or 0.0)
                degree_of_difficulty = request.POST.get('degree_of_difficulty', '')
                evergreen_index = request.POST.get('evergreen_index', '')
                new_or_pyq = request.POST.get('new_or_pyq', '')

                current_relevance = request.POST.get('current_relevance', '')
                current_relevance_topic = request.POST.get('current_relevance_topic', '')
                exam_year = request.POST.get('exam_year') or None

                # ✅ New topic creation if 'other' selected
                if 'other' in topic_ids:
                    topic_ids.remove('other')
                    if new_topic_name and chapter_ids:
                        selected_chapter = ChapterName.objects.get(chapter_number=chapter_ids[0])
                        new_topic = TopicName.objects.create(name=new_topic_name, chapter=selected_chapter)
                        topic_ids.append(new_topic.topic_SI_number)

                # ✅ Related M2M
                areas_qs = Area.objects.filter(area_SI_Code__in=area_ids)
                sections_qs = Section.objects.filter(section_Unit_SI__in=section_ids)
                parts_qs = PartName.objects.filter(part_serial__in=part_ids)
                chapters_qs = ChapterName.objects.filter(chapter_number__in=chapter_ids)
                topics_qs = TopicName.objects.filter(topic_SI_number__in=topic_ids)
                subtopics_qs = SubTopicName.objects.filter(sub_topic_SI_Number__in=subtopic_ids)
                exams_qs = ExamName.objects.filter(exam_SI_Number__in=exam_ids)

                # ✅ Update English version
                english_q.script = script
                english_q.question_sub_type = question_type
                english_q.question_part_first = question_part_first
                english_q.answer_option_a = answer_option_a
                english_q.answer_option_b = answer_option_b
                english_q.answer_option_c = answer_option_c
                english_q.answer_option_d = answer_option_d
                english_q.correct_answer_choice = correct_answer_choice
                english_q.correct_answer_description = correct_answer_description
                english_q.marks = marks
                english_q.negative_marks = negative_marks
                english_q.degree_of_difficulty = degree_of_difficulty
                english_q.evergreen_index = evergreen_index
                english_q.exam_year = exam_year
                english_q.new_or_pyq = new_or_pyq
                english_q.current_relevance = current_relevance
                english_q.current_relevance_topic = current_relevance_topic
                english_q.save()

                english_q.area_name.set(areas_qs)
                english_q.section_name.set(sections_qs)
                english_q.part_name.set(parts_qs)
                english_q.chapter_name.set(chapters_qs)
                english_q.topic_name.set(topics_qs)
                english_q.subtopic_name.set(subtopics_qs)
                english_q.exam_name.set(exams_qs)

                # ✅ Update Hindi version if exists
                if hindi_q:
                    hindi_q.script = script
                    hindi_q.question_sub_type = question_type
                    hindi_q.question_part_first_hi = question_part_first_hi
                    hindi_q.answer_option_a_hi = answer_option_a_hi
                    hindi_q.answer_option_b_hi = answer_option_b_hi
                    hindi_q.answer_option_c_hi = answer_option_c_hi
                    hindi_q.answer_option_d_hi = answer_option_d_hi
                    hindi_q.correct_answer_choice = correct_answer_choice
                    hindi_q.correct_answer_description_hi = correct_answer_description_hi
                    hindi_q.marks = marks
                    hindi_q.negative_marks = negative_marks
                    hindi_q.degree_of_difficulty = degree_of_difficulty
                    hindi_q.evergreen_index = evergreen_index
                    hindi_q.exam_year = exam_year
                    hindi_q.new_or_pyq = new_or_pyq
                    hindi_q.current_relevance = current_relevance
                    hindi_q.current_relevance_topic = current_relevance_topic
                    hindi_q.save()

                    hindi_q.area_name.set(areas_qs)
                    hindi_q.section_name.set(sections_qs)
                    hindi_q.part_name.set(parts_qs)
                    hindi_q.chapter_name.set(chapters_qs)
                    hindi_q.topic_name.set(topics_qs)
                    hindi_q.subtopic_name.set(subtopics_qs)
                    hindi_q.exam_name.set(exams_qs)

                messages.success(request, "✅ Simple Type Question updated successfully!")
                return redirect('view_questions')

        except Exception as e:
            messages.error(request, f"❌ An error occurred while editing: {e}")

    context = {
        'english_q': english_q,
        'hindi_q': hindi_q,
        'areas': areas,
        'sections': sections,
        'parts': parts,
        'chapters': chapters,
        'topics': topics,
        'subtopics': subtopics,
        'exams': exams,
        'selected_area_ids': [str(id) for id in english_q.area_name.values_list('area_SI_Code', flat=True)],
        'selected_section_ids': [str(id) for id in english_q.section_name.values_list('section_Unit_SI', flat=True)],
        'selected_part_ids': [str(id) for id in english_q.part_name.values_list('part_serial', flat=True)],
        'selected_chapter_ids': [str(id) for id in english_q.chapter_name.values_list('chapter_number', flat=True)],
        'selected_topic_ids': [str(id) for id in english_q.topic_name.values_list('topic_SI_number', flat=True)],
        'selected_subtopic_ids': [str(id) for id in english_q.subtopic_name.values_list('sub_topic_SI_Number', flat=True)],
        'selected_exam_ids': [str(id) for id in english_q.exam_name.values_list('exam_SI_Number', flat=True)],
        "mark_options": ["0.5", "1.0", "1.5", "2.0", "2.5", "3.0"],
        "negative_mark_options": ["1.0", "0.5", "0.33", "0.25", "0.2", "0"],
        "range_1_to_5": range(1, 6),
    }

    return render(request, 'question_bank/edit_question/edit_simple_type_form.html', context)

# ************************* Edit Simple Type Question End *********************************************

from django.contrib.auth.decorators import login_required
from django.shortcuts import get_object_or_404, redirect
from django.contrib import messages
from django.db import transaction
from .models import QuestionBank

# ************************* DELETE Simple Type Question *********************************************
@login_required
def delete_simple_type_question(request, pk):
    try:
        english_q = get_object_or_404(QuestionBank, pk=pk, language='e')
        hindi_q = QuestionBank.objects.filter(base_question_id=english_q.base_question_id, language='h').first()

        with transaction.atomic():
            if hindi_q:
                hindi_q.delete()
            english_q.delete()

        messages.success(request, "Simple Type Question deleted successfully!")
    except Exception as e:
        messages.error(request, f"An error occurred while deleting Simple Type Question: {e}")

    return redirect('view_questions')

# ************************* DELETE R and A Type Question *********************************************
@login_required
def delete_r_and_a_type_question(request, pk):
    try:
        english_q = get_object_or_404(QuestionBank, pk=pk, language='e')
        hindi_q = QuestionBank.objects.filter(base_question_id=english_q.base_question_id, language='h').first()

        with transaction.atomic():
            if hindi_q:
                hindi_q.delete()
            english_q.delete()

        messages.success(request, "R and A Type Question deleted successfully!")
    except Exception as e:
        messages.error(request, f"An error occurred while deleting R and A Type Question: {e}")

    return redirect('view_questions')

# ************************* DELETE List Type 1 Question *********************************************
@login_required
def delete_list_type_1_question(request, pk):
    try:
        english_q = get_object_or_404(QuestionBank, pk=pk, language='e')
        hindi_q = QuestionBank.objects.filter(base_question_id=english_q.base_question_id, language='h').first()

        with transaction.atomic():
            if hindi_q:
                hindi_q.delete()
            english_q.delete()

        messages.success(request, "List Type 1 Question deleted successfully!")
    except Exception as e:
        messages.error(request, f"An error occurred while deleting List Type 1 Question: {e}")

    return redirect('view_questions')

# ************************* DELETE List Type 2 Question *********************************************
@login_required
def delete_list_type_2_question(request, pk):
    try:
        english_q = get_object_or_404(QuestionBank, pk=pk, language='e')
        hindi_q = QuestionBank.objects.filter(base_question_id=english_q.base_question_id, language='h').first()

        with transaction.atomic():
            if hindi_q:
                hindi_q.delete()
            english_q.delete()

        messages.success(request, "List Type 2 Question deleted successfully!")
    except Exception as e:
        messages.error(request, f"An error occurred while deleting List Type 2 Question: {e}")

    return redirect('view_questions')

# ************************* DELETE True and False Type Question *********************************************
@login_required
def delete_true_and_false_type_question(request, pk):
    try:
        english_q = get_object_or_404(QuestionBank, pk=pk, language='e')
        hindi_q = QuestionBank.objects.filter(base_question_id=english_q.base_question_id, language='h').first()

        with transaction.atomic():
            if hindi_q:
                hindi_q.delete()
            english_q.delete()

        messages.success(request, "True and False Type Question deleted successfully!")
    except Exception as e:
        messages.error(request, f"An error occurred while deleting True and False Type Question: {e}")

    return redirect('view_questions')

# ************************* DELETE Fill in the Blank Type Question *********************************************
@login_required
def delete_fill_in_the_blank_question(request, pk):
    try:
        english_q = get_object_or_404(QuestionBank, pk=pk, language='e')
        hindi_q = QuestionBank.objects.filter(base_question_id=english_q.base_question_id, language='h').first()

        with transaction.atomic():
            if hindi_q:
                hindi_q.delete()
            english_q.delete()

        messages.success(request, "Fill in the Blank Type Question deleted successfully!")
    except Exception as e:
        messages.error(request, f"An error occurred while deleting Fill in the Blank Type Question: {e}")

    return redirect('view_questions')






# ************************* Create R and A Type Question Start *************************
from django.shortcuts import render, redirect
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.db import transaction
from .models import (
    Area, Section, PartName, ChapterName, TopicName, SubTopicName, ExamName, QuestionBank
)
import random

@login_required
def add_r_and_a_type_question(request, question=None):
    # ---------- Defaults for preselected data (from shortcode) ----------
    preselected = {
        'area_ids': [], 'section_ids': [], 'part_ids': [],
        'chapter_ids': [], 'topic_ids': [], 'subtopic_ids': [], 'exam_ids': []
    }

    # Multi-shortcode handling via GET
    codes = request.GET.getlist('codes[]')
    if not codes:
        shortcode_param = (request.GET.get('shortcode') or '').strip()
        if shortcode_param:
            codes = [c.strip() for c in shortcode_param.replace(',', ' ').split() if c.strip()]

    # Preselect hierarchy when adding via shortcode(s)
    if codes:
        subtopics = SubTopicName.objects.select_related(
            'topic__chapter__part__section__area'
        ).filter(sub_topic_short_Code__in=codes)

        if subtopics.exists():
            preselected['area_ids']     = list({st.topic.chapter.part.section.area.area_SI_Code for st in subtopics})
            preselected['section_ids']  = list({st.topic.chapter.part.section.section_Unit_SI for st in subtopics})
            preselected['part_ids']     = list({st.topic.chapter.part.part_serial for st in subtopics})
            preselected['chapter_ids']  = list({st.topic.chapter.chapter_number for st in subtopics})
            preselected['topic_ids']    = list({st.topic.topic_SI_number for st in subtopics})
            preselected['subtopic_ids'] = list({st.sub_topic_SI_Number for st in subtopics})
        else:
            messages.warning(request, "No matching subtopics found for provided shortcode(s).")

    # ---------- Handle POST ----------
    if request.method == "POST":
        try:
            # Hierarchy (accept both legacy & new names for subtopic)
            area_ids     = request.POST.getlist('area_name[]')
            section_ids  = request.POST.getlist('section_name[]')
            part_ids     = request.POST.getlist('part_name[]')
            chapter_ids  = request.POST.getlist('chapter_name[]')
            topic_ids    = request.POST.getlist('topic_name[]')
            subtopic_ids = request.POST.getlist('subtopic_name[]') or request.POST.getlist('sub_topic_name[]')
            exam_ids     = request.POST.getlist('exam_name[]')

            languages = request.POST.getlist('language')

            # Question content
            script                = request.POST.get('script', '')
            question_part_first   = request.POST.get('question_part_first', '')
            question_part_first_hi= request.POST.get('question_part_first_hi', '')
            assertion             = request.POST.get('assertion', '')
            assertion_hi          = request.POST.get('assertion_hi', '')
            reason                = request.POST.get('reason', '')
            reason_hi             = request.POST.get('reason_hi', '')

            # Options & answers
            correct_answer_choice        = request.POST.get('correct_answer_choice', '')
            correct_answer_description   = request.POST.get('correct_answer_description', '')
            correct_answer_description_hi= request.POST.get('correct_answer_description_hi', '')

            answer_option_a    = request.POST.get('answer_option_a', '')
            answer_option_b    = request.POST.get('answer_option_b', '')
            answer_option_c    = request.POST.get('answer_option_c', '')
            answer_option_d    = request.POST.get('answer_option_d', '')
            answer_option_a_hi = request.POST.get('answer_option_a_hi', '')
            answer_option_b_hi = request.POST.get('answer_option_b_hi', '')
            answer_option_c_hi = request.POST.get('answer_option_c_hi', '')
            answer_option_d_hi = request.POST.get('answer_option_d_hi', '')

            # Scoring / meta
            marks                = float(request.POST.get('marks') or 0)
            negative_marks       = float(request.POST.get('negative_marks') or 0)
            degree_of_difficulty = (request.POST.get('degree_of_difficulty') or '').strip()
            elim_tactics_degree  = (request.POST.get('elim_tactics_degree') or '').strip()
            # evergreen_index      = (request.POST.get('evergreen_index') or '').strip()

            # Source → type_of_question
            question_type_source = (request.POST.get('question_type_source') or '').strip()
            if question_type_source in ('pyq', 'osq'):
                type_of_question = question_type_source
            else:
                type_of_question = 'moq'

            # Current relevance
            current_relevance       = (request.POST.get('current_relevance') or '').strip()
            current_relevance_topic = request.POST.get('current_relevance_topic', '')

            # Exam year (optional int)
            exam_year_raw = request.POST.get('exam_year')
            if exam_year_raw in (None, '', 'null'):
                exam_year = None
            else:
                try:
                    exam_year = int(exam_year_raw)
                except ValueError:
                    messages.error(request, "Please provide a valid exam year.")
                    return redirect(request.path_info)

            # Basic validation
            if not languages or not question_part_first:
                messages.error(request, "Please fill in all required fields.")
                return redirect(request.path_info)

            # ---------- DB writes ----------
            with transaction.atomic():
                # Resolve M2M querysets
                areas_qs     = Area.objects.filter(area_SI_Code__in=area_ids)
                sections_qs  = Section.objects.filter(section_Unit_SI__in=section_ids)
                parts_qs     = PartName.objects.filter(part_serial__in=part_ids)
                chapters_qs  = ChapterName.objects.filter(chapter_number__in=chapter_ids)
                topics_qs    = TopicName.objects.filter(topic_SI_number__in=topic_ids)
                subtopics_qs = SubTopicName.objects.filter(sub_topic_SI_Number__in=subtopic_ids)
                exams_qs     = ExamName.objects.filter(exam_SI_Number__in=exam_ids)

                # Unique IDs
                while True:
                    base_question_id = random.randint(100000, 999999)
                    if not QuestionBank.objects.filter(base_question_id=base_question_id).exists():
                        break

                last_q = QuestionBank.objects.order_by('-question_number').first()
                next_question_number = last_q.question_number + 1 if last_q else 1

                # Create per-language rows
                for lang in languages:
                    is_en = (lang == 'e')
                    is_hi = (lang == 'h')

                    q = QuestionBank.objects.create(
                        question_number=next_question_number,
                        base_question_id=base_question_id,
                        question_sub_type='r_and_a_type',
                        language=lang,
                        script=script,

                        # Content (per language)
                        question_part_first   = question_part_first if is_en else '',
                        question_part_first_hi= question_part_first_hi if is_hi else '',
                        assertion             = assertion if is_en else '',
                        assertion_hi          = assertion_hi if is_hi else '',
                        reason                = reason if is_en else '',
                        reason_hi             = reason_hi if is_hi else '',

                        # Options (per language)
                        answer_option_a    = answer_option_a if is_en else '',
                        answer_option_b    = answer_option_b if is_en else '',
                        answer_option_c    = answer_option_c if is_en else '',
                        answer_option_d    = answer_option_d if is_en else '',
                        answer_option_a_hi = answer_option_a_hi if is_hi else '',
                        answer_option_b_hi = answer_option_b_hi if is_hi else '',
                        answer_option_c_hi = answer_option_c_hi if is_hi else '',
                        answer_option_d_hi = answer_option_d_hi if is_hi else '',

                        # Answers
                        correct_answer_choice        = correct_answer_choice,
                        correct_answer_description   = correct_answer_description if is_en else '',
                        correct_answer_description_hi= correct_answer_description_hi if is_hi else '',

                        # Scoring / meta
                        marks=marks,
                        negative_marks=negative_marks,
                        degree_of_difficulty=degree_of_difficulty,
                        elim_tactics_degree=elim_tactics_degree,
                        # evergreen_index=evergreen_index,
                        current_relevance=current_relevance,
                        current_relevance_topic=current_relevance_topic,  # store same text for both langs
                        exam_year=exam_year,

                        # Source
                        type_of_question=type_of_question,

                        created_by=request.user
                    )

                    # M2M sets
                    q.area_name.set(areas_qs)
                    q.section_name.set(sections_qs)
                    q.part_name.set(parts_qs)
                    q.chapter_name.set(chapters_qs)
                    q.topic_name.set(topics_qs)
                    q.subtopic_name.set(subtopics_qs)
                    q.exam_name.set(exams_qs)

            messages.success(request, "✅ R & A Type Question saved successfully!")
            return redirect('add-r-and-a-type-question')

        except Exception as e:
            messages.error(request, f"❌ Error: {e}")

    # ---------- Render (GET) ----------
    context = {
        'areas': Area.objects.all(),
        'sections': Section.objects.all(),
        'parts': PartName.objects.all(),
        'chapters': ChapterName.objects.all(),
        'topics': TopicName.objects.all(),
        'subtopics': SubTopicName.objects.all(),
        'exams': ExamName.objects.all(),
        **preselected,
    }
    return render(request, 'question_bank/add_question/r_and_a_type_form.html', context)
# ************************* Create R and A Type Question End *************************






# ************************* Edit R and A Type Question Start *********************************************
from django.shortcuts import get_object_or_404, render, redirect
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.db import transaction
from .models import (
    Area, Section, PartName, ChapterName, TopicName, SubTopicName, ExamName, QuestionBank
)

@login_required
def edit_r_and_a_type_question(request, pk):
    english_q = get_object_or_404(QuestionBank, pk=pk, language='e')
    hindi_q = QuestionBank.objects.filter(base_question_id=english_q.base_question_id, language='h').first()

    areas = Area.objects.all()
    sections = Section.objects.all()
    parts = PartName.objects.all()
    chapters = ChapterName.objects.all()
    topics = TopicName.objects.all()
    subtopics = SubTopicName.objects.all()
    exams = ExamName.objects.all()

    if request.method == 'POST':
        try:
            with transaction.atomic():
                # ✅ IDs
                area_ids = request.POST.getlist('area_name[]')
                section_ids = request.POST.getlist('section_name[]')
                part_ids = request.POST.getlist('part_name[]')
                chapter_ids = request.POST.getlist('chapter_name[]')
                topic_ids = request.POST.getlist('topic_name[]')
                subtopic_ids = request.POST.getlist('sub_topic_name[]')
                exam_ids = request.POST.getlist('exam_name[]')
                new_topic_name = request.POST.get('new_topic_name', None)

                # ✅ Fields
                script = request.POST.get('script', '')
                question_part_first = request.POST.get('question_part_first', '')
                question_part_first_hi = request.POST.get('question_part_first_hi', '')

                reason = request.POST.get('reason', '')
                reason_hi = request.POST.get('reason_hi', '')
                assertion = request.POST.get('assertion', '')
                assertion_hi = request.POST.get('assertion_hi', '')
                question_part_third = request.POST.get('question_part_third', '')
                question_part_third_hi = request.POST.get('question_part_third_hi', '')

                correct_answer_choice = request.POST.get('correct_answer_choice', '')
                correct_answer_description = request.POST.get('correct_answer_description', '')
                correct_answer_description_hi = request.POST.get('correct_answer_description_hi', '')

                answer_option_a = request.POST.get('answer_option_a', '')
                answer_option_a_hi = request.POST.get('answer_option_a_hi', '')
                answer_option_b = request.POST.get('answer_option_b', '')
                answer_option_b_hi = request.POST.get('answer_option_b_hi', '')
                answer_option_c = request.POST.get('answer_option_c', '')
                answer_option_c_hi = request.POST.get('answer_option_c_hi', '')
                answer_option_d = request.POST.get('answer_option_d', '')
                answer_option_d_hi = request.POST.get('answer_option_d_hi', '')

                marks = float(request.POST.get('marks') or 0.0)
                negative_marks = float(request.POST.get('negative_marks') or 0.0)
                degree_of_difficulty = request.POST.get('degree_of_difficulty', '')
                evergreen_index = request.POST.get('evergreen_index', '')
                current_relevance = request.POST.get('current_relevance', '')
                current_relevance_topic = request.POST.get('current_relevance_topic', '')
                exam_year = request.POST.get('exam_year') or None

                # ✅ New topic if 'other'
                if 'other' in topic_ids:
                    topic_ids.remove('other')
                    if new_topic_name and chapter_ids:
                        selected_chapter = ChapterName.objects.get(chapter_number=chapter_ids[0])
                        new_topic = TopicName.objects.create(name=new_topic_name, chapter=selected_chapter)
                        topic_ids.append(new_topic.topic_SI_number)

                # ✅ Related objects
                areas_qs = Area.objects.filter(area_SI_Code__in=area_ids)
                sections_qs = Section.objects.filter(section_Unit_SI__in=section_ids)
                parts_qs = PartName.objects.filter(part_serial__in=part_ids)
                chapters_qs = ChapterName.objects.filter(chapter_number__in=chapter_ids)
                topics_qs = TopicName.objects.filter(topic_SI_number__in=topic_ids)
                subtopics_qs = SubTopicName.objects.filter(sub_topic_SI_Number__in=subtopic_ids)
                exams_qs = ExamName.objects.filter(exam_SI_Number__in=exam_ids)

                # ✅ Update English
                english_q.script = script
                english_q.question_part_first = question_part_first
                english_q.reason = reason
                english_q.assertion = assertion
                english_q.question_part_third = question_part_third
                english_q.correct_answer_choice = correct_answer_choice
                english_q.correct_answer_description = correct_answer_description
                english_q.answer_option_a = answer_option_a
                english_q.answer_option_b = answer_option_b
                english_q.answer_option_c = answer_option_c
                english_q.answer_option_d = answer_option_d
                english_q.marks = marks
                english_q.negative_marks = negative_marks
                english_q.degree_of_difficulty = degree_of_difficulty
                english_q.evergreen_index = evergreen_index
                english_q.exam_year = exam_year
                english_q.current_relevance = current_relevance
                english_q.current_relevance_topic = current_relevance_topic
                english_q.save()

                english_q.area_name.set(areas_qs)
                english_q.section_name.set(sections_qs)
                english_q.part_name.set(parts_qs)
                english_q.chapter_name.set(chapters_qs)
                english_q.topic_name.set(topics_qs)
                english_q.subtopic_name.set(subtopics_qs)
                english_q.exam_name.set(exams_qs)

                # ✅ Update Hindi
                if hindi_q:
                    hindi_q.script = script
                    hindi_q.question_part_first_hi = question_part_first_hi
                    hindi_q.reason_hi = reason_hi
                    hindi_q.assertion_hi = assertion_hi
                    hindi_q.question_part_third_hi = question_part_third_hi
                    hindi_q.correct_answer_choice = correct_answer_choice
                    hindi_q.correct_answer_description_hi = correct_answer_description_hi
                    hindi_q.answer_option_a_hi = answer_option_a_hi
                    hindi_q.answer_option_b_hi = answer_option_b_hi
                    hindi_q.answer_option_c_hi = answer_option_c_hi
                    hindi_q.answer_option_d_hi = answer_option_d_hi
                    hindi_q.marks = marks
                    hindi_q.negative_marks = negative_marks
                    hindi_q.degree_of_difficulty = degree_of_difficulty
                    hindi_q.evergreen_index = evergreen_index
                    hindi_q.exam_year = exam_year
                    hindi_q.current_relevance = current_relevance
                    hindi_q.current_relevance_topic = current_relevance_topic
                    hindi_q.save()

                    hindi_q.area_name.set(areas_qs)
                    hindi_q.section_name.set(sections_qs)
                    hindi_q.part_name.set(parts_qs)
                    hindi_q.chapter_name.set(chapters_qs)
                    hindi_q.topic_name.set(topics_qs)
                    hindi_q.subtopic_name.set(subtopics_qs)
                    hindi_q.exam_name.set(exams_qs)

                messages.success(request, "✅ R & A Type Question updated successfully!")
                return redirect('view_questions')

        except Exception as e:
            messages.error(request, f"❌ An error occurred while editing: {e}")

    context = {
        'english_q': english_q,
        'hindi_q': hindi_q,
        'areas': areas,
        'sections': sections,
        'parts': parts,
        'chapters': chapters,
        'topics': topics,
        'subtopics': subtopics,
        'exams': exams,
        'selected_area_ids': [str(id) for id in english_q.area_name.values_list('area_SI_Code', flat=True)],
        'selected_section_ids': [str(id) for id in english_q.section_name.values_list('section_Unit_SI', flat=True)],
        'selected_part_ids': [str(id) for id in english_q.part_name.values_list('part_serial', flat=True)],
        'selected_chapter_ids': [str(id) for id in english_q.chapter_name.values_list('chapter_number', flat=True)],
        'selected_topic_ids': [str(id) for id in english_q.topic_name.values_list('topic_SI_number', flat=True)],
        'selected_subtopic_ids': [str(id) for id in english_q.subtopic_name.values_list('sub_topic_SI_Number', flat=True)],
        'selected_exam_ids': [str(id) for id in english_q.exam_name.values_list('exam_SI_Number', flat=True)],
        "mark_options": ["0.5", "1.0", "1.5", "2.0", "2.5", "3.0"],
        "negative_mark_options": ["1.0", "0.5", "0.33", "0.25", "0.2", "0"],
        "range_1_to_5": range(1, 6),
    }

    return render(request, 'question_bank/edit_question/edit_r_and_a_type_form.html', context)
# ************************* Edit R and A Type Question End *********************************************





# ************************* Create List-I Type Question Start *************************
from django.shortcuts import render, redirect
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.db import transaction
from .models import (
    Area, Section, PartName, ChapterName, TopicName, SubTopicName, ExamName, QuestionBank
)
import random

@login_required
def add_list_type_1_question(request):
    context = {
        'areas': Area.objects.all(),
        'sections': Section.objects.all(),
        'parts': PartName.objects.all(),
        'chapters': ChapterName.objects.all(),
        'topics': TopicName.objects.all(),
        'subtopics': SubTopicName.objects.all(),
    }

    if request.method == 'POST':
        try:
            with transaction.atomic():
                # ----- Hierarchy selections -----
                area_ids    = request.POST.getlist('area_name[]')
                section_ids = request.POST.getlist('section_name[]')
                part_ids    = request.POST.getlist('part_name[]')
                chapter_ids = request.POST.getlist('chapter_name[]')
                topic_ids   = request.POST.getlist('topic_name[]')

                # ✅ Accept both legacy and new param names for subtopics
                subtopic_ids = request.POST.getlist('sub_topic_name[]') or request.POST.getlist('subtopic_name[]')
                subtopic_ids = [str(s).strip() for s in subtopic_ids if str(s).strip()]

                # ✅ Exams (M2M)
                exam_ids = [str(x).strip() for x in request.POST.getlist('exam_name[]') if str(x).strip()]

                new_topic_name = request.POST.get('new_topic_name')
                script = request.POST.get('script', '')

                # Validate/parse exam year
                exam_year_raw = request.POST.get('exam_year')
                if exam_year_raw in (None, '', 'null'):
                    exam_year = None
                else:
                    try:
                        exam_year = int(exam_year_raw)
                    except ValueError:
                        messages.error(request, "Please provide a valid exam year.")
                        return redirect(request.path_info)

                # ----- Stems / text -----
                question_part_first     = request.POST.get('question_part_first', '')
                question_part_first_hi  = request.POST.get('question_part_first_hi', '')
                question_part_third     = request.POST.get('question_part_third', '')
                question_part_third_hi  = request.POST.get('question_part_third_hi', '')

                # ----- Answers -----
                correct_answer_choice        = request.POST.get('correct_answer_choice', '')
                correct_answer_description   = request.POST.get('correct_answer_description', '')
                correct_answer_description_hi= request.POST.get('correct_answer_description_hi', '')

                # ----- Meta -----
                marks                 = float(request.POST.get('marks') or 0.0)
                negative_marks        = float(request.POST.get('negative_marks') or 0.0)
                degree_of_difficulty  = (request.POST.get('degree_of_difficulty') or '').strip()
                elim_tactics_degree   = (request.POST.get('elim_tactics_degree') or '').strip()  # ✅ NEW
                evergreen_index       = (request.POST.get('evergreen_index') or '').strip()
                question_type_source  = (request.POST.get('question_type_source') or 'modern').strip()
                current_relevance     = (request.POST.get('current_relevance') or '').strip()
                current_relevance_topic = request.POST.get('current_relevance_topic', '')

                # Map source to type
                if question_type_source == 'pyq':
                    type_of_question = 'pyq'
                elif question_type_source == 'osq':
                    type_of_question = 'osq'
                else:
                    type_of_question = 'moq'

                # ----- M2M QuerySets -----
                areas_qs    = Area.objects.filter(area_SI_Code__in=area_ids)
                sections_qs = Section.objects.filter(section_Unit_SI__in=section_ids)
                parts_qs    = PartName.objects.filter(part_serial__in=part_ids)
                chapters_qs = ChapterName.objects.filter(chapter_number__in=chapter_ids)
                topics_qs   = TopicName.objects.filter(topic_SI_number__in=topic_ids)

                # Handle new topic creation if "other" selected
                if 'other' in topic_ids and new_topic_name and chapter_ids:
                    topic_ids.remove('other')
                    selected_chapter = ChapterName.objects.get(chapter_number=chapter_ids[0])
                    new_topic = TopicName.objects.create(name=new_topic_name, chapter=selected_chapter)
                    topic_ids.append(str(new_topic.topic_SI_number))

                # Refresh topics after potential creation
                topics_qs = TopicName.objects.filter(topic_SI_number__in=topic_ids)

                # ✅ Subtopics queryset (works with either param name)
                subtopics_qs = SubTopicName.objects.filter(sub_topic_SI_Number__in=subtopic_ids)

                # ✅ Exams queryset
                exams_qs = ExamName.objects.filter(exam_SI_Number__in=exam_ids)

                # Guard — if you require subtopics
                if not subtopics_qs.exists():
                    messages.error(request, "Please choose at least one Sub Topic.")
                    return redirect(request.path_info)

                # ----- IDs / numbering -----
                while True:
                    base_question_id = random.randint(100000, 999999)
                    if not QuestionBank.objects.filter(base_question_id=base_question_id).exists():
                        break

                last_q = QuestionBank.objects.order_by('-question_number').first()
                next_question_number = last_q.question_number + 1 if last_q else 1

                # ----- List items (List-I) -----
                list_items = {}
                for i in range(1, 9):
                    list_items[f"list_1_row{i}"]    = request.POST.get(f"list_1_row{i}", '')
                    list_items[f"list_1_row{i}_hi"] = request.POST.get(f"list_1_row{i}_hi", '')

                # ----- Options A-D -----
                answer_options = {}
                for opt in 'abcd':
                    answer_options[f"answer_option_{opt}"]     = request.POST.get(f"answer_option_{opt}", '')
                    answer_options[f"answer_option_{opt}_hi"]  = request.POST.get(f"answer_option_{opt}_hi", '')

                # ----- Create English question -----
                question_en = QuestionBank.objects.create(
                    question_number=next_question_number,
                    base_question_id=base_question_id,
                    question_sub_type='list_type_1',
                    type_of_question=type_of_question,
                    language='e',
                    script=script,
                    question_part_first=question_part_first,
                    question_part_third=question_part_third,
                    correct_answer_choice=correct_answer_choice,
                    correct_answer_description=correct_answer_description,
                    marks=marks,
                    negative_marks=negative_marks,
                    degree_of_difficulty=degree_of_difficulty,
                    elim_tactics_degree=elim_tactics_degree,  # ✅ NEW
                    evergreen_index=evergreen_index,
                    exam_year=exam_year,
                    current_relevance=current_relevance,
                    current_relevance_topic=current_relevance_topic,
                    created_by=request.user,
                    **{f"list_1_row{i}": list_items[f"list_1_row{i}"] for i in range(1, 9)},
                    **{f"answer_option_{opt}": answer_options[f"answer_option_{opt}"] for opt in 'abcd'},
                )
                # M2M sets (subtopics + exams included)
                question_en.area_name.set(areas_qs)
                question_en.section_name.set(sections_qs)
                question_en.part_name.set(parts_qs)
                question_en.chapter_name.set(chapters_qs)
                question_en.topic_name.set(topics_qs)
                question_en.subtopic_name.set(subtopics_qs)
                question_en.exam_name.set(exams_qs)  # ✅ NEW

                # ----- Create Hindi question -----
                question_hi = QuestionBank.objects.create(
                    question_number=next_question_number,
                    base_question_id=base_question_id,
                    question_sub_type='list_type_1',
                    type_of_question=type_of_question,
                    language='h',
                    script=script,
                    question_part_first_hi=question_part_first_hi,
                    question_part_third_hi=question_part_third_hi,
                    correct_answer_choice=correct_answer_choice,
                    correct_answer_description_hi=correct_answer_description_hi,
                    marks=marks,
                    negative_marks=negative_marks,
                    degree_of_difficulty=degree_of_difficulty,
                    elim_tactics_degree=elim_tactics_degree,  # ✅ NEW
                    evergreen_index=evergreen_index,
                    exam_year=exam_year,
                    current_relevance=current_relevance,
                    current_relevance_topic=current_relevance_topic,
                    created_by=request.user,
                    **{f"list_1_row{i}_hi": list_items[f"list_1_row{i}_hi"] for i in range(1, 9)},
                    **{f"answer_option_{opt}_hi": answer_options[f"answer_option_{opt}_hi"] for opt in 'abcd'},
                )
                question_hi.area_name.set(areas_qs)
                question_hi.section_name.set(sections_qs)
                question_hi.part_name.set(parts_qs)
                question_hi.chapter_name.set(chapters_qs)
                question_hi.topic_name.set(topics_qs)
                question_hi.subtopic_name.set(subtopics_qs)
                question_hi.exam_name.set(exams_qs)  # ✅ NEW

                messages.success(request, '✅ List-I Type Question saved successfully!')
                return redirect('add-list-type-1-question')

        except Exception as e:
            messages.error(request, f'❌ An error occurred: {str(e)}')

    return render(request, 'question_bank/add_question/list_type_1_form.html', context)
# ************************* Create List-I Type Question End *************************






# ************************* Edit List-I Type Question Updated *********************************************
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.db import transaction
from .models import (
    Area, Section, PartName, ChapterName, TopicName, SubTopicName, ExamName, QuestionBank
)

@login_required
def edit_list_type_1_question(request, pk):
    english_q = get_object_or_404(QuestionBank, pk=pk, language='e')
    hindi_q = QuestionBank.objects.filter(
        base_question_id=english_q.base_question_id, language='h'
    ).first()

    areas = Area.objects.all()
    sections = Section.objects.all()
    parts = PartName.objects.all()
    chapters = ChapterName.objects.all()
    topics = TopicName.objects.all()
    subtopics = SubTopicName.objects.all()
    exams = ExamName.objects.all()  # ✅ include exams for select

    if request.method == 'POST':
        try:
            with transaction.atomic():
                # ===== Basic fields =====
                script = request.POST.get('script', '')
                exam_year = request.POST.get('exam_year') or None

                question_part_first = request.POST.get('question_part_first', '')
                question_part_first_hi = request.POST.get('question_part_first_hi', '')
                question_part_third = request.POST.get('question_part_third', '')
                question_part_third_hi = request.POST.get('question_part_third_hi', '')

                correct_answer_choice = request.POST.get('correct_answer_choice', '')
                correct_answer_description = request.POST.get('correct_answer_description', '')
                correct_answer_description_hi = request.POST.get('correct_answer_description_hi', '')

                marks = float(request.POST.get('marks') or 0.0)
                negative_marks = float(request.POST.get('negative_marks') or 0.0)

                # ✅ Difficulty & Elimination tactics (CharFields)
                degree_of_difficulty = (request.POST.get('degree_of_difficulty') or '').strip()
                elim_tactics_degree = (request.POST.get('elim_tactics_degree') or '').strip()

                # ✅ Current relevance (parity with Simple Type)
                current_relevance = (request.POST.get('current_relevance') or '').strip()
                current_relevance_topic = request.POST.get('current_relevance_topic', '')

                evergreen_index = request.POST.get('evergreen_index', '')
                type_of_question = (request.POST.get('question_type_source') or '').strip()

                # ===== Answer options and list items =====
                answer_options = {
                    f"answer_option_{opt}": request.POST.get(f"answer_option_{opt}", '')
                    for opt in "abcd"
                }
                answer_options_hi = {
                    f"answer_option_{opt}_hi": request.POST.get(f"answer_option_{opt}_hi", '')
                    for opt in "abcd"
                }

                list_items = {f"list_1_row{i}": request.POST.get(f"list_1_row{i}", '') for i in range(1, 9)}
                list_items_hi = {f"list_1_row{i}_hi": request.POST.get(f"list_1_row{i}_hi", '') for i in range(1, 9)}

                # ===== Relationships =====
                area_ids = request.POST.getlist('area_name[]')
                section_ids = request.POST.getlist('section_name[]')
                part_ids = request.POST.getlist('part_name[]')
                chapter_ids = request.POST.getlist('chapter_name[]')
                topic_ids = request.POST.getlist('topic_name[]')
                subtopic_ids = request.POST.getlist('sub_topic_name[]')
                exam_ids = request.POST.getlist('exam_name[]')  # ✅ exams

                areas_qs = Area.objects.filter(area_SI_Code__in=area_ids)
                sections_qs = Section.objects.filter(section_Unit_SI__in=section_ids)
                parts_qs = PartName.objects.filter(part_serial__in=part_ids)
                chapters_qs = ChapterName.objects.filter(chapter_number__in=chapter_ids)
                topics_qs = TopicName.objects.filter(topic_SI_number__in=topic_ids)
                subtopics_qs = SubTopicName.objects.filter(sub_topic_SI_Number__in=subtopic_ids)
                exams_qs = ExamName.objects.filter(exam_SI_Number__in=exam_ids)  # ✅

                # ===== Update English =====
                english_q.script = script
                english_q.exam_year = exam_year
                english_q.question_part_first = question_part_first
                english_q.question_part_third = question_part_third
                english_q.correct_answer_choice = correct_answer_choice
                english_q.correct_answer_description = correct_answer_description
                english_q.marks = marks
                english_q.negative_marks = negative_marks
                english_q.degree_of_difficulty = degree_of_difficulty
                english_q.elim_tactics_degree = elim_tactics_degree
                english_q.current_relevance = current_relevance            # ✅
                english_q.current_relevance_topic = current_relevance_topic # ✅
                english_q.evergreen_index = evergreen_index
                english_q.type_of_question = type_of_question

                for k, v in list_items.items():
                    setattr(english_q, k, v)
                for k, v in answer_options.items():
                    setattr(english_q, k, v)

                english_q.save()
                english_q.area_name.set(areas_qs)
                english_q.section_name.set(sections_qs)
                english_q.part_name.set(parts_qs)
                english_q.chapter_name.set(chapters_qs)
                english_q.topic_name.set(topics_qs)
                english_q.subtopic_name.set(subtopics_qs)
                english_q.exam_name.set(exams_qs)  # ✅

                # ===== Update Hindi (if exists) =====
                if hindi_q:
                    hindi_q.script = script
                    hindi_q.exam_year = exam_year
                    hindi_q.question_part_first_hi = question_part_first_hi
                    hindi_q.question_part_third_hi = question_part_third_hi
                    hindi_q.correct_answer_choice = correct_answer_choice
                    hindi_q.correct_answer_description_hi = correct_answer_description_hi
                    hindi_q.marks = marks
                    hindi_q.negative_marks = negative_marks
                    hindi_q.degree_of_difficulty = degree_of_difficulty
                    hindi_q.elim_tactics_degree = elim_tactics_degree
                    hindi_q.current_relevance = current_relevance            # ✅
                    hindi_q.current_relevance_topic = current_relevance_topic # ✅
                    hindi_q.evergreen_index = evergreen_index
                    hindi_q.type_of_question = type_of_question

                    for k, v in list_items_hi.items():
                        setattr(hindi_q, k, v)
                    for k, v in answer_options_hi.items():
                        setattr(hindi_q, k, v)

                    hindi_q.save()
                    hindi_q.area_name.set(areas_qs)
                    hindi_q.section_name.set(sections_qs)
                    hindi_q.part_name.set(parts_qs)
                    hindi_q.chapter_name.set(chapters_qs)
                    hindi_q.topic_name.set(topics_qs)
                    hindi_q.subtopic_name.set(subtopics_qs)
                    hindi_q.exam_name.set(exams_qs)  # ✅

                messages.success(request, "✅ List-I Type Question updated successfully!")
                return redirect('view_questions')  # ✅ fixed: underscore to match your Simple Type code

        except Exception as e:
            messages.error(request, f"❌ An error occurred while editing: {e}")

    # ===== GET: prefill context =====
    list_items = {f"list_1_row{i}": getattr(english_q, f"list_1_row{i}", '') for i in range(1, 9)}
    list_items_hi = {f"list_1_row{i}_hi": getattr(hindi_q, f"list_1_row{i}_hi", '') for i in range(1, 9)} if hindi_q else {}

    context = {
        'english_q': english_q,
        'hindi_q': hindi_q,
        'areas': areas,
        'sections': sections,
        'parts': parts,
        'chapters': chapters,
        'topics': topics,
        'subtopics': subtopics,
        'exams': exams,  # ✅ provide to template
        'list_items': list_items,
        'list_items_hi': list_items_hi,
        'selected_area_ids': [str(v) for v in english_q.area_name.values_list('area_SI_Code', flat=True)],
        'selected_section_ids': [str(v) for v in english_q.section_name.values_list('section_Unit_SI', flat=True)],
        'selected_part_ids': [str(v) for v in english_q.part_name.values_list('part_serial', flat=True)],
        'selected_chapter_ids': [str(v) for v in english_q.chapter_name.values_list('chapter_number', flat=True)],
        'selected_topic_ids': [str(v) for v in english_q.topic_name.values_list('topic_SI_number', flat=True)],
        'selected_subtopic_ids': [str(v) for v in english_q.subtopic_name.values_list('sub_topic_SI_Number', flat=True)],
        'selected_exam_ids': [str(v) for v in english_q.exam_name.values_list('exam_SI_Number', flat=True)],  # ✅
        "mark_options": ["0.5", "1.0", "1.5", "2.0", "2.5", "3.0"],
        "negative_mark_options": ["1.0", "0.5", "0.33", "0.25", "0.2", "0"],
        "range_1_to_5": range(1, 6),  # for degree_of_difficulty & elim_tactics_degree
        "range_1_to_8": range(1, 9),  # for list rows
    }

    return render(request, 'question_bank/edit_question/edit_list_type_1_form.html', context)
# ************************* Edit List-I Type Question Updated *********************************************

# ************************* Create Statement Type Question Start *************************
from django.shortcuts import render, redirect
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.db import transaction
from .models import (
    Area, Section, PartName, ChapterName, TopicName, SubTopicName, ExamName, QuestionBank
)
import random

@login_required
def add_statement_type_question(request):
    context = {
        'areas': Area.objects.all(),
        'sections': Section.objects.all(),
        'parts': PartName.objects.all(),
        'chapters': ChapterName.objects.all(),
        'topics': TopicName.objects.all(),
        'subtopics': SubTopicName.objects.all(),
    }

    if request.method == 'POST':
        try:
            with transaction.atomic():
                # ----- Hierarchy selections -----
                area_ids    = request.POST.getlist('area_name[]')
                section_ids = request.POST.getlist('section_name[]')
                part_ids    = request.POST.getlist('part_name[]')
                chapter_ids = request.POST.getlist('chapter_name[]')
                topic_ids   = request.POST.getlist('topic_name[]')

                # Accept both param names for subtopics
                subtopic_ids = request.POST.getlist('sub_topic_name[]') or request.POST.getlist('subtopic_name[]')
                subtopic_ids = [str(s).strip() for s in subtopic_ids if str(s).strip()]

                # Exams (M2M)
                exam_ids = [str(x).strip() for x in request.POST.getlist('exam_name[]') if str(x).strip()]

                new_topic_name = request.POST.get('new_topic_name')
                script = request.POST.get('script', '')

                # Exam year
                exam_year_raw = request.POST.get('exam_year')
                if exam_year_raw in (None, '', 'null'):
                    exam_year = None
                else:
                    try:
                        exam_year = int(exam_year_raw)
                    except ValueError:
                        messages.error(request, "Please provide a valid exam year.")
                        return redirect(request.path_info)

                # ----- Stems / text -----
                question_part_first     = request.POST.get('question_part_first', '')
                question_part_first_hi  = request.POST.get('question_part_first_hi', '')
                question_part_third     = request.POST.get('question_part_third', '')
                question_part_third_hi  = request.POST.get('question_part_third_hi', '')

                # ----- Answers (A-D) -----
                correct_answer_choice        = request.POST.get('correct_answer_choice', '')
                correct_answer_description   = request.POST.get('correct_answer_description', '')
                correct_answer_description_hi= request.POST.get('correct_answer_description_hi', '')

                # ----- Meta -----
                marks                 = float(request.POST.get('marks') or 0.0)
                negative_marks        = float(request.POST.get('negative_marks') or 0.0)
                degree_of_difficulty  = (request.POST.get('degree_of_difficulty') or '').strip()
                elim_tactics_degree   = (request.POST.get('elim_tactics_degree') or '').strip()
                # evergreen_index       = (request.POST.get('evergreen_index') or '').strip()
                question_type_source  = (request.POST.get('question_type_source') or 'modern').strip()
                current_relevance     = (request.POST.get('current_relevance') or '').strip()
                current_relevance_topic = request.POST.get('current_relevance_topic', '')

                # Map source -> type_of_question
                if question_type_source == 'pyq':
                    type_of_question = 'pyq'
                elif question_type_source == 'osq':
                    type_of_question = 'osq'
                else:
                    type_of_question = 'moq'

                # ----- M2M QuerySets -----
                areas_qs    = Area.objects.filter(area_SI_Code__in=area_ids)
                sections_qs = Section.objects.filter(section_Unit_SI__in=section_ids)
                parts_qs    = PartName.objects.filter(part_serial__in=part_ids)
                chapters_qs = ChapterName.objects.filter(chapter_number__in=chapter_ids)
                topics_qs   = TopicName.objects.filter(topic_SI_number__in=topic_ids)

                # Allow new topic creation if "other" chosen
                if 'other' in topic_ids and new_topic_name and chapter_ids:
                    topic_ids.remove('other')
                    selected_chapter = ChapterName.objects.get(chapter_number=chapter_ids[0])
                    new_topic = TopicName.objects.create(name=new_topic_name, chapter=selected_chapter)
                    topic_ids.append(str(new_topic.topic_SI_number))

                topics_qs = TopicName.objects.filter(topic_SI_number__in=topic_ids)

                subtopics_qs = SubTopicName.objects.filter(sub_topic_SI_Number__in=subtopic_ids)
                exams_qs     = ExamName.objects.filter(exam_SI_Number__in=exam_ids)

                if not subtopics_qs.exists():
                    messages.error(request, "Please choose at least one Sub Topic.")
                    return redirect(request.path_info)

                # ----- IDs / numbering -----
                while True:
                    base_question_id = random.randint(100000, 999999)
                    if not QuestionBank.objects.filter(base_question_id=base_question_id).exists():
                        break

                last_q = QuestionBank.objects.order_by('-question_number').first()
                next_question_number = last_q.question_number + 1 if last_q else 1

                # ----- Statement lines & List-II (rows 1..9) -----
                stmt_lines = {}
                list2_rows = {}
                for i in range(1, 10):  # 1..9
                    stmt_lines[f"stmt_line_row{i}"]     = request.POST.get(f"stmt_line_row{i}", '')
                    stmt_lines[f"stmt_line_row{i}_hi"]  = request.POST.get(f"stmt_line_row{i}_hi", '')
                    list2_rows[f"list_2_row{i}"]        = request.POST.get(f"list_2_row{i}", '')
                    list2_rows[f"list_2_row{i}_hi"]     = request.POST.get(f"list_2_row{i}_hi", '')

                # ----- Options A-D -----
                answer_options = {}
                for opt in 'abcd':
                    answer_options[f"answer_option_{opt}"]     = request.POST.get(f"answer_option_{opt}", '')
                    answer_options[f"answer_option_{opt}_hi"]  = request.POST.get(f"answer_option_{opt}_hi", '')

                # ===== Create English question =====
                question_en = QuestionBank.objects.create(
                    question_number=next_question_number,
                    base_question_id=base_question_id,
                    question_sub_type='statement_type',
                    type_of_question=type_of_question,
                    language='e',
                    script=script,
                    question_part_first=question_part_first,
                    question_part_third=question_part_third,
                    correct_answer_choice=correct_answer_choice,
                    correct_answer_description=correct_answer_description,
                    marks=marks,
                    negative_marks=negative_marks,
                    degree_of_difficulty=degree_of_difficulty,
                    elim_tactics_degree=elim_tactics_degree,
                    # evergreen_index=evergreen_index,
                    exam_year=exam_year,
                    current_relevance=current_relevance,
                    current_relevance_topic=current_relevance_topic,
                    created_by=request.user,
                    # stmt_line_row1..9 + list_2_row1..9
                    **{f"stmt_line_row{i}": stmt_lines[f"stmt_line_row{i}"] for i in range(1, 10)},
                    **{f"list_2_row{i}":  list2_rows[f"list_2_row{i}"]  for i in range(1, 10)},
                    # options
                    **{f"answer_option_{opt}": answer_options[f"answer_option_{opt}"] for opt in 'abcd'},
                )
                question_en.area_name.set(areas_qs)
                question_en.section_name.set(sections_qs)
                question_en.part_name.set(parts_qs)
                question_en.chapter_name.set(chapters_qs)
                question_en.topic_name.set(topics_qs)
                question_en.subtopic_name.set(subtopics_qs)
                question_en.exam_name.set(exams_qs)

                # ===== Create Hindi question =====
                question_hi = QuestionBank.objects.create(
                    question_number=next_question_number,
                    base_question_id=base_question_id,
                    question_sub_type='statement_type',
                    type_of_question=type_of_question,
                    language='h',
                    script=script,
                    question_part_first_hi=question_part_first_hi,
                    question_part_third_hi=question_part_third_hi,
                    correct_answer_choice=correct_answer_choice,
                    correct_answer_description_hi=correct_answer_description_hi,
                    marks=marks,
                    negative_marks=negative_marks,
                    degree_of_difficulty=degree_of_difficulty,
                    elim_tactics_degree=elim_tactics_degree,
                    # evergreen_index=evergreen_index,
                    exam_year=exam_year,
                    current_relevance=current_relevance,
                    current_relevance_topic=current_relevance_topic,
                    created_by=request.user,
                    # stmt_line_row1..9_hi + list_2_row1..9_hi
                    **{f"stmt_line_row{i}_hi": stmt_lines[f"stmt_line_row{i}_hi"] for i in range(1, 10)},
                    **{f"list_2_row{i}_hi":  list2_rows[f"list_2_row{i}_hi"]  for i in range(1, 10)},
                    # options
                    **{f"answer_option_{opt}_hi": answer_options[f"answer_option_{opt}_hi"] for opt in 'abcd'},
                )
                question_hi.area_name.set(areas_qs)
                question_hi.section_name.set(sections_qs)
                question_hi.part_name.set(parts_qs)
                question_hi.chapter_name.set(chapters_qs)
                question_hi.topic_name.set(topics_qs)
                question_hi.subtopic_name.set(subtopics_qs)
                question_hi.exam_name.set(exams_qs)

                messages.success(request, '✅ Statement Type Question saved successfully!')
                return redirect('add-statement-type-question')

        except Exception as e:
            messages.error(request, f'❌ An error occurred: {str(e)}')

    return render(request, 'question_bank/add_question/statement_type_form.html', context)
# ************************* Create Statement Type Question End *************************

# ************************* Edit Statement Type Question *********************************************
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.db import transaction
from .models import (
    Area, Section, PartName, ChapterName, TopicName, SubTopicName, ExamName, QuestionBank
)

@login_required
def edit_statement_type_question(request, pk):
    # Only English PK is passed; we pair it with the Hindi twin via base_question_id
    english_q = get_object_or_404(
        QuestionBank,
        pk=pk,
        language='e',
        question_sub_type='statement_type'
    )
    hindi_q = QuestionBank.objects.filter(
        base_question_id=english_q.base_question_id, language='h'
    ).first()

    areas = Area.objects.all()
    sections = Section.objects.all()
    parts = PartName.objects.all()
    chapters = ChapterName.objects.all()
    topics = TopicName.objects.all()
    subtopics = SubTopicName.objects.all()
    exams = ExamName.objects.all()

    if request.method == 'POST':
        try:
            with transaction.atomic():
                # ===== Basic fields =====
                script = request.POST.get('script', '')

                exam_year_raw = request.POST.get('exam_year')
                try:
                    exam_year = int(exam_year_raw) if exam_year_raw not in (None, '', 'null') else None
                except (TypeError, ValueError):
                    exam_year = None

                question_part_first     = request.POST.get('question_part_first', '')
                question_part_first_hi  = request.POST.get('question_part_first_hi', '')
                question_part_third     = request.POST.get('question_part_third', '')
                question_part_third_hi  = request.POST.get('question_part_third_hi', '')

                correct_answer_choice           = request.POST.get('correct_answer_choice', '')
                correct_answer_description      = request.POST.get('correct_answer_description', '')
                correct_answer_description_hi   = request.POST.get('correct_answer_description_hi', '')

                marks = float(request.POST.get('marks') or 0.0)
                negative_marks = float(request.POST.get('negative_marks') or 0.0)

                # CharFields (stored as strings, consistent with your other views)
                degree_of_difficulty = (request.POST.get('degree_of_difficulty') or '').strip()
                elim_tactics_degree  = (request.POST.get('elim_tactics_degree') or '').strip()
                current_relevance     = (request.POST.get('current_relevance') or '').strip()
                current_relevance_topic = request.POST.get('current_relevance_topic', '')

                # Optional, only set if your model actually has the field
                evergreen_index = request.POST.get('evergreen_index', '')

                # Source mapping (kept simple to mirror your other edit views)
                type_of_question = (request.POST.get('question_type_source') or english_q.type_of_question).strip()

                # ===== Answer options =====
                answer_options_en = {f"answer_option_{opt}": request.POST.get(f"answer_option_{opt}", '') for opt in "abcd"}
                answer_options_hi = {f"answer_option_{opt}_hi": request.POST.get(f"answer_option_{opt}_hi", '') for opt in "abcd"}

                # ===== Statement lines & List-II rows (1..9) =====
                stmt_lines_en = {f"stmt_line_row{i}": request.POST.get(f"stmt_line_row{i}", '') for i in range(1, 10)}
                stmt_lines_hi = {f"stmt_line_row{i}_hi": request.POST.get(f"stmt_line_row{i}_hi", '') for i in range(1, 10)}

                list2_rows_en = {f"list_2_row{i}": request.POST.get(f"list_2_row{i}", '') for i in range(1, 10)}
                list2_rows_hi = {f"list_2_row{i}_hi": request.POST.get(f"list_2_row{i}_hi", '') for i in range(1, 10)}

                # ===== Relationships =====
                area_ids    = request.POST.getlist('area_name[]')
                section_ids = request.POST.getlist('section_name[]')
                part_ids    = request.POST.getlist('part_name[]')
                chapter_ids = request.POST.getlist('chapter_name[]')
                topic_ids   = request.POST.getlist('topic_name[]')

                # support both names for subtopics
                subtopic_ids = request.POST.getlist('sub_topic_name[]') or request.POST.getlist('subtopic_name[]')
                exam_ids     = request.POST.getlist('exam_name[]')

                areas_qs    = Area.objects.filter(area_SI_Code__in=area_ids)
                sections_qs = Section.objects.filter(section_Unit_SI__in=section_ids)
                parts_qs    = PartName.objects.filter(part_serial__in=part_ids)
                chapters_qs = ChapterName.objects.filter(chapter_number__in=chapter_ids)
                topics_qs   = TopicName.objects.filter(topic_SI_number__in=topic_ids)
                subtopics_qs= SubTopicName.objects.filter(sub_topic_SI_Number__in=subtopic_ids)
                exams_qs    = ExamName.objects.filter(exam_SI_Number__in=exam_ids)

                # ===== Update English =====
                english_q.script = script
                english_q.exam_year = exam_year
                english_q.question_part_first = question_part_first
                english_q.question_part_third = question_part_third
                english_q.correct_answer_choice = correct_answer_choice
                english_q.correct_answer_description = correct_answer_description
                english_q.marks = marks
                english_q.negative_marks = negative_marks
                english_q.degree_of_difficulty = degree_of_difficulty
                english_q.elim_tactics_degree = elim_tactics_degree
                english_q.current_relevance = current_relevance
                english_q.current_relevance_topic = current_relevance_topic
                english_q.type_of_question = type_of_question

                # Set evergreen_index only if the field exists on the model
                if hasattr(english_q, 'evergreen_index'):
                    english_q.evergreen_index = evergreen_index

                # Bulk set statement lines and list-II rows
                for k, v in {**stmt_lines_en, **list2_rows_en, **answer_options_en}.items():
                    setattr(english_q, k, v)

                english_q.save()
                english_q.area_name.set(areas_qs)
                english_q.section_name.set(sections_qs)
                english_q.part_name.set(parts_qs)
                english_q.chapter_name.set(chapters_qs)
                english_q.topic_name.set(topics_qs)
                english_q.subtopic_name.set(subtopics_qs)
                english_q.exam_name.set(exams_qs)

                # ===== Update Hindi (if exists) =====
                if hindi_q:
                    hindi_q.script = script
                    hindi_q.exam_year = exam_year
                    hindi_q.question_part_first_hi = question_part_first_hi
                    hindi_q.question_part_third_hi = question_part_third_hi
                    hindi_q.correct_answer_choice = correct_answer_choice
                    hindi_q.correct_answer_description_hi = correct_answer_description_hi
                    hindi_q.marks = marks
                    hindi_q.negative_marks = negative_marks
                    hindi_q.degree_of_difficulty = degree_of_difficulty
                    hindi_q.elim_tactics_degree = elim_tactics_degree
                    hindi_q.current_relevance = current_relevance
                    hindi_q.current_relevance_topic = current_relevance_topic
                    hindi_q.type_of_question = type_of_question

                    if hasattr(hindi_q, 'evergreen_index'):
                        hindi_q.evergreen_index = evergreen_index

                    for k, v in {**stmt_lines_hi, **list2_rows_hi, **answer_options_hi}.items():
                        setattr(hindi_q, k, v)

                    hindi_q.save()
                    hindi_q.area_name.set(areas_qs)
                    hindi_q.section_name.set(sections_qs)
                    hindi_q.part_name.set(parts_qs)
                    hindi_q.chapter_name.set(chapters_qs)
                    hindi_q.topic_name.set(topics_qs)
                    hindi_q.subtopic_name.set(subtopics_qs)
                    hindi_q.exam_name.set(exams_qs)

                messages.success(request, "✅ Statement Type Question updated successfully!")
                return redirect('view_questions')

        except Exception as e:
            messages.error(request, f"❌ An error occurred while editing: {e}")

    # ===== GET: prefill context for the form =====
    stmt_lines_en = {f"stmt_line_row{i}": getattr(english_q, f"stmt_line_row{i}", '') for i in range(1, 10)}
    stmt_lines_hi = {f"stmt_line_row{i}_hi": (getattr(hindi_q, f"stmt_line_row{i}_hi", '') if hindi_q else '') for i in range(1, 10)}

    list2_rows_en = {f"list_2_row{i}": getattr(english_q, f"list_2_row{i}", '') for i in range(1, 10)}
    list2_rows_hi = {f"list_2_row{i}_hi": (getattr(hindi_q, f"list_2_row{i}_hi", '') if hindi_q else '') for i in range(1, 10)}

    answer_options_hi = {f"answer_option_{opt}_hi": (getattr(hindi_q, f"answer_option_{opt}_hi", '') if hindi_q else '') for opt in "abcd"}

    context = {
        'english_q': english_q,
        'hindi_q': hindi_q,
        'areas': areas,
        'sections': sections,
        'parts': parts,
        'chapters': chapters,
        'topics': topics,
        'subtopics': subtopics,
        'exams': exams,

        # prefill maps
        'stmt_lines_en': stmt_lines_en,
        'stmt_lines_hi': stmt_lines_hi,
        'list2_rows_en': list2_rows_en,
        'list2_rows_hi': list2_rows_hi,
        'answer_options_hi': answer_options_hi,

        # selected M2M ids (strings for template compatibility)
        'selected_area_ids':     [str(v) for v in english_q.area_name.values_list('area_SI_Code', flat=True)],
        'selected_section_ids':  [str(v) for v in english_q.section_name.values_list('section_Unit_SI', flat=True)],
        'selected_part_ids':     [str(v) for v in english_q.part_name.values_list('part_serial', flat=True)],
        'selected_chapter_ids':  [str(v) for v in english_q.chapter_name.values_list('chapter_number', flat=True)],
        'selected_topic_ids':    [str(v) for v in english_q.topic_name.values_list('topic_SI_number', flat=True)],
        'selected_subtopic_ids': [str(v) for v in english_q.subtopic_name.values_list('sub_topic_SI_Number', flat=True)],
        'selected_exam_ids':     [str(v) for v in english_q.exam_name.values_list('exam_SI_Number', flat=True)],

        "mark_options": ["0.5", "1.0", "1.5", "2.0", "2.5", "3.0"],
        "negative_mark_options": ["1.0", "0.5", "0.33", "0.25", "0.2", "0"],
        "range_1_to_5": range(1, 6),   # for degree_of_difficulty & elim_tactics_degree
        "range_1_to_9": range(1, 10),  # for statement/list-II rows
    }

    return render(request, 'question_bank/edit_question/edit_statement_type_form.html', context)
# ************************* Edit Statement Type Question *********************************************




# ************************* Create List-II Type Question Start *********************************************
from django.shortcuts import render, redirect
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.db import transaction
from .models import (
    Area, Section, PartName, ChapterName, TopicName, SubTopicName, ExamName, QuestionBank
)
import random

@login_required
def add_list_type_2_question(request):
    context = {
        'areas': Area.objects.all(),
        'sections': Section.objects.all(),
        'parts': PartName.objects.all(),
        'chapters': ChapterName.objects.all(),
        'topics': TopicName.objects.all(),
        'subtopics': SubTopicName.objects.all(),
    }

    if request.method == 'POST':
        try:
            with transaction.atomic():
                # ----- Hierarchy selections -----
                area_ids    = request.POST.getlist('area_name[]')
                section_ids = request.POST.getlist('section_name[]')
                part_ids    = request.POST.getlist('part_name[]')
                chapter_ids = request.POST.getlist('chapter_name[]')
                topic_ids   = request.POST.getlist('topic_name[]')

                # ✅ Accept both legacy and new param names for subtopics
                subtopic_ids = request.POST.getlist('sub_topic_name[]') or request.POST.getlist('subtopic_name[]')
                subtopic_ids = [str(s).strip() for s in subtopic_ids if str(s).strip()]

                # ✅ Exams (M2M)
                exam_ids = [str(x).strip() for x in request.POST.getlist('exam_name[]') if str(x).strip()]

                new_topic_name = request.POST.get('new_topic_name', '').strip()
                script = request.POST.get('script', '')

                # ----- Validate/parse exam year -----
                exam_year_raw = request.POST.get('exam_year')
                if exam_year_raw in (None, '', 'null'):
                    exam_year = None
                else:
                    try:
                        exam_year = int(exam_year_raw)
                    except ValueError:
                        messages.error(request, "Please provide a valid exam year.")
                        return redirect(request.path_info)

                # ----- Text / stems -----
                question_part_first     = request.POST.get('question_part_first', '')
                question_part_first_hi  = request.POST.get('question_part_first_hi', '')
                question_part_third     = request.POST.get('question_part_third', '')
                question_part_third_hi  = request.POST.get('question_part_third_hi', '')

                # ----- Answers -----
                correct_answer_choice        = request.POST.get('correct_answer_choice', '')
                correct_answer_description   = request.POST.get('correct_answer_description', '')
                correct_answer_description_hi= request.POST.get('correct_answer_description_hi', '')

                # ----- Meta -----
                marks                 = float(request.POST.get('marks') or 0.0)
                negative_marks        = float(request.POST.get('negative_marks') or 0.0)
                degree_of_difficulty  = (request.POST.get('degree_of_difficulty') or '').strip()
                elim_tactics_degree   = (request.POST.get('elim_tactics_degree') or '').strip()  # ✅ NEW
                evergreen_index       = (request.POST.get('evergreen_index') or '').strip()

                list_1_name    = request.POST.get('list_1_name', '')
                list_2_name    = request.POST.get('list_2_name', '')
                list_1_name_hi = request.POST.get('list_1_name_hi', '')
                list_2_name_hi = request.POST.get('list_2_name_hi', '')

                # ✅ Current Relevance
                current_relevance       = (request.POST.get('current_relevance') or '').strip()
                current_relevance_topic = request.POST.get('current_relevance_topic', '')

                # Question source -> type
                question_type_source = (request.POST.get('question_type_source', 'modern') or '').strip()
                if question_type_source == 'pyq':
                    type_of_question = 'pyq'
                elif question_type_source == 'osq':
                    type_of_question = 'osq'
                else:
                    type_of_question = 'moq'

                # ----- M2M QuerySets -----
                areas_qs    = Area.objects.filter(area_SI_Code__in=area_ids)
                sections_qs = Section.objects.filter(section_Unit_SI__in=section_ids)
                parts_qs    = PartName.objects.filter(part_serial__in=part_ids)
                chapters_qs = ChapterName.objects.filter(chapter_number__in=chapter_ids)

                # New topic handling (if "other")
                if 'other' in topic_ids and new_topic_name and chapter_ids:
                    topic_ids.remove('other')
                    selected_chapter = ChapterName.objects.get(chapter_number=chapter_ids[0])
                    new_topic = TopicName.objects.create(name=new_topic_name, chapter=selected_chapter)
                    topic_ids.append(str(new_topic.topic_SI_number))

                topics_qs    = TopicName.objects.filter(topic_SI_number__in=topic_ids)
                subtopics_qs = SubTopicName.objects.filter(sub_topic_SI_Number__in=subtopic_ids)
                exams_qs     = ExamName.objects.filter(exam_SI_Number__in=exam_ids)

                # (Optional) enforce subtopics selected
                if not subtopics_qs.exists():
                    messages.error(request, "Please choose at least one Sub Topic.")
                    return redirect(request.path_info)

                # ----- Generate identifiers -----
                last_question = QuestionBank.objects.order_by('-question_number').first()
                next_question_number = last_question.question_number + 1 if last_question else 1

                while True:
                    base_question_id = random.randint(100000, 999999)
                    if not QuestionBank.objects.filter(base_question_id=base_question_id).exists():
                        break

                # ----- List items & options -----
                list_items    = {}
                list_items_hi = {}
                for i in range(1, 6):
                    list_items[f"list_1_row{i}"]    = request.POST.get(f"list_1_row{i}", '')
                    list_items[f"list_2_row{i}"]    = request.POST.get(f"list_2_row{i}", '')
                    list_items_hi[f"list_1_row{i}_hi"] = request.POST.get(f"list_1_row{i}_hi", '')
                    list_items_hi[f"list_2_row{i}_hi"] = request.POST.get(f"list_2_row{i}_hi", '')

                answer_options    = {f"answer_option_{opt}": request.POST.get(f"answer_option_{opt}", '') for opt in 'abcd'}
                answer_options_hi = {f"answer_option_{opt}_hi": request.POST.get(f"answer_option_{opt}_hi", '') for opt in 'abcd'}

                # ----- Create EN row -----
                question_en = QuestionBank.objects.create(
                    question_number=next_question_number,
                    question_sub_type='list_type_2',
                    base_question_id=base_question_id,
                    type_of_question=type_of_question,
                    language='e',
                    script=script,
                    question_part_first=question_part_first,
                    question_part_third=question_part_third,
                    correct_answer_choice=correct_answer_choice,
                    correct_answer_description=correct_answer_description,
                    marks=marks,
                    negative_marks=negative_marks,
                    degree_of_difficulty=degree_of_difficulty,
                    elim_tactics_degree=elim_tactics_degree,  # ✅ NEW
                    evergreen_index=evergreen_index,
                    list_1_name=list_1_name,
                    list_2_name=list_2_name,
                    exam_year=exam_year,
                    current_relevance=current_relevance,
                    current_relevance_topic=current_relevance_topic,
                    created_by=request.user,
                    **list_items,
                    **answer_options
                )
                question_en.area_name.set(areas_qs)
                question_en.section_name.set(sections_qs)
                question_en.part_name.set(parts_qs)
                question_en.chapter_name.set(chapters_qs)
                question_en.topic_name.set(topics_qs)
                question_en.subtopic_name.set(subtopics_qs)
                question_en.exam_name.set(exams_qs)  # ✅ NEW

                # ----- Create HI row -----
                question_hi = QuestionBank.objects.create(
                    question_number=next_question_number,
                    question_sub_type='list_type_2',
                    base_question_id=base_question_id,
                    type_of_question=type_of_question,
                    language='h',
                    script=script,
                    question_part_first_hi=question_part_first_hi,
                    question_part_third_hi=question_part_third_hi,
                    correct_answer_choice=correct_answer_choice,
                    correct_answer_description_hi=correct_answer_description_hi,
                    marks=marks,
                    negative_marks=negative_marks,
                    degree_of_difficulty=degree_of_difficulty,
                    elim_tactics_degree=elim_tactics_degree,  # ✅ NEW
                    evergreen_index=evergreen_index,
                    list_1_name=list_1_name_hi,
                    list_2_name=list_2_name_hi,
                    exam_year=exam_year,
                    current_relevance=current_relevance,
                    current_relevance_topic=current_relevance_topic,
                    created_by=request.user,
                    **list_items_hi,
                    **answer_options_hi
                )
                question_hi.area_name.set(areas_qs)
                question_hi.section_name.set(sections_qs)
                question_hi.part_name.set(parts_qs)
                question_hi.chapter_name.set(chapters_qs)
                question_hi.topic_name.set(topics_qs)
                question_hi.subtopic_name.set(subtopics_qs)
                question_hi.exam_name.set(exams_qs)  # ✅ NEW

                messages.success(request, '✅ List-II Type Question added successfully!')
                return redirect('add-list-type-2-question')

        except Exception as e:
            messages.error(request, f'❌ An unexpected error occurred: {str(e)}')

    return render(request, 'question_bank/add_question/list_type_2_form.html', context)

# ************************* Create List-II Type Question End *********************************************



# ************************* Edit List-II Type Question Updated *********************************************
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.db import transaction
from .models import (
    Area, Section, PartName, ChapterName, TopicName, SubTopicName, ExamName, QuestionBank
)

@login_required
def edit_list_type_2_question(request, pk):
    english_q = get_object_or_404(QuestionBank, pk=pk, language='e')
    hindi_q = QuestionBank.objects.filter(base_question_id=english_q.base_question_id, language='h').first()

    areas = Area.objects.all()
    sections = Section.objects.all()
    parts = PartName.objects.all()
    chapters = ChapterName.objects.all()
    topics = TopicName.objects.all()
    subtopics = SubTopicName.objects.all()
    exams = ExamName.objects.all()

    if request.method == 'POST':
        try:
            with transaction.atomic():
                # ---------- Basics ----------
                script = request.POST.get('script', '')
                exam_year = request.POST.get('exam_year') or None

                question_part_first = request.POST.get('question_part_first', '')
                question_part_first_hi = request.POST.get('question_part_first_hi', '')

                question_part_third = request.POST.get('question_part_third', '')
                question_part_third_hi = request.POST.get('question_part_third_hi', '')

                correct_answer_choice = request.POST.get('correct_answer_choice', '')
                correct_answer_description = request.POST.get('correct_answer_description', '')
                correct_answer_description_hi = request.POST.get('correct_answer_description_hi', '')

                list_1_name = request.POST.get('list_1_name', '')
                list_2_name = request.POST.get('list_2_name', '')
                list_1_name_hi = request.POST.get('list_1_name_hi', '')
                list_2_name_hi = request.POST.get('list_2_name_hi', '')

                marks = float(request.POST.get('marks') or 0.0)
                negative_marks = float(request.POST.get('negative_marks') or 0.0)

                # ---------- Difficulty & Elimination tactics ----------
                degree_of_difficulty = (request.POST.get('degree_of_difficulty') or '').strip()
                elim_tactics_degree = (request.POST.get('elim_tactics_degree') or '').strip()

                # ---------- Other meta ----------
                evergreen_index = request.POST.get('evergreen_index', '')
                current_relevance = request.POST.get('current_relevance', '')
                current_relevance_topic = request.POST.get('current_relevance_topic', '')
                type_of_question = (request.POST.get('question_type_source') or '').strip()

                # ---------- Answers & lists ----------
                answer_options = {f"answer_option_{opt}": request.POST.get(f"answer_option_{opt}", '') for opt in "abcd"}
                answer_options_hi = {f"answer_option_{opt}_hi": request.POST.get(f"answer_option_{opt}_hi", '') for opt in "abcd"}

                list_items = {f"list_1_row{i}": request.POST.get(f"list_1_row{i}", '') for i in range(1, 6)}
                list_items.update({f"list_2_row{i}": request.POST.get(f"list_2_row{i}", '') for i in range(1, 6)})

                list_items_hi = {f"list_1_row{i}_hi": request.POST.get(f"list_1_row{i}_hi", '') for i in range(1, 6)}
                list_items_hi.update({f"list_2_row{i}_hi": request.POST.get(f"list_2_row{i}_hi", '') for i in range(1, 6)})

                # ---------- Relationships ----------
                area_ids = request.POST.getlist('area_name[]')
                section_ids = request.POST.getlist('section_name[]')
                part_ids = request.POST.getlist('part_name[]')
                chapter_ids = request.POST.getlist('chapter_name[]')
                topic_ids = request.POST.getlist('topic_name[]')
                subtopic_ids = request.POST.getlist('sub_topic_name[]')
                exam_ids = request.POST.getlist('exam_name[]')

                areas_qs = Area.objects.filter(area_SI_Code__in=area_ids)
                sections_qs = Section.objects.filter(section_Unit_SI__in=section_ids)
                parts_qs = PartName.objects.filter(part_serial__in=part_ids)
                chapters_qs = ChapterName.objects.filter(chapter_number__in=chapter_ids)
                topics_qs = TopicName.objects.filter(topic_SI_number__in=topic_ids)
                subtopics_qs = SubTopicName.objects.filter(sub_topic_SI_Number__in=subtopic_ids)
                exams_qs = ExamName.objects.filter(exam_SI_Number__in=exam_ids)

                # ---------- Update English ----------
                english_q.script = script
                english_q.question_part_first = question_part_first
                english_q.question_part_third = question_part_third
                english_q.correct_answer_choice = correct_answer_choice
                english_q.correct_answer_description = correct_answer_description
                english_q.list_1_name = list_1_name
                english_q.list_2_name = list_2_name
                english_q.marks = marks
                english_q.negative_marks = negative_marks
                english_q.degree_of_difficulty = degree_of_difficulty
                english_q.elim_tactics_degree = elim_tactics_degree
                english_q.evergreen_index = evergreen_index
                english_q.current_relevance = current_relevance
                english_q.current_relevance_topic = current_relevance_topic
                english_q.exam_year = exam_year
                english_q.type_of_question = type_of_question

                for k, v in list_items.items():
                    setattr(english_q, k, v)
                for k, v in answer_options.items():
                    setattr(english_q, k, v)

                english_q.save()
                english_q.area_name.set(areas_qs)
                english_q.section_name.set(sections_qs)
                english_q.part_name.set(parts_qs)
                english_q.chapter_name.set(chapters_qs)
                english_q.topic_name.set(topics_qs)
                english_q.subtopic_name.set(subtopics_qs)
                english_q.exam_name.set(exams_qs)

                # ---------- Update Hindi (if exists) ----------
                if hindi_q:
                    hindi_q.script = script
                    hindi_q.question_part_first_hi = question_part_first_hi
                    hindi_q.question_part_third_hi = question_part_third_hi
                    hindi_q.correct_answer_choice = correct_answer_choice
                    hindi_q.correct_answer_description_hi = correct_answer_description_hi
                    hindi_q.list_1_name = list_1_name_hi
                    hindi_q.list_2_name = list_2_name_hi
                    hindi_q.marks = marks
                    hindi_q.negative_marks = negative_marks
                    hindi_q.degree_of_difficulty = degree_of_difficulty
                    hindi_q.elim_tactics_degree = elim_tactics_degree
                    hindi_q.evergreen_index = evergreen_index
                    hindi_q.current_relevance = current_relevance
                    hindi_q.current_relevance_topic = current_relevance_topic
                    hindi_q.exam_year = exam_year
                    hindi_q.type_of_question = type_of_question

                    for k, v in list_items_hi.items():
                        setattr(hindi_q, k, v)
                    for k, v in answer_options_hi.items():
                        setattr(hindi_q, k, v)

                    hindi_q.save()
                    hindi_q.area_name.set(areas_qs)
                    hindi_q.section_name.set(sections_qs)
                    hindi_q.part_name.set(parts_qs)
                    hindi_q.chapter_name.set(chapters_qs)
                    hindi_q.topic_name.set(topics_qs)
                    hindi_q.subtopic_name.set(subtopics_qs)
                    hindi_q.exam_name.set(exams_qs)

                messages.success(request, "✅ List-II Type Question updated successfully!")
                return redirect('view_questions')

        except Exception as e:
            messages.error(request, f"❌ An error occurred while editing: {e}")

    # ---------- GET: prefill context ----------
    list_1_items = {f"list_1_row{i}": getattr(english_q, f"list_1_row{i}", '') for i in range(1, 6)}
    list_2_items = {f"list_2_row{i}": getattr(english_q, f"list_2_row{i}", '') for i in range(1, 6)}
    list_1_items_hi = {f"list_1_row{i}_hi": getattr(hindi_q, f"list_1_row{i}_hi", '') for i in range(1, 6)} if hindi_q else {}
    list_2_items_hi = {f"list_2_row{i}_hi": getattr(hindi_q, f"list_2_row{i}_hi", '') for i in range(1, 6)} if hindi_q else {}

    context = {
        'english_q': english_q,
        'hindi_q': hindi_q,
        'areas': areas,
        'sections': sections,
        'parts': parts,
        'chapters': chapters,
        'topics': topics,
        'subtopics': subtopics,
        'exams': exams,
        'list_1_items': list_1_items,
        'list_2_items': list_2_items,
        'list_1_items_hi': list_1_items_hi,
        'list_2_items_hi': list_2_items_hi,
        'selected_area_ids': [str(v) for v in english_q.area_name.values_list('area_SI_Code', flat=True)],
        'selected_section_ids': [str(v) for v in english_q.section_name.values_list('section_Unit_SI', flat=True)],
        'selected_part_ids': [str(v) for v in english_q.part_name.values_list('part_serial', flat=True)],
        'selected_chapter_ids': [str(v) for v in english_q.chapter_name.values_list('chapter_number', flat=True)],
        'selected_topic_ids': [str(v) for v in english_q.topic_name.values_list('topic_SI_number', flat=True)],
        'selected_subtopic_ids': [str(v) for v in english_q.subtopic_name.values_list('sub_topic_SI_Number', flat=True)],
        'selected_exam_ids': [str(v) for v in english_q.exam_name.values_list('exam_SI_Number', flat=True)],
        "mark_options": ["0.5", "1.0", "1.5", "2.0", "2.5", "3.0"],
        "negative_mark_options": ["1.0", "0.5", "0.33", "0.25", "0.2", "0"],
        "range_1_to_5": range(1, 6),  # for degree_of_difficulty & elim_tactics_degree selects
    }

    return render(request, 'question_bank/edit_question/edit_list_type_2_form.html', context)
# ************************* Edit List-II Type Question Updated *********************************************




# ************************* Create True and False Type Question Start *********************************************
from django.shortcuts import render, redirect
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.db import transaction
from .models import (
    Area, Section, PartName, ChapterName, TopicName, SubTopicName,
    QuestionBank, ExamName  # ✅ include ExamName
)
import random

@login_required
def add_true_and_false_type_question(request):
    context = {
        'areas': Area.objects.all(),
        'sections': Section.objects.all(),
        'parts': PartName.objects.all(),
        'chapters': ChapterName.objects.all(),
        'topics': TopicName.objects.all(),
        'subtopics': SubTopicName.objects.all(),
    }

    if request.method == 'POST':
        try:
            with transaction.atomic():
                # ----- Hierarchy selections -----
                area_ids    = request.POST.getlist('area_name[]')
                section_ids = request.POST.getlist('section_name[]')
                part_ids    = request.POST.getlist('part_name[]')
                chapter_ids = request.POST.getlist('chapter_name[]')
                topic_ids   = request.POST.getlist('topic_name[]')

                # ✅ Accept both legacy and new param names for subtopics
                subtopic_ids = request.POST.getlist('sub_topic_name[]') or request.POST.getlist('subtopic_name[]')
                subtopic_ids = [str(s).strip() for s in subtopic_ids if str(s).strip()]

                # ✅ Exams (M2M)
                exam_ids = [str(x).strip() for x in request.POST.getlist('exam_name[]') if str(x).strip()]

                new_topic_name = (request.POST.get('new_topic_name') or '').strip()

                # ----- Validate/parse exam year -----
                exam_year_raw = request.POST.get('exam_year')
                if exam_year_raw in (None, '', 'null'):
                    exam_year = None
                else:
                    try:
                        exam_year = int(exam_year_raw)
                    except ValueError:
                        messages.error(request, "Please provide a valid exam year.")
                        return redirect(request.path_info)

                script = request.POST.get('script', '')

                # ----- Stems -----
                question_part_first    = request.POST.get('question_part_first', '')
                question_part_first_hi = request.POST.get('question_part_first_hi', '')

                # ----- Answer / explanation -----
                correct_answer_choice         = request.POST.get('correct_answer_choice', '')
                correct_answer_description    = request.POST.get('correct_answer_description', '')
                correct_answer_description_hi = request.POST.get('correct_answer_description_hi', '')

                # ----- Meta -----
                marks                = float(request.POST.get('marks') or 0.0)
                negative_marks       = float(request.POST.get('negative_marks') or 0.0)
                degree_of_difficulty = (request.POST.get('degree_of_difficulty') or '').strip()
                elim_tactics_degree  = (request.POST.get('elim_tactics_degree') or '').strip()  # ✅ NEW
                evergreen_index      = (request.POST.get('evergreen_index') or '').strip()

                current_relevance       = (request.POST.get('current_relevance') or '').strip()
                current_relevance_topic = request.POST.get('current_relevance_topic', '')

                # Source → type
                question_type_source = (request.POST.get('question_type_source', 'modern') or '').strip()
                if question_type_source == 'pyq':
                    type_of_question = 'pyq'
                elif question_type_source == 'osq':
                    type_of_question = 'osq'
                else:
                    type_of_question = 'moq'

                # ----- New topic (if "other") -----
                if 'other' in topic_ids and new_topic_name and chapter_ids:
                    topic_ids.remove('other')
                    selected_chapter = ChapterName.objects.get(chapter_number=chapter_ids[0])
                    new_topic = TopicName.objects.create(name=new_topic_name, chapter=selected_chapter)
                    topic_ids.append(str(new_topic.topic_SI_number))

                # ----- M2M QuerySets -----
                areas_qs     = Area.objects.filter(area_SI_Code__in=area_ids)
                sections_qs  = Section.objects.filter(section_Unit_SI__in=section_ids)
                parts_qs     = PartName.objects.filter(part_serial__in=part_ids)
                chapters_qs  = ChapterName.objects.filter(chapter_number__in=chapter_ids)
                topics_qs    = TopicName.objects.filter(topic_SI_number__in=topic_ids)
                subtopics_qs = SubTopicName.objects.filter(sub_topic_SI_Number__in=subtopic_ids)
                exams_qs     = ExamName.objects.filter(exam_SI_Number__in=exam_ids)

                # (Optional) enforce at least one subtopic
                if not subtopics_qs.exists():
                    messages.error(request, "Please choose at least one Sub Topic.")
                    return redirect(request.path_info)

                # ----- IDs -----
                last_q = QuestionBank.objects.order_by('-question_number').first()
                next_question_number = last_q.question_number + 1 if last_q else 1

                while True:
                    base_question_id = random.randint(100000, 999999)
                    if not QuestionBank.objects.filter(base_question_id=base_question_id).exists():
                        break

                # ----- Create EN -----
                question_en = QuestionBank.objects.create(
                    question_number=next_question_number,
                    question_sub_type='true_and_false_type',
                    base_question_id=base_question_id,
                    type_of_question=type_of_question,
                    language='e',
                    script=script,
                    question_part_first=question_part_first,
                    correct_answer_choice=correct_answer_choice,
                    correct_answer_description=correct_answer_description,
                    marks=marks,
                    negative_marks=negative_marks,
                    degree_of_difficulty=degree_of_difficulty,
                    elim_tactics_degree=elim_tactics_degree,   # ✅ NEW
                    evergreen_index=evergreen_index,
                    exam_year=exam_year,
                    current_relevance=current_relevance,
                    current_relevance_topic=current_relevance_topic,
                    # Force True/False options
                    answer_option_a="True",
                    answer_option_b="False",
                    created_by=request.user
                )
                question_en.area_name.set(areas_qs)
                question_en.section_name.set(sections_qs)
                question_en.part_name.set(parts_qs)
                question_en.chapter_name.set(chapters_qs)
                question_en.topic_name.set(topics_qs)
                question_en.subtopic_name.set(subtopics_qs)
                question_en.exam_name.set(exams_qs)  # ✅ save Exams

                # ----- Create HI -----
                question_hi = QuestionBank.objects.create(
                    question_number=next_question_number,
                    question_sub_type='true_and_false_type',
                    base_question_id=base_question_id,
                    type_of_question=type_of_question,
                    language='h',
                    script=script,
                    question_part_first_hi=question_part_first_hi,
                    correct_answer_choice=correct_answer_choice,
                    correct_answer_description_hi=correct_answer_description_hi,
                    marks=marks,
                    negative_marks=negative_marks,
                    degree_of_difficulty=degree_of_difficulty,
                    elim_tactics_degree=elim_tactics_degree,   # ✅ NEW
                    evergreen_index=evergreen_index,
                    exam_year=exam_year,
                    current_relevance=current_relevance,
                    current_relevance_topic=current_relevance_topic,
                    # Force True/False options (Hindi)
                    answer_option_a="सही",
                    answer_option_b="गलत",
                    created_by=request.user
                )
                question_hi.area_name.set(areas_qs)
                question_hi.section_name.set(sections_qs)
                question_hi.part_name.set(parts_qs)
                question_hi.chapter_name.set(chapters_qs)
                question_hi.topic_name.set(topics_qs)
                question_hi.subtopic_name.set(subtopics_qs)
                question_hi.exam_name.set(exams_qs)  # ✅ save Exams

                messages.success(request, '✅ True & False Type Question added successfully!')
                return redirect('add-true-and-false-type-question')

        except Exception as e:
            messages.error(request, f'❌ An unexpected error occurred: {str(e)}')

    return render(request, 'question_bank/add_question/true_false_type_form.html', context)

# ************************* Create True and False Type Question End *********************************************




# ************************* Edit True and False Type Question Updated *********************************************
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.db import transaction
from .models import (
    Area, Section, PartName, ChapterName, TopicName, SubTopicName, ExamName, QuestionBank
)

@login_required
def edit_true_and_false_type_question(request, base_question_id):
    english_q = get_object_or_404(QuestionBank, base_question_id=base_question_id, language='e')
    hindi_q = QuestionBank.objects.filter(base_question_id=base_question_id, language='h').first()

    areas = Area.objects.all()
    sections = Section.objects.all()
    parts = PartName.objects.all()
    chapters = ChapterName.objects.all()
    topics = TopicName.objects.all()
    subtopics = SubTopicName.objects.all()

    if request.method == 'POST':
        try:
            with transaction.atomic():
                script = request.POST.get('script', '')
                question_part_first = request.POST.get('question_part_first', '')
                question_part_first_hi = request.POST.get('question_part_first_hi', '')

                correct_answer_choice = request.POST.get('correct_answer_choice', '')
                correct_answer_description = request.POST.get('correct_answer_description', '')
                correct_answer_description_hi = request.POST.get('correct_answer_description_hi', '')

                marks = float(request.POST.get('marks') or 0.0)
                negative_marks = float(request.POST.get('negative_marks') or 0.0)
                degree_of_difficulty = request.POST.get('degree_of_difficulty', '')
                evergreen_index = request.POST.get('evergreen_index', '')
                current_relevance = request.POST.get('current_relevance', '')
                current_relevance_topic = request.POST.get('current_relevance_topic', '')

                exam_year = request.POST.get('exam_year') or None

                area_ids = request.POST.getlist('area_name[]')
                section_ids = request.POST.getlist('section_name[]')
                part_ids = request.POST.getlist('part_name[]')
                chapter_ids = request.POST.getlist('chapter_name[]')
                topic_ids = request.POST.getlist('topic_name[]')
                subtopic_ids = request.POST.getlist('sub_topic_name[]')

                areas_qs = Area.objects.filter(area_SI_Code__in=area_ids)
                sections_qs = Section.objects.filter(section_Unit_SI__in=section_ids)
                parts_qs = PartName.objects.filter(part_serial__in=part_ids)
                chapters_qs = ChapterName.objects.filter(chapter_number__in=chapter_ids)
                topics_qs = TopicName.objects.filter(topic_SI_number__in=topic_ids)
                subtopics_qs = SubTopicName.objects.filter(sub_topic_SI_Number__in=subtopic_ids)

                # ✅ Update English version
                english_q.script = script
                english_q.question_part_first = question_part_first
                english_q.correct_answer_choice = correct_answer_choice
                english_q.correct_answer_description = correct_answer_description
                english_q.marks = marks
                english_q.negative_marks = negative_marks
                english_q.degree_of_difficulty = degree_of_difficulty
                english_q.evergreen_index = evergreen_index
                english_q.current_relevance = current_relevance
                english_q.current_relevance_topic = current_relevance_topic
                english_q.exam_year = exam_year
                english_q.answer_option_a = "True"
                english_q.answer_option_b = "False"
                english_q.save()

                english_q.area_name.set(areas_qs)
                english_q.section_name.set(sections_qs)
                english_q.part_name.set(parts_qs)
                english_q.chapter_name.set(chapters_qs)
                english_q.topic_name.set(topics_qs)
                english_q.subtopic_name.set(subtopics_qs)

                # ✅ Update Hindi version if exists
                if hindi_q:
                    hindi_q.script = script
                    hindi_q.question_part_first_hi = question_part_first_hi
                    hindi_q.correct_answer_choice = correct_answer_choice
                    hindi_q.correct_answer_description_hi = correct_answer_description_hi
                    hindi_q.marks = marks
                    hindi_q.negative_marks = negative_marks
                    hindi_q.degree_of_difficulty = degree_of_difficulty
                    hindi_q.evergreen_index = evergreen_index
                    hindi_q.current_relevance = current_relevance
                    hindi_q.current_relevance_topic = current_relevance_topic
                    hindi_q.exam_year = exam_year
                    hindi_q.answer_option_a = "सही"
                    hindi_q.answer_option_b = "गलत"
                    hindi_q.save()

                    hindi_q.area_name.set(areas_qs)
                    hindi_q.section_name.set(sections_qs)
                    hindi_q.part_name.set(parts_qs)
                    hindi_q.chapter_name.set(chapters_qs)
                    hindi_q.topic_name.set(topics_qs)
                    hindi_q.subtopic_name.set(subtopics_qs)

                messages.success(request, "✅ True & False Type Question updated successfully!")
                return redirect('view_questions')

        except Exception as e:
            messages.error(request, f"❌ An error occurred while editing: {e}")

    context = {
        'english_q': english_q,
        'hindi_q': hindi_q,
        'areas': areas,
        'sections': sections,
        'parts': parts,
        'chapters': chapters,
        'topics': topics,
        'subtopics': subtopics,
        "mark_options": ["0.5", "1.0", "1.5", "2.0", "2.5", "3.0"],
        "negative_mark_options": ["1.0", "0.5", "0.33", "0.25", "0.2", "0"],
    }

    return render(request, 'question_bank/edit_question/edit_true_false_type_form.html', context)
# ************************* Edit True and False Type Question Updated *********************************************



# ************************* Create Fill in the Blank Type Question Start *********************************************
from django.shortcuts import render, redirect
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.db import transaction
from .models import (
    Area, Section, PartName, ChapterName, TopicName, SubTopicName, QuestionBank
)
import random

@login_required
def add_fill_in_the_blank_question(request):
    context = {
        'areas': Area.objects.all(),
        'sections': Section.objects.all(),
        'parts': PartName.objects.all(),
        'chapters': ChapterName.objects.all(),
        'topics': TopicName.objects.all(),
        'subtopics': SubTopicName.objects.all(),
    }

    if request.method == 'POST':
        try:
            with transaction.atomic():
                # ✅ Extract fields
                area_ids = request.POST.getlist('area_name[]')
                section_ids = request.POST.getlist('section_name[]')
                part_ids = request.POST.getlist('part_name[]')
                chapter_ids = request.POST.getlist('chapter_name[]')
                topic_ids = request.POST.getlist('topic_name[]')
                subtopic_ids = request.POST.getlist('sub_topic_name[]')

                new_topic_name = request.POST.get('new_topic_name', '').strip()
                new_subtopic_name = request.POST.get('new_subtopic_name', '').strip()

                exam_year = request.POST.get('exam_year') or None
                script = request.POST.get('script', '')

                question_part_first = request.POST.get('question_part_first', '')
                question_part_first_hi = request.POST.get('question_part_first_hi', '')
                correct_answer_choice = request.POST.get('correct_answer_choice', '')
                correct_answer_description = request.POST.get('correct_answer_description', '')
                correct_answer_description_hi = request.POST.get('correct_answer_description_hi', '')

                marks = float(request.POST.get('marks') or 0.0)
                negative_marks = float(request.POST.get('negative_marks') or 0.0)
                degree_of_difficulty = request.POST.get('degree_of_difficulty', '')
                evergreen_index = request.POST.get('evergreen_index', '')

                current_relevance = request.POST.get('current_relevance', '')
                current_relevance_topic = request.POST.get('current_relevance_topic', '')

                question_type_source = request.POST.get('question_type_source', 'modern')
                type_of_question = 'pyq' if question_type_source == 'pyq' else (
                    'osq' if question_type_source == 'osq' else 'moq'
                )

                # ✅ Handle new topic creation
                if 'other' in topic_ids:
                    topic_ids.remove('other')
                    if new_topic_name and chapter_ids:
                        selected_chapter = ChapterName.objects.get(chapter_number=chapter_ids[0])
                        new_topic = TopicName.objects.create(
                            name=new_topic_name,
                            chapter=selected_chapter
                        )
                        topic_ids.append(new_topic.topic_SI_number)

                # ✅ Handle new subtopic creation
                if 'other' in subtopic_ids:
                    subtopic_ids.remove('other')
                    if new_subtopic_name and topic_ids:
                        selected_topic = TopicName.objects.get(topic_SI_number=topic_ids[0])
                        new_subtopic = SubTopicName.objects.create(
                            name=new_subtopic_name,
                            topic=selected_topic
                        )
                        subtopic_ids.append(new_subtopic.sub_topic_SI_Number)

                # ✅ Fetch related M2M objects
                areas_qs = Area.objects.filter(area_SI_Code__in=area_ids)
                sections_qs = Section.objects.filter(section_Unit_SI__in=section_ids)
                parts_qs = PartName.objects.filter(part_serial__in=part_ids)
                chapters_qs = ChapterName.objects.filter(chapter_number__in=chapter_ids)
                topics_qs = TopicName.objects.filter(topic_SI_number__in=topic_ids)
                subtopics_qs = SubTopicName.objects.filter(sub_topic_SI_Number__in=subtopic_ids)

                # ✅ Generate question number & unique base ID
                last_question = QuestionBank.objects.order_by('-question_number').first()
                next_question_number = last_question.question_number + 1 if last_question else 1

                while True:
                    base_question_id = random.randint(100000, 999999)
                    if not QuestionBank.objects.filter(base_question_id=base_question_id).exists():
                        break

                # ✅ Create English question
                question_en = QuestionBank.objects.create(
                    question_number=next_question_number,
                    question_sub_type='fill_in_the_blank_type',
                    base_question_id=base_question_id,
                    type_of_question=type_of_question,
                    language='e',
                    script=script,
                    question_part_first=question_part_first,
                    correct_answer_choice=correct_answer_choice,
                    correct_answer_description=correct_answer_description,
                    marks=marks,
                    negative_marks=negative_marks,
                    degree_of_difficulty=degree_of_difficulty,
                    evergreen_index=evergreen_index,
                    exam_year=exam_year,
                    current_relevance=current_relevance,
                    current_relevance_topic=current_relevance_topic,
                    created_by=request.user
                )

                # ✅ Create Hindi question
                question_hi = QuestionBank.objects.create(
                    question_number=next_question_number,
                    question_sub_type='fill_in_the_blank_type',
                    base_question_id=base_question_id,
                    type_of_question=type_of_question,
                    language='h',
                    script=script,
                    question_part_first_hi=question_part_first_hi,
                    correct_answer_choice=correct_answer_choice,
                    correct_answer_description_hi=correct_answer_description_hi,
                    marks=marks,
                    negative_marks=negative_marks,
                    degree_of_difficulty=degree_of_difficulty,
                    evergreen_index=evergreen_index,
                    exam_year=exam_year,
                    current_relevance=current_relevance,
                    current_relevance_topic=current_relevance_topic,
                    created_by=request.user
                )

                # ✅ Save M2M relations
                for q in [question_en, question_hi]:
                    q.area_name.set(areas_qs)
                    q.section_name.set(sections_qs)
                    q.part_name.set(parts_qs)
                    q.chapter_name.set(chapters_qs)
                    q.topic_name.set(topics_qs)
                    q.subtopic_name.set(subtopics_qs)

                messages.success(request, '✅ Fill in the Blank Question added successfully in English and Hindi!')
                return redirect('add-fill-in-the-blank-question')

        except Exception as e:
            messages.error(request, f'❌ An unexpected error occurred: {str(e)}')

    return render(request, 'question_bank/add_question/fill_in_the_blank_form.html', context)

# ************************* Create Fill in the Blank Type Question End *********************************************



# ************************* Edit Fill in the Blank Type Question Updated *********************************************
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.db import transaction
from .models import (
    Area, Section, PartName, ChapterName, TopicName, SubTopicName, ExamName, QuestionBank
)

@login_required
def edit_fill_in_the_blank_question(request, base_question_id):
    english_q = get_object_or_404(QuestionBank, base_question_id=base_question_id, language='e')
    hindi_q = QuestionBank.objects.filter(base_question_id=base_question_id, language='h').first()

    areas = Area.objects.all()
    sections = Section.objects.all()
    parts = PartName.objects.all()
    chapters = ChapterName.objects.all()
    topics = TopicName.objects.all()
    subtopics = SubTopicName.objects.all()

    if request.method == 'POST':
        try:
            with transaction.atomic():
                script = request.POST.get('script', '')
                question_part_first = request.POST.get('question_part_first', '')
                question_part_first_hi = request.POST.get('question_part_first_hi', '')

                correct_answer_choice = request.POST.get('correct_answer_choice', '')
                correct_answer_description = request.POST.get('correct_answer_description', '')
                correct_answer_description_hi = request.POST.get('correct_answer_description_hi', '')

                marks = float(request.POST.get('marks') or 0.0)
                negative_marks = float(request.POST.get('negative_marks') or 0.0)
                degree_of_difficulty = request.POST.get('degree_of_difficulty', '')
                evergreen_index = request.POST.get('evergreen_index', '')
                current_relevance = request.POST.get('current_relevance', '')
                current_relevance_topic = request.POST.get('current_relevance_topic', '')

                exam_year = request.POST.get('exam_year') or None

                area_ids = request.POST.getlist('area_name[]')
                section_ids = request.POST.getlist('section_name[]')
                part_ids = request.POST.getlist('part_name[]')
                chapter_ids = request.POST.getlist('chapter_name[]')
                topic_ids = request.POST.getlist('topic_name[]')
                subtopic_ids = request.POST.getlist('sub_topic_name[]')

                areas_qs = Area.objects.filter(area_SI_Code__in=area_ids)
                sections_qs = Section.objects.filter(section_Unit_SI__in=section_ids)
                parts_qs = PartName.objects.filter(part_serial__in=part_ids)
                chapters_qs = ChapterName.objects.filter(chapter_number__in=chapter_ids)
                topics_qs = TopicName.objects.filter(topic_SI_number__in=topic_ids)
                subtopics_qs = SubTopicName.objects.filter(sub_topic_SI_Number__in=subtopic_ids)

                # ✅ Update English Question
                english_q.script = script
                english_q.question_part_first = question_part_first
                english_q.correct_answer_choice = correct_answer_choice
                english_q.correct_answer_description = correct_answer_description
                english_q.marks = marks
                english_q.negative_marks = negative_marks
                english_q.degree_of_difficulty = degree_of_difficulty
                english_q.evergreen_index = evergreen_index
                english_q.current_relevance = current_relevance
                english_q.current_relevance_topic = current_relevance_topic
                english_q.exam_year = exam_year
                english_q.save()

                english_q.area_name.set(areas_qs)
                english_q.section_name.set(sections_qs)
                english_q.part_name.set(parts_qs)
                english_q.chapter_name.set(chapters_qs)
                english_q.topic_name.set(topics_qs)
                english_q.subtopic_name.set(subtopics_qs)

                # ✅ Update Hindi Question if exists
                if hindi_q:
                    hindi_q.script = script
                    hindi_q.question_part_first_hi = question_part_first_hi
                    hindi_q.correct_answer_choice = correct_answer_choice
                    hindi_q.correct_answer_description_hi = correct_answer_description_hi
                    hindi_q.marks = marks
                    hindi_q.negative_marks = negative_marks
                    hindi_q.degree_of_difficulty = degree_of_difficulty
                    hindi_q.evergreen_index = evergreen_index
                    hindi_q.current_relevance = current_relevance
                    hindi_q.current_relevance_topic = current_relevance_topic
                    hindi_q.exam_year = exam_year
                    hindi_q.save()

                    hindi_q.area_name.set(areas_qs)
                    hindi_q.section_name.set(sections_qs)
                    hindi_q.part_name.set(parts_qs)
                    hindi_q.chapter_name.set(chapters_qs)
                    hindi_q.topic_name.set(topics_qs)
                    hindi_q.subtopic_name.set(subtopics_qs)

                messages.success(request, "✅ Fill in the Blank Question updated successfully!")
                return redirect('view_questions')

        except Exception as e:
            messages.error(request, f"❌ An error occurred while editing: {e}")

    context = {
        'english_q': english_q,
        'hindi_q': hindi_q,
        'areas': areas,
        'sections': sections,
        'parts': parts,
        'chapters': chapters,
        'topics': topics,
        'subtopics': subtopics,
        "mark_options": ["0.5", "1.0", "1.5", "2.0", "2.5", "3.0"],
        "negative_mark_options": ["1.0", "0.5", "0.33", "0.25", "0.2", "0"],
    }

    return render(request, 'question_bank/edit_question/edit_fill_in_the_blank_form.html', context)
# ************************* Edit Fill in the Blank Type Question Updated *********************************************


from django.shortcuts import get_object_or_404, redirect
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from .models import QuestionBank

@login_required
def delete_question(request, pk):
    question = get_object_or_404(QuestionBank, pk=pk)
    question.delete()
    messages.success(request, "Question deleted successfully.")
    return redirect('view_questions')

# Lectures Notes Views

from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.db import transaction

from .models import Area, PartName, ChapterName, TopicName, SubTopicName, LectureNote
from .forms import LectureNoteForm


@login_required
def upload_lecture_note_view(request):
    if request.method == "POST":
        form = LectureNoteForm(request.POST, request.FILES)

        # ✅ Manually fetch files because they are not part of the model
        note_file_en = request.FILES.get("note_file_en")
        note_file_hi = request.FILES.get("note_file_hi")

        if not (note_file_en or note_file_hi):
            messages.error(request, "Please upload at least one file (English or Hindi).")
            return redirect(request.path_info)

        # ✅ Check required fields
        area = request.POST.get("area")
        part = request.POST.get("part")
        chapter = request.POST.get("chapter")
        ctpl = request.POST.get("ctpl")
        note_type = request.POST.get("note_type")

        if not area or not part or not chapter:
            messages.error(request, "Area, Part, and Chapter are required.")
            return redirect(request.path_info)

        topic = request.POST.get("topic") or None
        subtopic = request.POST.get("subtopic") or None

        try:
            with transaction.atomic():
                if note_file_en:
                    LectureNote.objects.create(
                        language="en",
                        note_type=note_type,
                        ctpl=ctpl,
                        note_file=note_file_en,
                        area_id=area,
                        part_id=part,
                        chapter_id=chapter,
                        topic_id=topic,
                        subtopic_id=subtopic,
                        created_by=request.user
                    )
                if note_file_hi:
                    LectureNote.objects.create(
                        language="hi",
                        note_type=note_type,
                        ctpl=ctpl,
                        note_file=note_file_hi,
                        area_id=area,
                        part_id=part,
                        chapter_id=chapter,
                        topic_id=topic,
                        subtopic_id=subtopic,
                        created_by=request.user
                    )
            messages.success(request, "Lecture note(s) uploaded successfully.")
            return redirect("upload_lecture_note")

        except Exception as e:
            messages.error(request, f"An error occurred: {e}")
            return redirect(request.path_info)

    else:
        form = LectureNoteForm()

    return render(request, "question_bank/upload_lecture_note.html", {
        "form": form,
        "areas": Area.objects.all(),
        "parts": PartName.objects.all(),
        "chapters": ChapterName.objects.all(),
        "topics": TopicName.objects.all(),
        "subtopics": SubTopicName.objects.all(),
    })


@login_required
def list_lecture_notes(request):
    notes = LectureNote.objects.all().order_by('-created_at')
    return render(request, "question_bank/list_notes.html", {"notes": notes})


@login_required
def view_lecture_note_detail(request, pk):
    note = get_object_or_404(LectureNote, pk=pk)
    return render(request, "question_bank/note_detail_flipbook.html", {"note": note})




# Input Suggestion Views Started 

from django.shortcuts import render, redirect
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.db import transaction
from django.db.utils import IntegrityError
import re

from .models import (
    InputSuggestion,
    InputSuggestionImage,
    InputSuggestionDocument,
    Area, PartName, ChapterName, TopicName, SubTopicName, ExamName, HashtagsName
)


@login_required
def add_input_suggestion(request):
    if request.method == "POST":
        try:
            # ✅ Basic fields
            languages = request.POST.getlist('language')
            language = ",".join(languages)

            brief_description = request.POST.get('brief_description')
            details = request.POST.get('details')
            other_text = request.POST.get('other_text')
            question_video = request.FILES.get('question_video')
            source = request.POST.get('source', 'self')
            approval_status = 'pending_faculty'

            # ✅ Links & credits
            question_links_list = request.POST.getlist('question_link[]')
            question_link = ', '.join([l.strip() for l in question_links_list if l.strip()])
            credits_list = request.POST.getlist('credit_or_courtesy[]')
            credit_or_courtesy = ', '.join([c.strip() for c in credits_list if c.strip()])

            # ✅ Relevance
            relevance_topics_list = request.POST.getlist('current_relevance_topic[]')
            current_relevance_topic = ' || '.join([t.strip() for t in relevance_topics_list if t.strip()])
            current_relevance = request.POST.get('current_relevance')

            # ✅ Dropdown IDs
            area_ids = request.POST.getlist('area_name[]')
            part_ids = request.POST.getlist('part_name[]')
            chapter_ids = request.POST.getlist('chapter_name[]')
            topic_ids = request.POST.getlist('topic_name[]')
            subtopic_ids = request.POST.getlist('sub_topic_name[]')
            exam_ids = request.POST.getlist('exam_name[]')

            # ✅ NEW: Shortcodes (split on comma OR space)
            shortcode_raw = request.POST.get('shortcodeInput', '')
            shortcodes = [s.strip() for s in re.split(r'[\s,]+', shortcode_raw) if s.strip()]

            # ✅ Hashtags
            hashtags_input = request.POST.get('hashtags') or request.POST.get('hashtagsHiddenInput', '')
            user_hashtags = [h.strip().lstrip('#') for h in hashtags_input.split(',') if h.strip()]

            if not languages or not brief_description or not details or not current_relevance:
                messages.error(request, "Please fill in all required fields.")
                return redirect(request.path_info)

            with transaction.atomic():
                suggestion = InputSuggestion.objects.create(
                    language=language,
                    brief_description=brief_description,
                    details=details,
                    question_link=question_link,
                    other_text=other_text,
                    question_video=question_video,
                    source=source,
                    approval_status=approval_status,
                    credit_or_courtesy=credit_or_courtesy,
                    current_relevance=current_relevance,
                    current_relevance_topic=current_relevance_topic,
                    created_by=request.user
                )

                # ✅ If shortcodes → resolve hierarchy
                if shortcodes:
                    resolved_subtopics = SubTopicName.objects.filter(sub_topic_short_Code__in=shortcodes)
                    if resolved_subtopics.exists():
                        suggestion.subtopic_name.set(resolved_subtopics)

                        # Cascade relations
                        topic_ids = set(resolved_subtopics.values_list("topic__topic_SI_number", flat=True))
                        chapter_ids = set(resolved_subtopics.values_list("topic__chapter__chapter_number", flat=True))
                        part_ids = set(resolved_subtopics.values_list("topic__chapter__part__part_serial", flat=True))
                        area_ids = set(resolved_subtopics.values_list("topic__chapter__part__section__area__area_SI_Code", flat=True))
                        exam_ids = set(
                            ExamName.objects.filter(subtopics__in=resolved_subtopics)
                            .values_list("exam_SI_Number", flat=True)
                        )

                        # Apply relations
                        suggestion.topic_name.set(TopicName.objects.filter(topic_SI_number__in=topic_ids))
                        suggestion.chapter_name.set(ChapterName.objects.filter(chapter_number__in=chapter_ids))
                        suggestion.part_name.set(PartName.objects.filter(part_serial__in=part_ids))
                        suggestion.area_name.set(Area.objects.filter(area_SI_Code__in=area_ids))
                        suggestion.exams.set(ExamName.objects.filter(exam_SI_Number__in=exam_ids))

                        # Auto hashtags
                        auto_hashtags = HashtagsName.objects.filter(subtopics__in=resolved_subtopics).distinct()
                        for tag in auto_hashtags:
                            suggestion.hashtags.add(tag)

                # ✅ Else → use manual dropdowns
                else:
                    if area_ids:
                        suggestion.area_name.set(Area.objects.filter(area_SI_Code__in=area_ids))
                    if part_ids:
                        suggestion.part_name.set(PartName.objects.filter(part_serial__in=part_ids))
                    if chapter_ids:
                        suggestion.chapter_name.set(ChapterName.objects.filter(chapter_number__in=chapter_ids))
                    if topic_ids:
                        suggestion.topic_name.set(TopicName.objects.filter(topic_SI_number__in=topic_ids))
                    if subtopic_ids:
                        subtopics = SubTopicName.objects.filter(sub_topic_SI_Number__in=subtopic_ids)
                        suggestion.subtopic_name.set(subtopics)

                        auto_hashtags = HashtagsName.objects.filter(subtopics__in=subtopics).distinct()
                        for tag in auto_hashtags:
                            suggestion.hashtags.add(tag)
                    if exam_ids:
                        suggestion.exams.set(ExamName.objects.filter(exam_SI_Number__in=exam_ids))

                # ✅ Create/link manual hashtags
                used_slugs = set(HashtagsName.objects.values_list('hashtags_SI_Number', flat=True))

                def generate_unique_slug(base_slug):
                    slug = base_slug
                    counter = 1
                    while slug in used_slugs:
                        slug = f"{base_slug}_{counter}"
                        counter += 1
                    used_slugs.add(slug)
                    return slug

                from django.utils.text import slugify

                for tag in user_hashtags:
                    tag = tag.strip().lstrip('#')
                    if not tag:
                        continue

                    obj, created = HashtagsName.objects.get_or_create(
                        name__iexact=tag,  # case-insensitive search
                        defaults={'name': tag}
                    )

                    suggestion.hashtags.add(obj)
                    if suggestion.subtopic_name.exists():
                        for subtopic in suggestion.subtopic_name.all():
                            subtopic.hashtags.add(obj)

                # ✅ Save files
                for image in request.FILES.getlist('question_images'):
                    InputSuggestionImage.objects.create(question=suggestion, image=image)
                for document in request.FILES.getlist('question_documents'):
                    InputSuggestionDocument.objects.create(question=suggestion, document=document)

                messages.success(request, "✅ Input Suggestion added successfully!")
                return redirect('input-suggestion-list')

        except Exception as e:
            messages.error(request, f"❌ An error occurred: {e}")
            return redirect(request.path_info)

    return render(request, 'question_bank/add_input_suggestion.html', {
        'areas': Area.objects.all(),
        'parts': PartName.objects.all(),
        'chapters': ChapterName.objects.all(),
        'topics': TopicName.objects.all(),
        'subtopics': SubTopicName.objects.all(),
        'exams': ExamName.objects.all(),
    })

from django.shortcuts import render, get_object_or_404
from django.contrib.auth.decorators import login_required
from .models import (
    InputSuggestion, QuestionBank, Area, PartName, ChapterName,
    TopicName, SubTopicName, HashtagsName, LectureNote  # ✅ include LectureNote
)

@login_required
def view_input_suggestion(request, pk=None):
    # ✅ Base queryset
    questions = InputSuggestion.objects.prefetch_related(
        'area_name', 'part_name', 'chapter_name',
        'topic_name', 'subtopic_name', 'hashtags'
    ).distinct()

    # ✅ Extract updated filters
    area_name = request.GET.get('area_name')
    part_name = request.GET.get('part_name')
    chapter_name = request.GET.get('chapter_name')
    topic_name = request.GET.get('topic_name')
    subtopic_name = request.GET.get('subtopic_name')
    hashtag = request.GET.get('hashtag')

    # ✅ Apply filters (uses ManyToMany robust match)
    if area_name:
        questions = questions.filter(area_name__area_SI_Code__in=[str(area_name)])
    if part_name:
        questions = questions.filter(part_name__part_serial__in=[str(part_name)])
    if chapter_name:
        questions = questions.filter(chapter_name__chapter_number__in=[str(chapter_name)])
    if topic_name:
        questions = questions.filter(topic_name__topic_SI_number__in=[str(topic_name)])
    if subtopic_name:
        questions = questions.filter(subtopic_name__sub_topic_SI_Number__in=[str(subtopic_name)])
    if hashtag:
        questions = questions.filter(hashtags__hashtags_SI_Number__in=[str(hashtag)])

    # ✅ Related subtopic and chapter lookup
    suggestion = None
    subtopic_ids = []
    chapter_ids = []
    lecture_notes = LectureNote.objects.none()

    if pk:
        suggestion = get_object_or_404(InputSuggestion, pk=pk)
        subtopic_ids = list(suggestion.subtopic_name.values_list('sub_topic_SI_Number', flat=True))
        chapter_ids = list(suggestion.chapter_name.values_list('chapter_number', flat=True))

        # ✅ Fetch lecture notes based on chapter match
        if chapter_ids:
            lecture_notes = LectureNote.objects.filter(chapter__chapter_number__in=chapter_ids).order_by('-created_at')[:10]
    else:
        subtopic_ids = (
            questions.filter(subtopic_name__isnull=False)
            .values_list('subtopic_name__sub_topic_SI_Number', flat=True)
            .distinct()
        )

    # ✅ Build grouped related QuestionBank pairs
    grouped_questions = {}
    if subtopic_ids:
        related_qs = (
            QuestionBank.objects.filter(subtopic_name__sub_topic_SI_Number__in=subtopic_ids)
            .distinct().order_by('base_question_id', 'language')
        )
        for q in related_qs:
            base_id = q.base_question_id
            if base_id not in grouped_questions:
                grouped_questions[base_id] = {'english': None, 'hindi': None}
            if q.language == 'e':
                grouped_questions[base_id]['english'] = q
            elif q.language == 'h':
                grouped_questions[base_id]['hindi'] = q

    context = {
        'questions': questions.distinct(),
        'suggestion': suggestion if pk else None,
        'areas': Area.objects.all(),
        'parts': PartName.objects.all(),
        'chapters': ChapterName.objects.all(),
        'topics': TopicName.objects.all(),
        'subtopics': SubTopicName.objects.all(),
        'hashtags': HashtagsName.objects.all(),
        'grouped_questions': grouped_questions,
        'lecture_notes': lecture_notes,  # ✅ include lecture notes
        'request': request,
    }
    return render(request, 'question_bank/input_suggestion_list.html', context)

from django.shortcuts import get_object_or_404, render
from .models import InputSuggestion, LectureNote

from django.shortcuts import get_object_or_404, render
from .models import QuestionBank, InputSuggestion, LectureNote

def question_blog_view(request, question_id):
    question = get_object_or_404(InputSuggestion, id=question_id)

    # ✅ Related LectureNotes by chapter
    chapter_ids = list(question.chapter_name.values_list('chapter_number', flat=True)) if hasattr(question.chapter_name, 'all') else [question.chapter_name.chapter_number]
    related_notes = LectureNote.objects.filter(chapter__chapter_number__in=chapter_ids).order_by('-created_at') if chapter_ids else []

    # ✅ Related Questions by subtopic
    subtopic_ids = list(question.subtopic_name.values_list('sub_topic_SI_Number', flat=True))
    grouped_questions = {}

    if subtopic_ids:
        related_qs = QuestionBank.objects.filter(subtopic_name__sub_topic_SI_Number__in=subtopic_ids).distinct().order_by('base_question_id', 'language')
        for q in related_qs:
            if q.base_question_id not in grouped_questions:
                grouped_questions[q.base_question_id] = {'english': None, 'hindi': None}
            if q.language == 'e':
                grouped_questions[q.base_question_id]['english'] = q
            elif q.language == 'h':
                grouped_questions[q.base_question_id]['hindi'] = q

    return render(
        request,
        'question_bank/view_input_suggestion.html',
        {
            'question': question,
            'grouped_questions': grouped_questions,
            'related_notes': related_notes,
        }
    )

import json
import re
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.contrib.auth.decorators import login_required
from django.utils.text import slugify
from openai import OpenAI
from decouple import config

from .models import HashtagsName, SubTopicName

# Initialize OpenAI client
client = OpenAI(api_key=config("OPENAI_API_KEY"))


# ============================================================
# 🔹 Generate AI Hashtags
# ============================================================
@login_required
@csrf_exempt
def generate_ai_hashtags(request):
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            brief = data.get("brief_description", "")
            details = data.get("details", "")
            resources = data.get("resources", [])
            existing_hashtags = data.get("existing_hashtags", "")

            prompt = f"""
            You are an expert content strategist. Based on the following input, suggest 5-10 new hashtags.

            Brief Description: {brief}
            Details: {details}
            Resource Links: {', '.join(resources)}
            Existing Hashtags: {existing_hashtags}

            Return only hashtags (without #), separated by commas.
            """

            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}]
            )

            hashtags_text = response.choices[0].message.content.strip()
            hashtags = [
                re.sub(r"[^a-zA-Z0-9_]", "", h.strip().lstrip("#"))
                for h in hashtags_text.split(",") if h.strip()
            ]

            return JsonResponse({"status": "success", "hashtags": hashtags})

        except Exception as e:
            return JsonResponse({"status": "error", "message": str(e)})

    return JsonResponse({"status": "error", "message": "Invalid request method"})


# ============================================================
# 🔹 Save Hashtags and Link with SubTopics
# ============================================================
@login_required
@csrf_exempt
def save_ai_hashtags(request):
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            hashtags = data.get("hashtags", [])
            subtopic_ids = data.get("subtopic_ids", [])

            saved = []

            for tag in hashtags:
                slug = slugify(tag)

                # Create or get by primary key (hashtags_SI_Number)
                obj, created = HashtagsName.objects.get_or_create(
                    hashtags_SI_Number=slug,
                    defaults={"name": tag}
                )

                for sid in subtopic_ids:
                    sub = SubTopicName.objects.filter(sub_topic_SI_Number=sid).first()
                    if sub and not sub.hashtags.filter(pk=obj.pk).exists():
                        sub.hashtags.add(obj)

                saved.append(tag)

            return JsonResponse({"status": "success", "saved": saved})

        except Exception as e:
            return JsonResponse({"status": "error", "message": str(e)})

    return JsonResponse({"status": "error", "message": "Invalid request method"})

# Input Suggestion Views End 

from django.contrib.auth.decorators import login_required
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.shortcuts import render
from .models import (
    QuestionBank,
    Area,
    Section,
    PartName,
    ChapterName,
    TopicName,
    SubTopicName
)

@login_required
def view_questions(request):
    # Filter questions created by the logged-in user
    questions = QuestionBank.objects.filter(
        created_by=request.user
    ).prefetch_related(
        'area_name', 'section_name', 'part_name',
        'chapter_name', 'topic_name', 'subtopic_name'
    ).order_by('-created_at')  # Order by creation date (latest first)

    # Get filter parameters from the request
    type_of_question = request.GET.get('question_source')
    question_sub_type = request.GET.get('question_type')
    area_name = request.GET.get('area_name')
    section_name = request.GET.get('section_name')
    part_name = request.GET.get('part_name')
    chapter_name = request.GET.get('chapter_name')
    topic_name = request.GET.get('topic_name')
    subtopic_name = request.GET.get('subtopic_name')

    # Apply filters based on the URL parameters
    if type_of_question:
        questions = questions.filter(type_of_question=type_of_question)
    if question_sub_type:
        questions = questions.filter(question_sub_type=question_sub_type)
    if area_name:
        questions = questions.filter(area_name__area_SI_Code=area_name)
    if section_name:
        questions = questions.filter(section_name__section_Unit_SI=section_name)
    if part_name:
        questions = questions.filter(part_name__part_serial=part_name)
    if chapter_name:
        questions = questions.filter(chapter_name__chapter_number=chapter_name)
    if topic_name:
        questions = questions.filter(topic_name__topic_SI_number=topic_name)
    if subtopic_name:
        questions = questions.filter(subtopic_name__sub_topic_SI_Number=subtopic_name)

    questions = questions.distinct()  # Ensure distinct questions

    # Paginate the questions, 50 per page
    paginator = Paginator(questions, 500)  # Show 50 questions per page
    page_number = request.GET.get('page')  # Get the current page number

    # Handle invalid page numbers
    try:
        page_obj = paginator.get_page(page_number)  # Get the page object
    except PageNotAnInteger:
        page_obj = paginator.get_page(1)  # Default to the first page if page number is not an integer
    except EmptyPage:
        page_obj = paginator.get_page(paginator.num_pages)  # Show the last page if the page number is out of range

    # Group questions by base_question_id
    grouped_questions = {}
    for question in page_obj:
        base_id = question.base_question_id
        if base_id not in grouped_questions:
            grouped_questions[base_id] = {'english': None, 'hindi': None}
        if question.language == 'e':
            grouped_questions[base_id]['english'] = question
        elif question.language == 'h':
            grouped_questions[base_id]['hindi'] = question

    context = {
        'grouped_questions': grouped_questions,
        'QUESTION_SOURCES': QuestionBank.QUESTION_SOURCES,
        'QUESTION_TYPES': QuestionBank.QUESTION_TYPES,
        'areas': Area.objects.all(),
        'sections': Section.objects.all(),
        'parts': PartName.objects.all(),
        'chapters': ChapterName.objects.all(),
        'topics': TopicName.objects.all(),
        'subtopics': SubTopicName.objects.all(),
        'request': request,
        'page_obj': page_obj,  # Add page_obj for pagination
    }
    return render(request, 'question_bank/add_question/view_questions.html', context)


def add_quote_idiom_phrase(request):
    if request.method == 'POST':
        # Get all form data
        type = request.POST.get('type')
        content = request.POST.get('content')
        meaning = request.POST.get('meaning', '')  # Optional meaning field
        author = request.POST.get('author', '')  # Optional author field
        exam_ids = request.POST.getlist('exam_name[]')  # Get multiple exam IDs
        subject_ids = request.POST.getlist('subject_name[]')  # Get multiple subject IDs
        area_ids = request.POST.getlist('area_name[]')  # Get multiple area IDs
        part_ids = request.POST.getlist('part_name[]')  # Get multiple part IDs
        chapter_ids = request.POST.getlist('chapter_name[]')  # Get multiple chapter IDs
        topic_ids = request.POST.getlist('topic_name[]')  # Get multiple topic IDs
        new_topic_name = request.POST.get('new_topic_name', '')  # New topic name if manually entered

        # Remove 'other' from topic_ids, as 'other' is not a valid ID
        if 'other' in topic_ids:
            topic_ids.remove('other')

        # Create the QuoteIdiomPhrase instance
        new_entry = QuoteIdiomPhrase.objects.create(
            type=type,
            content=content,
            meaning=meaning if type in ['idiom', 'phrase'] else '',  # Add meaning only for idioms or phrases
            author=author,  # Add author if provided
            created_by=request.user
        )

        # Handle ManyToMany fields after creating the instance
        # if exam_ids:
            # exams = ExamName.objects.filter(id__in=exam_ids)
            # new_entry.exams.set(exams)

        # if subject_ids:
            # subjects = Subject.objects.filter(id__in=subject_ids)
            # new_entry.subjects.set(subjects)

        if area_ids:
            areas = Area.objects.filter(id__in=area_ids)
            new_entry.areas.set(areas)

        if part_ids:
            parts = PartName.objects.filter(id__in=part_ids)
            new_entry.parts.set(parts)

        if chapter_ids:
            chapters = ChapterName.objects.filter(id__in=chapter_ids)
            new_entry.chapters.set(chapters)

        # Handle manually entered topic with association to the selected chapter
        if new_topic_name and chapter_ids:
            # Associate the manually entered topic with the first selected chapter
            related_chapter = ChapterName.objects.filter(id=chapter_ids[0]).first()
            new_topic = TopicName.objects.create(name=new_topic_name, chapter=related_chapter)
            new_entry.topics.add(new_topic)

        # Handle selected topics (without 'other')
        if topic_ids:
            selected_topics = TopicName.objects.filter(id__in=topic_ids)
            new_entry.topics.add(*selected_topics)

        # Display success message and redirect
        messages.success(request, 'Your Quote, Idiom, or Phrase has been added successfully!')
        return redirect('add_quote_idiom_phrase')

    # Fetch all exam names for the form
    # exam_names = ExamName.objects.all()
    return render(request, 'question_bank/add_quote_idiom_phrase.html')





def quotes_idioms_phrases_view(request):
    # Fetch all quotes, idioms, and phrases from the database
    quotes_idioms_phrases = QuoteIdiomPhrase.objects.all().order_by('-created_at')
    
    # Render the template with the fetched data
    return render(request, 'question_bank/quotes_idioms_phrases.html', {
        'quotes_idioms_phrases': quotes_idioms_phrases
    })





def analytics_dashboard(request):
    # Filter and annotate data for users with questions, suggestions, or quote/phrase/idiom entries
    users_data = User.objects.annotate(
        user_name=Concat('first_name', Value(' '), 'last_name'),
        total_input_suggestions=Count('inputsuggestion', filter=Q(inputsuggestion__isnull=False)),
        total_questions=Count('questionbank', filter=Q(questionbank__isnull=False)),
        simple_type_count=Count('questionbank', filter=Q(questionbank__question_sub_type='simple_type')),
        list1_type_count=Count('questionbank', filter=Q(questionbank__question_sub_type='list_type_1')),
        list2_type_count=Count('questionbank', filter=Q(questionbank__question_sub_type='list_type_2')),
        ra_type_count=Count('questionbank', filter=Q(questionbank__question_sub_type='r_and_a_type')),
        true_false_type_count=Count('questionbank', filter=Q(questionbank__question_sub_type='true_and_false_type')),
        fill_blank_type_count=Count('questionbank', filter=Q(questionbank__question_sub_type='fill_in_the_blank_type')),
        phrases_uploaded=Count('quoteidiomphrase', filter=Q(quoteidiomphrase__type='phrase')),
        idioms_uploaded=Count('quoteidiomphrase', filter=Q(quoteidiomphrase__type='idiom')),
        quotes_uploaded=Count('quoteidiomphrase', filter=Q(quoteidiomphrase__type='quote')),
    ).filter(
        Q(total_input_suggestions__gt=0) | 
        Q(total_questions__gt=0) | 
        Q(phrases_uploaded__gt=0) |
        Q(idioms_uploaded__gt=0) | 
        Q(quotes_uploaded__gt=0)
    ).values_list(
        'user_name', 'email', 'total_input_suggestions', 'total_questions', 
        'simple_type_count', 'list1_type_count', 'list2_type_count',
        'ra_type_count', 'true_false_type_count', 'fill_blank_type_count',
        'phrases_uploaded', 'idioms_uploaded', 'quotes_uploaded',
        named=True
    )
    # Prepare data for charts
    users = [user.user_name for user in users_data]
    emails = [user.email for user in users_data]
    suggestion_counts = [user.total_input_suggestions for user in users_data]
    question_counts = [user.total_questions for user in users_data]
    simple_counts = [user.simple_type_count for user in users_data]
    list1_counts = [user.list1_type_count for user in users_data]
    list2_counts = [user.list2_type_count for user in users_data]
    ra_counts = [user.ra_type_count for user in users_data]
    true_false_counts = [user.true_false_type_count for user in users_data]
    fill_blank_counts = [user.fill_blank_type_count for user in users_data]
    phrase_counts = [user.phrases_uploaded for user in users_data]
    idiom_counts = [user.idioms_uploaded for user in users_data]
    quote_counts = [user.quotes_uploaded for user in users_data]

    context = {
        'users': json.dumps(users),
        'emails': json.dumps(emails),
        'suggestion_counts': json.dumps(suggestion_counts),
        'question_counts': json.dumps(question_counts),
        'simple_counts': json.dumps(simple_counts),
        'list1_counts': json.dumps(list1_counts),
        'list2_counts': json.dumps(list2_counts),
        'ra_counts': json.dumps(ra_counts),
        'true_false_counts': json.dumps(true_false_counts),
        'fill_blank_counts': json.dumps(fill_blank_counts),
        'phrase_counts': json.dumps(phrase_counts),
        'idiom_counts': json.dumps(idiom_counts),
        'quote_counts': json.dumps(quote_counts),
    }

    return render(request, 'question_bank/analytics_dashboard.html', context)


@login_required
def new_dashboard_view(request):
    today = timezone.now().date()
    this_week_start = today - timedelta(days=7)  # Last 7 days
    
    # Date range for earlier week (8–14 days ago)
    earlier_week_start, earlier_week_end = today - timedelta(days=14), today - timedelta(days=7)

    user_is_admin = request.user.user_type == 'admin'

    # Prepare data for this week
    this_week_report_data = []
    earlier_week_report_data = []

    # Fetch users based on user type
    users = User.objects.all() if user_is_admin else User.objects.filter(id=request.user.id)

    # Collect This Week's Reports (last 7 days)
    for user in users:
        total_questions = QuestionBank.objects.filter(created_by=user, created_at__gte=this_week_start).count()
        total_phrases = QuoteIdiomPhrase.objects.filter(created_by=user, created_at__gte=this_week_start).count()
        total_suggestions = InputSuggestion.objects.filter(created_by=user, created_at__gte=this_week_start).count()

        simple_type_count = QuestionBank.objects.filter(created_by=user, created_at__gte=this_week_start, question_sub_type='simple_type').count()
        list_1_type_count = QuestionBank.objects.filter(created_by=user, created_at__gte=this_week_start, question_sub_type='list_type_1').count()
        list_2_type_count = QuestionBank.objects.filter(created_by=user, created_at__gte=this_week_start, question_sub_type='list_type_2').count()
        ra_type_count = QuestionBank.objects.filter(created_by=user, created_at__gte=this_week_start, question_sub_type='r_and_a_type').count()
        true_false_type_count = QuestionBank.objects.filter(created_by=user, created_at__gte=this_week_start, question_sub_type='true_and_false_type').count()
        fill_blank_count = QuestionBank.objects.filter(created_by=user, created_at__gte=this_week_start, question_sub_type='fill_in_the_blank_type').count()

        this_week_report_data.append({
            'user': user,
            'total_questions': total_questions,
            'total_phrases': total_phrases,
            'total_suggestions': total_suggestions,
            'simple_type_count': simple_type_count,
            'list_1_type_count': list_1_type_count,
            'list_2_type_count': list_2_type_count,
            'ra_type_count': ra_type_count,
            'true_false_type_count': true_false_type_count,
            'fill_blank_count': fill_blank_count,
        })

    # Collect Earlier Week's Reports (8–14 days ago)
    for user in users:
        total_questions = QuestionBank.objects.filter(created_by=user, created_at__gte=earlier_week_start, created_at__lt=earlier_week_end).count()
        total_phrases = QuoteIdiomPhrase.objects.filter(created_by=user, created_at__gte=earlier_week_start, created_at__lt=earlier_week_end).count()
        total_suggestions = InputSuggestion.objects.filter(created_by=user, created_at__gte=earlier_week_start, created_at__lt=earlier_week_end).count()

        simple_type_count = QuestionBank.objects.filter(created_by=user, created_at__gte=earlier_week_start, created_at__lt=earlier_week_end, question_sub_type='simple_type').count()
        list_1_type_count = QuestionBank.objects.filter(created_by=user, created_at__gte=earlier_week_start, created_at__lt=earlier_week_end, question_sub_type='list_type_1').count()
        list_2_type_count = QuestionBank.objects.filter(created_by=user, created_at__gte=earlier_week_start, created_at__lt=earlier_week_end, question_sub_type='list_type_2').count()
        ra_type_count = QuestionBank.objects.filter(created_by=user, created_at__gte=earlier_week_start, created_at__lt=earlier_week_end, question_sub_type='r_and_a_type').count()
        true_false_type_count = QuestionBank.objects.filter(created_by=user, created_at__gte=earlier_week_start, created_at__lt=earlier_week_end, question_sub_type='true_and_false_type').count()
        fill_blank_count = QuestionBank.objects.filter(created_by=user, created_at__gte=earlier_week_start, created_at__lt=earlier_week_end, question_sub_type='fill_in_the_blank_type').count()

        earlier_week_report_data.append({
            'user': user,
            'total_questions': total_questions,
            'total_phrases': total_phrases,
            'total_suggestions': total_suggestions,
            'simple_type_count': simple_type_count,
            'list_1_type_count': list_1_type_count,
            'list_2_type_count': list_2_type_count,
            'ra_type_count': ra_type_count,
            'true_false_type_count': true_false_type_count,
            'fill_blank_count': fill_blank_count,
        })

    # Get the last generated dates
    this_week_generated_date = Report.objects.filter(report_type='this_week').order_by('-report_date').first()
    earlier_week_generated_date = Report.objects.filter(report_type='earlier').order_by('-report_date').first()

    context = {
        'this_week_report_data': this_week_report_data,
        'earlier_week_report_data': earlier_week_report_data,
        'this_week_generated_date': this_week_generated_date.report_date if this_week_generated_date else None,
        'earlier_week_generated_date': earlier_week_generated_date.report_date if earlier_week_generated_date else None,
    }
    return render(request, 'question_bank/dashboard.html', context)





@login_required
def generate_this_week_csv(request):
    today = timezone.now().date()
    this_week_start = today - timedelta(days=7)

    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="this_week_report.csv"'
    writer = csv.writer(response)

    writer.writerow(['Sr. No.', 'Email', 'User Name', 'No. of Questions Uploaded', 'No. of Phrases/Idioms/Quotes', 'No. of Input Suggestions',
                     'Simple Type', 'List I Type', 'List II Type', 'R & A Type', 'True & False', 'Fill in the Blank'])

    users = User.objects.all() if request.user.user_type == 'admin' else User.objects.filter(id=request.user.id)

    for idx, user in enumerate(users, 1):
        total_questions = QuestionBank.objects.filter(created_by=user, created_at__gte=this_week_start).count()
        total_phrases = QuoteIdiomPhrase.objects.filter(created_by=user, created_at__gte=this_week_start).count()
        total_suggestions = InputSuggestion.objects.filter(created_by=user, created_at__gte=this_week_start).count()

        simple_type_count = QuestionBank.objects.filter(created_by=user, created_at__gte=this_week_start, question_sub_type='simple_type').count()
        list_1_type_count = QuestionBank.objects.filter(created_by=user, created_at__gte=this_week_start, question_sub_type='list_type_1').count()
        list_2_type_count = QuestionBank.objects.filter(created_by=user, created_at__gte=this_week_start, question_sub_type='list_type_2').count()
        ra_type_count = QuestionBank.objects.filter(created_by=user, created_at__gte=this_week_start, question_sub_type='r_and_a_type').count()
        true_false_type_count = QuestionBank.objects.filter(created_by=user, created_at__gte=this_week_start, question_sub_type='true_and_false_type').count()
        fill_blank_count = QuestionBank.objects.filter(created_by=user, created_at__gte=this_week_start, question_sub_type='fill_in_the_blank_type').count()

        writer.writerow([idx, user.email, f"{user.first_name} {user.last_name}", total_questions, total_phrases, total_suggestions,
                         simple_type_count, list_1_type_count, list_2_type_count, ra_type_count, true_false_type_count, fill_blank_count])

        # Save the report in the Report model
        Report.objects.create(
            report_type='this_week',
            report_date=today,
            created_by=user,
            total_questions=total_questions,
            total_phrases=total_phrases,
            total_suggestions=total_suggestions,
            simple_type_count=simple_type_count,
            list_1_type_count=list_1_type_count,
            list_2_type_count=list_2_type_count,
            ra_type_count=ra_type_count,
            true_false_type_count=true_false_type_count,
            fill_blank_count=fill_blank_count,
        )

    return response



@login_required
def generate_earlier_week_csv(request):
    earlier_week_start, earlier_week_end = timezone.now() - timedelta(days=14), timezone.now() - timedelta(days=7)

    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="earlier_week_report.csv"'
    writer = csv.writer(response)

    writer.writerow(['Sr. No.', 'Email', 'User Name', 'No. of Questions Uploaded', 'No. of Phrases/Idioms/Quotes', 'No. of Input Suggestions',
                     'Simple Type', 'List I Type', 'List II Type', 'R & A Type', 'True & False', 'Fill in the Blank'])

    users = User.objects.all() if request.user.user_type == 'admin' else User.objects.filter(id=request.user.id)

    for idx, user in enumerate(users, 1):
        total_questions = QuestionBank.objects.filter(created_by=user, created_at__gte=earlier_week_start, created_at__lt=earlier_week_end).count()
        total_phrases = QuoteIdiomPhrase.objects.filter(created_by=user, created_at__gte=earlier_week_start, created_at__lt=earlier_week_end).count()
        total_suggestions = InputSuggestion.objects.filter(created_by=user, created_at__gte=earlier_week_start, created_at__lt=earlier_week_end).count()

        simple_type_count = QuestionBank.objects.filter(created_by=user, created_at__gte=earlier_week_start, created_at__lt=earlier_week_end, question_sub_type='simple_type').count()
        list_1_type_count = QuestionBank.objects.filter(created_by=user, created_at__gte=earlier_week_start, created_at__lt=earlier_week_end, question_sub_type='list_type_1').count()
        list_2_type_count = QuestionBank.objects.filter(created_by=user, created_at__gte=earlier_week_start, created_at__lt=earlier_week_end, question_sub_type='list_type_2').count()
        ra_type_count = QuestionBank.objects.filter(created_by=user, created_at__gte=earlier_week_start, created_at__lt=earlier_week_end, question_sub_type='r_and_a_type').count()
        true_false_type_count = QuestionBank.objects.filter(created_by=user, created_at__gte=earlier_week_start, created_at__lt=earlier_week_end, question_sub_type='true_and_false_type').count()
        fill_blank_count = QuestionBank.objects.filter(created_by=user, created_at__gte=earlier_week_start, created_at__lt=earlier_week_end, question_sub_type='fill_in_the_blank_type').count()

        writer.writerow([idx, user.email, f"{user.first_name} {user.last_name}", total_questions, total_phrases, total_suggestions,
                         simple_type_count, list_1_type_count, list_2_type_count, ra_type_count, true_false_type_count, fill_blank_count])

        # Save the report in the Report model
        Report.objects.create(
            report_type='earlier',
            report_date=timezone.now(),
            created_by=user,
            total_questions=total_questions,
            total_phrases=total_phrases,
            total_suggestions=total_suggestions,
            simple_type_count=simple_type_count,
            list_1_type_count=list_1_type_count,
            list_2_type_count=list_2_type_count,
            ra_type_count=ra_type_count,
            true_false_type_count=true_false_type_count,
            fill_blank_count=fill_blank_count,
        )

    return response
    


# ✅ Your standard imports
from django.shortcuts import render, redirect
from django.http import HttpResponse
from .models import (
    Batch, Area, Section, PartName, ChapterName, TopicName, SubTopicName,
    QuestionBank, BatchGeneratedQuestion
)
from django.core.paginator import Paginator
from docx import Document
from docx.shared import Pt
import os
from io import BytesIO
import zipfile
from datetime import datetime
from django.conf import settings
import random

# ✅ ✅ FORM VIEW — fetch pairs always
def generate_test_form(request):
    batches = Batch.objects.all()
    areas = Area.objects.all()
    sections = Section.objects.all()
    parts = PartName.objects.all()
    chapters = ChapterName.objects.all()
    topics = TopicName.objects.all()
    subtopics = SubTopicName.objects.all()

    area_id = request.GET.get("area")
    section_id = request.GET.get("section")
    part_id = request.GET.get("part")
    chapter_id = request.GET.get("chapter")
    topic_id = request.GET.get("topic")
    subtopic_id = request.GET.get("subtopic")
    question_sub_type = request.GET.get("question_sub_type")
    exam_year = request.GET.get("exam_year")
    batch_id = request.GET.get("batch_id")
    page_number = request.GET.get("page")

    en_questions = QuestionBank.objects.filter(language='e')

    if area_id:
        en_questions = en_questions.filter(area_name__area_SI_Code=area_id)
    if section_id:
        en_questions = en_questions.filter(section_name__section_Unit_SI=section_id)
    if part_id:
        en_questions = en_questions.filter(part_name__part_serial=part_id)
    if chapter_id:
        en_questions = en_questions.filter(chapter_name__chapter_number=chapter_id)
    if topic_id:
        en_questions = en_questions.filter(topic_name__topic_SI_number=topic_id)
    if subtopic_id:
        en_questions = en_questions.filter(subtopic_name__sub_topic_SI_Number=subtopic_id)
    if question_sub_type:
        en_questions = en_questions.filter(question_sub_type=question_sub_type)
    if exam_year:
        en_questions = en_questions.filter(exam_year=exam_year)

    en_questions = en_questions.distinct()

    paired_questions = []
    for q_en in en_questions:
        q_hi = QuestionBank.objects.filter(base_question_id=q_en.base_question_id, language='h').first()
        paired_questions.append({'english': q_en, 'hindi': q_hi})

    paginator = Paginator(paired_questions, 100)
    page_obj = paginator.get_page(page_number)

    return render(request, "question_bank/generate_test.html", {
        "batches": batches,
        "questions": page_obj,
        "paginator": paginator,
        "page_obj": page_obj,
        "is_paginated": page_obj.has_other_pages(),
        "areas": areas,
        "sections": sections,
        "parts": parts,
        "chapters": chapters,
        "topics": topics,
        "subtopics": subtopics,
        "selected_batch_id": batch_id,
    })


# -------------------------------
# ✅ MANUAL GENERATE TEST VIEW
# -------------------------------
from io import BytesIO
import zipfile
import os
from docx import Document
from docx.shared import Pt
from datetime import datetime
from django.conf import settings
from django.shortcuts import redirect
from django.http import HttpResponse

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def generate_test(request):
    if request.method == 'POST':
        batch_id = request.POST.get('batch_id')
        if not batch_id:
            return HttpResponse("Error: Batch not selected.", status=400)

        try:
            batch = Batch.objects.get(id=batch_id)
        except Batch.DoesNotExist:
            return HttpResponse("Error: Batch does not exist.", status=404)

        action = request.POST.get('action')

        if action == 'reset':
            BatchGeneratedQuestion.objects.filter(batch=batch).delete()
            return redirect(f'/question-bank/generate-test/?batch_id={batch_id}')

        elif action == 'manual_generate':
            selected_ids = request.POST.getlist('manual_questions')
            selected_questions = QuestionBank.objects.filter(id__in=selected_ids)

            # ✅ Always use pairs
            base_ids = selected_questions.values_list('base_question_id', flat=True).distinct()
            questions = QuestionBank.objects.filter(base_question_id__in=base_ids)

            for q in questions:
                BatchGeneratedQuestion.objects.get_or_create(batch=batch, question=q)

            english_questions = questions.filter(language='e')
            hindi_questions = questions.filter(language='h')

            en_file = generate_word_response(batch, english_questions, lang='EN')
            hi_file = generate_word_response(batch, hindi_questions, lang='HI')

            zip_buffer = BytesIO()
            with zipfile.ZipFile(zip_buffer, "w") as zf:
                zf.writestr(en_file['filename'], en_file['content'].getvalue())
                zf.writestr(hi_file['filename'], hi_file['content'].getvalue())

            response = HttpResponse(zip_buffer.getvalue(), content_type="application/zip")
            response['Content-Disposition'] = f'attachment; filename="Test_Papers_{batch.name}.zip"'
            return response

        else:
            return HttpResponse("Please use manual selection for pair-based export.", status=400)

# -------------------------------
# ✅ FORMAT QUESTION TEXT (NO SOLUTION / CORRECT)
# -------------------------------
from docx.shared import Inches

def format_question_text(q, display_number, document=None):
    is_hindi = q.language == 'h'
    sub_type = q.question_sub_type or 'simple_type'

    text = f"{display_number}. "

    if sub_type == 'list_type_2' and document:
        # Add heading text
        part = q.question_part_first_hi if is_hindi else q.question_part_first
        para = document.add_paragraph()
        para.add_run(text + (part or '')).bold = False

        # Build table for List-I and List-II
        list1 = q.list_1_items_hi if is_hindi else q.list_1_items
        list2 = q.list_2_items_hi if is_hindi else q.list_2_items

        table = document.add_table(rows=max(len(list1), len(list2)) + 1, cols=2)
        table.autofit = True
        remove_table_borders(table)

        # Header row
        table.cell(0, 0).text = 'List-I'
        table.cell(0, 1).text = 'List-II'

        # Fill rows
        for i in range(max(len(list1), len(list2))):
            li = f"{chr(65 + i)}. {list1[i]}" if i < len(list1) else ''
            lii = f"{i + 1}. {list2[i]}" if i < len(list2) else ''
            table.cell(i + 1, 0).text = li
            table.cell(i + 1, 1).text = lii

        # Add options below
        options = [
            (q.answer_option_a_hi if is_hindi else q.answer_option_a),
            (q.answer_option_b_hi if is_hindi else q.answer_option_b),
            (q.answer_option_c_hi if is_hindi else q.answer_option_c),
            (q.answer_option_d_hi if is_hindi else q.answer_option_d)
        ]
        for idx, opt in enumerate(options):
            if opt:
                document.add_paragraph(f"({chr(97 + idx)}) {opt}")

        return ""  # Content already added via document

    else:
        # For all other types
        part = (
            q.question_part_first_hi if is_hindi else
            (q.question_part_first or q.question_part or q.script)
        )
        text += f"{part or ''}\n"

        if sub_type == 'r_and_a_type':
            assertion = q.assertion_hi if is_hindi else q.assertion
            reason = q.reason_hi if is_hindi else q.reason
            text += f"\nAssertion: {assertion or ''}\nReason: {reason or ''}\n"

        a = q.answer_option_a_hi if is_hindi else q.answer_option_a
        b = q.answer_option_b_hi if is_hindi else q.answer_option_b
        c = q.answer_option_c_hi if is_hindi else q.answer_option_c
        d = q.answer_option_d_hi if is_hindi else q.answer_option_d

        if a: text += f"\n(a) {a}"
        if b: text += f"\n(b) {b}"
        if c: text += f"\n(c) {c}"
        if d: text += f"\n(d) {d}"

        return text

# -------------------------------
# ✅ REMOVE TABLE BORDERS HELPER
# -------------------------------
def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        tblBorders.append(border)
    tblPr.append(tblBorders)

from io import BytesIO
import os
from docx import Document
from docx.shared import Pt
from datetime import datetime
from django.conf import settings
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ✅ Utility to remove table borders
def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        tblBorders.append(border)
    tblPr.append(tblBorders)

# ✅ ✅ FINAL GENERATOR — calls your updated format_question_text
def generate_word_response(batch, questions, lang='EN'):
    template_path = os.path.join(
        settings.BASE_DIR, 
        'static', 'assets', 'templates', 'lecture_notes_template.docx'
    )
    document = Document(template_path)

    document.add_heading(f"Hajela’s IAS - Test Paper ({batch.name}) [{lang}]", level=1)

    table = document.add_table(rows=0, cols=2)
    table.autofit = True
    remove_table_borders(table)

    for i in range(0, len(questions), 2):
        row = table.add_row().cells

        q1 = questions[i]
        build_question_in_cell(q1, i + 1, row[0], lang)

        if i + 1 < len(questions):
            q2 = questions[i + 1]
            build_question_in_cell(q2, i + 2, row[1], lang)

        for cell in row:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(11)

    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    subject_code = "GS4"
    exam_type = "MPPSC"
    stage = "PRE"
    set_number = f"SET{datetime.now().strftime('%d%m%H%M%S')}"
    unique_filename = f"{subject_code}_{exam_type}_{stage}_{lang}_{set_number}.docx"

    return {
        "filename": unique_filename,
        "content": file_stream
    }

def build_question_in_cell(q, display_number, cell, lang):
    is_hindi = q.language == 'h'
    sub_type = q.question_sub_type or 'simple_type'

    para = cell.add_paragraph(f"{display_number}. ")

    if sub_type == 'list_type_2':
        part = q.question_part_first_hi if is_hindi else q.question_part_first
        para.add_run(part or "")

        list1 = q.list_1_items_hi if is_hindi else q.list_1_items
        list2 = q.list_2_items_hi if is_hindi else q.list_2_items

        nested_table = cell.add_table(rows=len(list1)+1, cols=2)
        remove_table_borders(nested_table)

        hdr = nested_table.rows[0].cells
        hdr[0].text = "List-I"
        hdr[1].text = "List-II"

        for idx in range(len(list1)):
            row_cells = nested_table.rows[idx+1].cells
            row_cells[0].text = f"{chr(65+idx)}. {list1[idx]}"
            if idx < len(list2):
                row_cells[1].text = f"{idx+1}. {list2[idx]}"

        options = [
            q.answer_option_a_hi if is_hindi else q.answer_option_a,
            q.answer_option_b_hi if is_hindi else q.answer_option_b,
            q.answer_option_c_hi if is_hindi else q.answer_option_c,
            q.answer_option_d_hi if is_hindi else q.answer_option_d,
        ]

        for idx, option in enumerate(options):
            if option:
                cell.add_paragraph(f"({chr(97 + idx)}) {option}")

    else:
        part = q.question_part_hi if is_hindi else (q.question_part or q.script)
        para.add_run(part or "")

        options = [
            q.answer_option_a_hi if is_hindi else q.answer_option_a,
            q.answer_option_b_hi if is_hindi else q.answer_option_b,
            q.answer_option_c_hi if is_hindi else q.answer_option_c,
            q.answer_option_d_hi if is_hindi else q.answer_option_d,
        ]

        for idx, option in enumerate(options):
            if option:
                cell.add_paragraph(f"({chr(97 + idx)}) {option}")


# Image Generation Chat View
from django.shortcuts import render, redirect
from .openai_image import generate_image  # This should return a list of GeneratedImage objects
from datetime import datetime

from datetime import datetime
from django.shortcuts import render, redirect
from .openai_image import generate_image
from datetime import datetime

def generate_image_view(request):
    session_id = request.session.get('active_chat_session_id')
    selected_session_id = request.GET.get('session_id')

    # Initialize session storage
    if 'chat_sessions' not in request.session:
        request.session['chat_sessions'] = {}
    if 'chat_history_sessions' not in request.session:
        request.session['chat_history_sessions'] = {}

    chat_sessions = request.session['chat_sessions']
    chat_history_sessions = request.session['chat_history_sessions']

    # Start a new session if not exists
    if not session_id:
        session_id = str(datetime.now().timestamp())
        request.session['active_chat_session_id'] = session_id
        chat_sessions[session_id] = []

    # Handle POST (generate new image and store to current session)
    if request.method == 'POST':
        user_prompt = request.POST.get('prompt')
        if user_prompt:
            current_chat = chat_sessions.get(session_id, [])
            last_prompt = current_chat[-1]['prompt'] if current_chat else ""
            final_prompt = f"{last_prompt}. {user_prompt}" if last_prompt else user_prompt

            image_objs = generate_image(final_prompt)

            for img in image_objs:
                current_chat.append({
                    'prompt': final_prompt,
                    'image_url': img.image.url,
                    'timestamp': str(datetime.now())
                })

            chat_sessions[session_id] = current_chat

            # Save session info if first time
            if session_id not in chat_history_sessions:
                chat_history_sessions[session_id] = {
                    'id': session_id,
                    'first_prompt': final_prompt,
                    'created_at': datetime.now().isoformat()
                }

            request.session.modified = True
            return redirect('generate_image')

    # ✅ Load history: use selected session_id if present
    chat_history = chat_sessions.get(selected_session_id, []) if selected_session_id else chat_sessions.get(session_id, [])

    return render(request, 'question_bank/generate_image.html', {
        'chat_history': chat_history,
        'session_list': list(chat_history_sessions.values())
    })




def clear_image_chat(request):
    if 'chat_history' in request.session:
        request.session['chat_history'] = []
    return redirect('generate_image')


def new_image_chat(request):
    request.session.pop('chat_history', None)
    request.session.pop('active_chat_session_id', None)
    return redirect('generate_image')



from django.shortcuts import render
from .models import Area

def hierarchy_view(request):
    areas = Area.objects.prefetch_related(
        'sections__parts__chapters__topics__sub_topics'
    ).all()
    return render(request, 'question_bank/hierarchy_page.html', {'areas': areas})



import os
import re
import tempfile
from django.shortcuts import render
from django.http import HttpResponse
from decouple import config
from openai import OpenAI
from docx import Document
from django.conf import settings

client = OpenAI(api_key=config("OPENAI_API_KEY"))

def slugify_filename(text):
    """
    Replace invalid filename characters with underscores.
    This makes your filename safe for Windows and cross-platform.
    """
    text = re.sub(r'[\\/*?:"<>|]', "", text)
    text = text.strip().replace(' ', '_')
    return text

def generate_notes(request):
    if request.method == "POST":
        topic = request.POST.get("topic")
        safe_topic = slugify_filename(topic)

        # ✅ Structured UPSC Current Affairs prompt
        prompt = f"""
You are an expert UPSC mentor. Give me comprehensive UPSC notes for the topic: "{topic}"

Your output must include:
1. Short introduction (background + context)
2. Key facts, data, agreements, or dates (bullet points)
3. Strategic significance for India (national security, economy, diplomacy, GS2 linkages)
4. Stakeholders involved (India + other countries/institutions)
5. Possible benefits and opportunities for India
6. Challenges, concerns or criticisms
7. Expert views, reports, or quotes to enrich analysis
8. Diagrams, maps, or flowcharts if relevant (describe)
9. Answer framework for UPSC Mains: Intro – Body – Conclusion
10. 3-5 potential UPSC Prelims MCQs with correct answers
11. 3-5 possible UPSC Mains questions (GS2/GS3/Essay)

Use clear bullet points, formal language, and keep it crisp and exam-oriented.
Provide headings in both Hindi and English where suitable.
"""

        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}]
        )
        notes_text = response.choices[0].message.content.strip()

        # ✅ Load Word template
        template_path = os.path.join(
            settings.BASE_DIR,
            'static',
            'assets',
            'templates',
            'lecture_notes_template.docx'
        )
        if not os.path.exists(template_path):
            return HttpResponse(f"Template not found: {template_path}", status=500)

        doc = Document(template_path)
        doc.add_heading(topic, level=1)

        # ✅ Add structured notes content
        for line in notes_text.split('\n'):
            line = line.strip()
            if line and (
                (line[0].isdigit() and line[1] == '.') or
                line.endswith(":") or
                line.startswith(("✅", "🔸", "•", "-", "*"))
            ):
                doc.add_paragraph(line, style='Heading 2')
            elif line:
                doc.add_paragraph(line)

        # ✅ Use safe filename
        output_path = os.path.join(tempfile.gettempdir(), f"{safe_topic}_UPSC_Notes.docx")
        doc.save(output_path)
        with open(output_path, 'rb') as f:
            response = HttpResponse(
                f.read(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'attachment; filename="{safe_topic}_UPSC_Notes.docx"'
            return response

    return render(request, 'question_bank/generate_notes.html')



# Database Views to show all the data at one place 


def database(request):

    return render(request,'question_bank/database.html')




from django.contrib.auth.decorators import login_required
from django.shortcuts import render, get_object_or_404
from collections import defaultdict
from .models import QuestionBank, LectureNote, InputSuggestion, SubTopicName

@login_required
def content_by_subtopic(request):
    sub_topic_short_code = request.GET.get("sub_topic_short_Code", None)

    pyqs_grouped = []
    moqs_grouped = []
    osqs_grouped = []
    lecture_notes = []
    input_suggestions = []

    if sub_topic_short_code:
        subtopic = get_object_or_404(SubTopicName, sub_topic_short_Code=sub_topic_short_code)

        # Helper function to group English & Hindi
        def group_queryset(qs):
            grouped = defaultdict(dict)
            for q in qs:
                if q.language == "e":
                    grouped[q.base_question_id]["english"] = q
                elif q.language == "h":
                    grouped[q.base_question_id]["hindi"] = q
            return list(grouped.items())  # ✅ return as list of tuples

        # --- PYQs ---
        pyqs = QuestionBank.objects.filter(
            type_of_question="pyq", subtopic_name=subtopic
        ).order_by("-created_at")
        pyqs_grouped = group_queryset(pyqs)

        # --- MOQs ---
        moqs = QuestionBank.objects.filter(
            type_of_question="moq", subtopic_name=subtopic
        ).order_by("-created_at")
        moqs_grouped = group_queryset(moqs)

        # --- OSQs ---
        osqs = QuestionBank.objects.filter(
            type_of_question="osq", subtopic_name=subtopic
        ).order_by("-created_at")
        osqs_grouped = group_queryset(osqs)

        # --- Notes & Suggestions ---
        lecture_notes = LectureNote.objects.filter(subtopic=subtopic)
        input_suggestions = InputSuggestion.objects.filter(subtopic_name=subtopic)

    context = {
        "sub_topic_short_code": sub_topic_short_code,
        "pyqs_grouped": pyqs_grouped,
        "moqs_grouped": moqs_grouped,
        "osqs_grouped": osqs_grouped,
        "lecture_notes": lecture_notes,
        "input_suggestions": input_suggestions,
    }
    return render(request, "question_bank/content_by_subtopic.html", context)






# from django.http import JsonResponse
# from django.shortcuts import get_object_or_404
# from .models import QuestionBank

# def fetch_tagged_subtopics(request, question_id):
#     """
#     Fetch manually tagged subtopics, AI subtopics, and keywords for a specific question.
#     """
#     question = get_object_or_404(QuestionBank, id=question_id)

#     # Manual Subtopics (if you have ManyToMany field for manual tagging)
#     manual_subtopics = []
#     if hasattr(question, 'manual_subtopics'):
#         manual_subtopics = [st.name for st in question.manual_subtopics.all()]

#     # AI Subtopics
#     ai_subtopics = []
#     if hasattr(question, 'ai_subtopics'):
#         ai_subtopics = [st.name for st in question.ai_subtopics.all()]

#     # Keywords (English + Hindi)
#     keywords_en = []
#     keywords_hi = []
#     if hasattr(question, 'keywords'):
#         keywords_en = [kw.name_en for kw in question.keywords.all() if hasattr(kw, 'name_en')]
#         keywords_hi = [kw.name_hi for kw in question.keywords.all() if hasattr(kw, 'name_hi')]

#     data = {
#         "manual_subtopics": manual_subtopics,
#         "ai_subtopics": ai_subtopics,
#         "keywords_en": keywords_en,
#         "keywords_hi": keywords_hi
#     }

#     return JsonResponse(data)
